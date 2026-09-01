import streamlit as st
import pandas as pd
from streamlit_gsheets import GSheetsConnection
from google import genai
from google.genai import types
import docx
from docx.text.paragraph import Paragraph
from docx.shared import Pt
import pypdf
import json
import re
import os
import copy
import time
import subprocess
import smtplib
import hashlib
from email.message import EmailMessage
from pdf2docx import Converter
from copy import deepcopy
from datetime import date

# ====================================================================
# --- STREAMLIT UI & PASSWORD LOGIC ---
# ====================================================================

st.set_page_config(page_title="Resume Formatter", layout="wide")

# --- CUSTOM CSS FOR INPUT BOXES ---
st.markdown(
    """
    <style>
    /* Add a visible outline and subtle background shading to all inputs, text areas, and dropdowns */
    div[data-baseweb="input"], 
    div[data-baseweb="base-input"], 
    div[data-baseweb="select"] {
        border: 1px solid #7c7c8c !important; /* Gray border */
        background-color: rgba(128, 128, 128, 0.1) !important; /* Slight transparent shading */
        border-radius: 6px !important; /* Smooth rounded corners */
    }
    </style>
    """, 
    unsafe_allow_html=True
)
# ----------------------------------

def check_password():
    """Returns `True` if the user entered the correct password."""
    def password_entered():
        if st.session_state["password"] == st.secrets["APP_PASSWORD"]:
            st.session_state["password_correct"] = True
            del st.session_state["password"]  
        else:
            st.session_state["password_correct"] = False

    if st.session_state.get("password_correct", False):
        return True

    st.title("🔒 Access Restricted")
    st.text_input(
        "Please enter the team password to use this tool:", 
        type="password", 
        on_change=password_entered, 
        key="password"
    )
    
    if "password_correct" in st.session_state:
        st.error("❌ Incorrect password.")
        
    return False

# 🛑 STOP THE APP HERE IF THEY DON'T HAVE THE PASSWORD
if not check_password():
    st.stop()


# ====================================================================
# --- UNIVERSAL HELPER FUNCTIONS (Shared by all clients) ---
# ====================================================================

def clean_bullets(bullet_list):
    if isinstance(bullet_list, str):
        bullet_list = bullet_list.split('\n')
    cleaned = []
    for line in bullet_list:
        line = re.sub(r'^[\s\-\•\d\.\*]+', '', line).strip()
        if line:
            cleaned.append(line)
    return "\n".join(cleaned)

def clean_school(name):
    for delimiter in [' - ', ' – ', ' , ', ', ']:
        if delimiter in name:
            name = name.split(delimiter)[0]
    return name.strip()

def standardize_dates(date_str):
    """
    Normalize only date detail that is actually present in the source.

    IMPORTANT:
    - Never invent a month when the source provides only a year.
    - "2018 - 2020" must remain year-only.
    - "2018 - Present" must remain year-only on the start side.
    - Month/year inputs may still be standardized normally.
    """
    month_map = {
        "01": "Jan", "02": "Feb", "03": "Mar", "04": "Apr", "05": "May", "06": "Jun",
        "07": "Jul", "08": "Aug", "09": "Sep", "10": "Oct", "11": "Nov", "12": "Dec",
        "1": "Jan", "2": "Feb", "3": "Mar", "4": "Apr", "5": "May", "6": "Jun",
        "7": "Jul", "8": "Aug", "9": "Sep",
        "January": "Jan", "February": "Feb", "March": "Mar", "April": "Apr",
        "May": "May", "June": "Jun", "July": "Jul", "August": "Aug",
        "September": "Sep", "Sept": "Sep", "October": "Oct", "November": "Nov", "December": "Dec"
    }
    date_str = re.sub(r"['’](\d{2})", lambda m: ("20" if int(m.group(1)) < 50 else "19") + m.group(1), date_str)
    date_str = re.sub(r'(\d{1,2})/(\d{4})', lambda m: f"{month_map.get(m.group(1).zfill(2), m.group(1))} {m.group(2)}", date_str)
    for full, short in month_map.items():
        if not full.isdigit():
            date_str = re.sub(full, short, date_str, flags=re.IGNORECASE)
    return date_str

def format_peraton_dates(date_str):
    """
    Format Peraton dates without inventing missing month detail.

    Examples:
    - "10/2019 - 07/2024" -> "10/2019 to 07/2024"
    - "2019 - 2024"       -> "2019 to 2024"
    - "2019 - Present"    -> "2019 to Present"

    If the original resume gives only a year, keep only the year.
    """
    if not date_str:
        return ""
    # Swap common hyphens/dashes to the required " to "
    date_str = re.sub(r'\s*[-–—]\s*', ' to ', date_str)
    # Ensure "Present" or "Current" is properly formatted
    date_str = re.sub(r'(?i)present', 'Present', date_str)
    date_str = re.sub(r'(?i)current', 'Present', date_str)
    return date_str.strip()

def extract_text(file_path):
    text_parts = []
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.pdf':
            pdf = pypdf.PdfReader(file_path)
            for page in pdf.pages:
                content = page.extract_text()
                if content:
                    text_parts.append(content)
        elif ext in ['.docx', '.doc']:
            doc = docx.Document(file_path)
            text_box_found = False
            for node in doc.element.xpath('//w:txbxContent//w:t'):
                if node.text and node.text.strip():
                    text_parts.append(node.text.strip())
                    text_box_found = True
            if text_box_found:
                text_parts.append("--- END OF TEXT BOXES ---")
            for section in doc.sections:
                headers = [section.header, section.first_page_header, section.even_page_header]
                for hdr in headers:
                    if hdr: 
                        for p in hdr.paragraphs:
                            if p.text.strip():
                                text_parts.append(p.text)
                        for table in hdr.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    if cell.text.strip():
                                        text_parts.append(cell.text)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
    except:
        with open(file_path, 'rb') as f:
            text_parts.append(f.read().decode('utf-8', errors='ignore'))
    return "\n".join(text_parts)

def repair_and_load_json(raw_text):
    text = raw_text.strip()
    text = re.sub(r'^```[a-zA-Z]*\n|\n```$', '', text, flags=re.MULTILINE).strip()
    match = re.search(r'\{.*\}', text, re.DOTALL)
    if match:
        text = match.group()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        text = re.sub(r'"\s*\n\s*"', '",\n"', text)
        text = re.sub(r'\}\s*\n\s*\{', '},\n{', text)
        text = re.sub(r'\]\s*\n\s*"', '],\n"', text)
        try:
            return json.loads(text)
        except Exception as e:
            st.warning(f"⚠️ Minor Error: AI formatting glitch on this resume. Bypassing crash... ({e})")
            return {}

def replace_tag_safely(p, tag, value, unbold=False, force_bold=False, font_name=None):
    if tag.lower() not in p.text.lower():
        return False
    replaced = False
    for run in p.runs:
        if tag.lower() in run.text.lower():
            pattern = re.compile(re.escape(tag), re.IGNORECASE)
            run.text = pattern.sub(str(value), run.text)
            if unbold:
                run.font.bold = False
            if force_bold:
                run.font.bold = True
            if font_name:
                run.font.name = font_name
            replaced = True
    if not replaced:
        full_text = p.text
        pattern = re.compile(re.escape(tag), re.IGNORECASE)
        new_text = pattern.sub(str(value), full_text)
        if p.runs:
            p.runs[0].text = new_text
            if unbold:
                p.runs[0].font.bold = False
            if force_bold:
                p.runs[0].font.bold = True
            if font_name:
                p.runs[0].font.name = font_name
            for i in range(1, len(p.runs)):
                p.runs[i].text = ""
    return True

def insert_bullet_line_after(paragraph, text, font_name=None):
    new_p_xml = copy.deepcopy(paragraph._p)
    paragraph._p.addnext(new_p_xml)
    new_p = Paragraph(new_p_xml, paragraph._parent)
    for run in new_p.runs:
        run._element.getparent().remove(run._element)
    if paragraph.runs:
        new_run = new_p.add_run(text)
        ref_font = paragraph.runs[0].font
        new_run.font.bold = ref_font.bold
        new_run.font.italic = ref_font.italic
        if ref_font.color and ref_font.color.rgb:
            new_run.font.color.rgb = ref_font.color.rgb
        new_run.font.size = ref_font.size
        if font_name:
            new_run.font.name = font_name
        else:
            new_run.font.name = ref_font.name
    else:
        new_p.add_run(text)
    return new_p

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

def collapse_extra_blank_paragraphs(doc):
    """
    Collapse consecutive empty BODY paragraphs to a single blank paragraph.

    This is intentionally conservative:
    - Only operates on top-level document-body paragraphs.
    - Does NOT touch paragraphs inside tables.
    - Does NOT remove page-break or section-break paragraphs.
    - Preserves one intentional blank paragraph between sections.
    """

    body = doc.element.body
    previous_was_blank = False

    for element in list(body):

        # A table or other non-paragraph element breaks the blank-paragraph run.
        if not element.tag.endswith("}p"):
            previous_was_blank = False
            continue

        # ------------------------------------------------------------
        # DO NOT TOUCH STRUCTURAL PARAGRAPHS
        # ------------------------------------------------------------

        # Preserve section breaks.
        if element.xpath(".//w:sectPr"):
            previous_was_blank = False
            continue

        # Preserve explicit page/column/line breaks.
        if element.xpath(".//w:br"):
            previous_was_blank = False
            continue

        # Preserve drawings/images/shapes.
        if (
            element.xpath(".//w:drawing")
            or element.xpath(".//w:pict")
            or element.xpath(".//w:object")
        ):
            previous_was_blank = False
            continue

        # ------------------------------------------------------------
        # DETERMINE WHETHER PARAGRAPH IS VISUALLY EMPTY
        # ------------------------------------------------------------

        text_nodes = element.xpath(".//w:t")
        text = "".join(
            node.text or ""
            for node in text_nodes
        ).strip()

        is_blank = not text

        # ------------------------------------------------------------
        # COLLAPSE ONLY CONSECUTIVE BLANKS
        # ------------------------------------------------------------

        if is_blank:
            if previous_was_blank:
                body.remove(element)
            else:
                previous_was_blank = True
        else:
            previous_was_blank = False

def ensure_experience_spacing(doc, mapping):
    """
    Ensure one visible blank paragraph between consecutive work-history entries.

    This specifically protects dynamically generated jobs beyond the original
    seven template slots, including abbreviated older roles that have
    Company/Title/Dates but no bullets. Existing spacing is preserved; a new
    blank paragraph is added only when one is missing.
    """
    paragraphs = list(doc.paragraphs)
    company_pattern = re.compile(
        r"\{\{company(\d+)\}\}",
        flags=re.IGNORECASE,
    )

    role_locations = []

    for idx, paragraph in enumerate(paragraphs):
        match = company_pattern.search(paragraph.text or "")
        if not match:
            continue

        role_num = int(match.group(1))
        company_value = str(
            mapping.get(f"Company{role_num}", "") or ""
        ).strip()

        if company_value:
            role_locations.append((role_num, idx))

    if len(role_locations) < 2:
        return

    # Bottom-up so inserting a paragraph does not invalidate earlier indexes.
    for _, idx in reversed(role_locations[1:]):
        paragraphs = list(doc.paragraphs)

        if idx <= 0 or idx >= len(paragraphs):
            continue

        current_p = paragraphs[idx]
        previous_p = paragraphs[idx - 1]

        # Do nothing when the template/output already has a separator.
        if not previous_p.text.strip():
            continue

        blank = current_p.insert_paragraph_before("")

        try:
            blank.style = doc.styles["Normal"]
        except Exception:
            pass

        blank.paragraph_format.space_before = Pt(0)
        blank.paragraph_format.space_after = Pt(0)


def expand_experience_placeholders(doc, mapping):
    """
    Dynamically extend a template's work-history placeholder block when the
    candidate has more roles than the template was originally built to hold.

    Existing templates can stay exactly as they are (for example, through
    {{Company7}} / {{Dates7}} / {{Title7}} / {{Bullets7}}). If mapping contains
    Company8, Company9, etc., this function clones the final existing experience
    block and renumbers the placeholders before the normal replacement logic runs.

    This is deliberately template-driven rather than capped at an arbitrary
    number such as 20.
    """

    # Determine how many experience entries the mapping actually contains.
    mapped_indices = []
    for key in mapping.keys():
        match = re.fullmatch(
            r"(?i)(?:company|dates|title|bullets|environment|responsible)(\d+)",
            str(key),
        )
        if match:
            mapped_indices.append(int(match.group(1)))

    if not mapped_indices:
        return

    required_max = max(mapped_indices)

    # Locate the highest-numbered complete experience block already present
    # in the template body. Current client templates use paragraph-based
    # Company/Dates -> Title -> Bullets blocks.
    paragraphs = list(doc.paragraphs)
    placeholder_pattern = re.compile(
        r"\{\{(company|dates|title|bullets|environment|responsible)(\d+)\}\}",
        flags=re.IGNORECASE,
    )

    locations = {}
    for idx, paragraph in enumerate(paragraphs):
        for kind, number in placeholder_pattern.findall(paragraph.text or ""):
            number = int(number)
            locations.setdefault(number, {}).setdefault(kind.lower(), []).append(idx)

    complete_indices = [
        number
        for number, parts in locations.items()
        if "company" in parts and "title" in parts and "bullets" in parts
    ]

    if not complete_indices:
        return

    template_max = max(complete_indices)

    if required_max <= template_max:
        return

    parts = locations[template_max]
    block_start = min(
        parts["company"] + parts.get("dates", []) + parts["title"] + parts["bullets"]
    )
    block_end = max(
        parts["company"] + parts.get("dates", []) + parts["title"] + parts["bullets"]
    )

    # Include any Environment placeholder that belongs inside the same final block.
    if parts.get("environment"):
        block_end = max(block_end, max(parts["environment"]))

    # Preserve the template's inter-role spacing. Peraton, for example,
    # uses one blank Normal paragraph between experience blocks, while
    # the Fannie-style templates do not.
    if (
        block_start > 0
        and not paragraphs[block_start - 1].text.strip()
    ):
        block_start -= 1

    source_block = paragraphs[block_start:block_end + 1]
    if not source_block:
        return

    insert_after = source_block[-1]._p

    # Clone the final template block once for each additional candidate role.
    for new_index in range(template_max + 1, required_max + 1):
        cloned_elements = []

        for source_paragraph in source_block:
            new_p = copy.deepcopy(source_paragraph._p)

            # Renumber every placeholder in the cloned block that belongs to
            # the original final template index.
            text_nodes = new_p.xpath(".//w:t")
            combined_text = "".join(
                node.text or ""
                for node in text_nodes
            )

            # Word frequently splits one placeholder across multiple runs
            # (for example "{{" + "company" + "7" + "}}").  Once we confirm
            # this cloned paragraph belongs to the final template role, safely
            # renumber the placeholder fragments across those runs.
            if re.search(
                rf"(?i)\{{\{{(?:company|dates|title|bullets|environment|responsible){template_max}\}}\}}",
                combined_text,
            ):
                for text_node in text_nodes:
                    if not text_node.text:
                        continue

                    text_node.text = re.sub(
                        rf"(?i)(company|dates|title|bullets|environment|responsible){template_max}",
                        lambda m: f"{m.group(1)}{new_index}",
                        text_node.text,
                    )

                    if text_node.text.strip() == str(template_max):
                        text_node.text = text_node.text.replace(
                            str(template_max),
                            str(new_index),
                        )

            insert_after.addnext(new_p)
            insert_after = new_p
            cloned_elements.append(new_p)


def process_word_doc(temp_path, mapping, out_path):
    doc = docx.Document(temp_path)
    expand_experience_placeholders(doc, mapping)
    ensure_experience_spacing(doc, mapping)
    
    # Check if this is the Peraton tool
    is_peraton = "CERTIFICATION1" in mapping
    target_font = 'Calibri' if is_peraton else None
    env_font = 'Calibri' if is_peraton else 'Times New Roman'

    # --- TWEAK 1: Remove "Certifications" Table & Header (All Clients) ---
    has_certifications = True
    if is_peraton:
        if not mapping.get("CERTIFICATION1"):
            has_certifications = False
    else:
        if not mapping.get("Certifications") or not str(mapping.get("Certifications")).strip():
            has_certifications = False

    if not has_certifications:
        tables_to_delete = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Made case-insensitive to ensure it catches the tag even if typed differently in template
                    if "{{certifications}}" in cell.text.lower() or "{{certification1}}" in cell.text.lower():
                        if table not in tables_to_delete:
                            tables_to_delete.append(table)
        for table in tables_to_delete:
            table._element.getparent().remove(table._element)
            
        paras = list(doc.paragraphs)
        for i, p in enumerate(paras):
            # Made case-insensitive to ensure the header text is caught
            if p.text.strip().lower() == "certifications":
                delete_paragraph(p)
                # Prevent a double-space gap by deleting the empty line above it if it exists
                if i > 0 and not paras[i-1].text.strip():
                    try:
                        delete_paragraph(paras[i-1])
                    except:
                        pass

    # --- TWEAK 2: Remove Q&A Section for ALL clients if empty (Including Fannie Mae) ---
    has_qa = any(mapping.get(k) and str(mapping.get(k)).strip() for k in ["Q1", "A1", "Q2", "A2", "Q3", "A3", "Q4", "A4", "Q5", "A5"])

    if not has_qa:
        paras = list(doc.paragraphs)
        for i, p in enumerate(paras):
            if "SUPPLIER TECHNICAL INTERVIEW RESULTS" in p.text.upper():
                delete_paragraph(p)
                # Prevent a double-space gap by deleting the empty line above it if it exists
                if i > 0 and not paras[i-1].text.strip():
                    try:
                        delete_paragraph(paras[i-1])
                    except:
                        pass
        
        tables_to_delete = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # Made case-insensitive for safety
                    if "{{q1}}" in cell.text.lower():
                        if table not in tables_to_delete:
                            tables_to_delete.append(table)
        for table in tables_to_delete:
            table._element.getparent().remove(table._element)
                
    # --- PERATON TWEAK 2: Space before Relevant Professional Experience ---
    if is_peraton:
        for p in list(doc.paragraphs):
            if "Relevant Professional Experience" in p.text:
                p.insert_paragraph_before("")
                break # Only do it once
            
    has_any_education = bool(mapping.get("School1") or mapping.get("School2") or mapping.get("School3"))
    for table in doc.tables:
        rows_to_delete = []
        for row in table.rows:
            row_text = "".join(cell.text for cell in row.cells).strip()
            for i in range(1, 4):
                tag = f"{{{{School{i}}}}}"
                if tag.lower() in row_text.lower() and not mapping.get(f"School{i}"):
                    if has_any_education and row not in rows_to_delete:
                        rows_to_delete.append(row)
        for row in rows_to_delete:
            table._tbl.remove(row._tr)
            
    paras_to_remove = []
    # Updated kill_keys with 'Certification', 'School', and 'Responsible'
    kill_keys = ['Company', 'Title', 'Bullets', 'Dates', 'Q1', 'Q2', 'Q3', 'Q4', 'Q5', 'A1', 'A2', 'A3', 'A4', 'A5', 'Summary', 'Skill', 'Years', 'Certification', 'School', 'Responsible', 'BillRate', 'WorkStatus']
    for p in list(doc.paragraphs):
        for key, value in mapping.items():
            tag = f"{{{{{key}}}}}"
            if tag.lower() in p.text.lower():
                if not value or not str(value).strip():
                    if any(k.lower() in key.lower() for k in kill_keys):
                        if p not in paras_to_remove:
                            paras_to_remove.append(p)
                    else:
                        replace_tag_safely(p, tag, "", font_name=target_font)
                else:
                    if "Bullets" in key:
                        for run in p.runs:
                            if '\n' in run.text:
                                run.text = run.text.replace('\n', '')
                        lines = str(value).split('\n')
                        replace_tag_safely(p, tag, lines[0], font_name=target_font)
                        curr_p = p
                        for line in lines[1:]:
                            curr_p = insert_bullet_line_after(curr_p, line, font_name=target_font)
                        num_match = re.search(r'\d+', key)
                        if num_match:
                            env_val = mapping.get(f"Environment{num_match.group()}")
                            if env_val and str(env_val).strip():
                                env_p = curr_p.insert_paragraph_before("")
                                curr_p._p.addnext(env_p._p)
                                try:
                                    env_p.style = doc.styles['Normal']
                                except:
                                    pass
                                env_p.paragraph_format.left_indent = Pt(0)
                                env_p.paragraph_format.space_after = Pt(0)
                                clean_env = re.sub(r'^Environment\s*:\s*', '', str(env_val).strip(), flags=re.IGNORECASE)
                                b_run = env_p.add_run("Environment: ")
                                b_run.bold = True
                                b_run.font.name = env_font
                                b_run.font.size = Pt(12)
                                n_run = env_p.add_run(clean_env)
                                n_run.bold = False
                                n_run.font.name = env_font
                                n_run.font.size = Pt(12)
                                curr_p = env_p
                        
                        # --- PERATON TWEAK 3: Disable auto-spacer for Peraton ---
                        if not is_peraton:
                            spacer = curr_p.insert_paragraph_before("")
                            curr_p._p.addnext(spacer._p)
                            try:
                                spacer.style = doc.styles['Normal']
                            except:
                                pass
                            spacer.paragraph_format.space_after = Pt(0)
                            spacer.paragraph_format.space_before = Pt(0)
                    else:
                        is_answer = key in ['A1', 'A2', 'A3', 'A4', 'A5']
                        is_question = key in ['Q1', 'Q2', 'Q3', 'Q4', 'Q5']
                        replace_tag_safely(p, tag, str(value).strip(), unbold=is_answer, force_bold=is_question, font_name=target_font)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in mapping.items():
                        tag = f"{{{{{key}}}}}"
                        if tag.lower() in p.text.lower():
                            is_answer = key in ['A1', 'A2', 'A3', 'A4', 'A5']
                            is_question = key in ['Q1', 'Q2', 'Q3', 'Q4', 'Q5']
                            replace_tag_safely(p, tag, str(value) if value else "", unbold=is_answer, force_bold=is_question, font_name=target_font)
    for p in paras_to_remove:
        try:
            delete_paragraph(p)
        except:
            pass

    # Collapse extra blank lines created when optional sections/tables were removed.
    collapse_extra_blank_paragraphs(doc)
    
    doc.save(out_path)
    return out_path


# ====================================================================
# --- 🔴 CLIENT APP 1: FANNIE MAE 🔴 ---
# ====================================================================
def fannie_mae_app():
    TEMPLATE_FILENAME = "Template.docx"
    st.title("Fannie Mae Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="fm_loc")
        Remote_or_Onsite = st.selectbox("Remote or Onsite", ["Remote", "Onsite"], index=1, key="fm_rem")
    with col2:
        Former_FM = st.selectbox("Former FM", ["Y - Per CRC, this candidate is eligible for rehire", "N"], index=1, key="fm_form")
        LinkedIn_GitHub_Portfolio_Link = st.text_input("LinkedIn/GitHub/Portfolio Link", key="fm_link")

    st.header("🎤 Supplier Technical Interview Results")
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="fm_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="fm_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="fm_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="fm_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="fm_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="fm_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="fm_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="fm_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="fm_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="fm_a5")

    st.header("📝 Job Description & Notes")

    # --- NEW GOOGLE SHEET LOGIC ---
    # IMPORTANT: Replace this placeholder with your actual Google Sheet link. 
    # The sheet's share settings MUST be set to "Anyone with the link can view".
    SHEET_URL = "https://docs.google.com/spreadsheets/d/1swKaDhGdlmO0ILv7tngkRCbsK-7jF89cdCpf5bypzwM/edit?usp=sharing"

    @st.cache_data(ttl=600)
    def load_reqs(url):
        try:
            # Converts standard Google Sheet URL to a direct CSV download link
            if "/edit" in url:
                csv_url = url.split("/edit")[0] + "/export?format=csv"
            else:
                csv_url = url
            df = pd.read_csv(csv_url)
            return df.fillna("")
        except Exception:
            return pd.DataFrame()

    df_reqs = load_reqs(SHEET_URL)
    options = ["None"]
    req_mapping = {}

    if not df_reqs.empty:
        for idx, row in df_reqs.iterrows():
            req_id = str(row.get("ID", "")).replace(".0", "")
            vms_id = str(row.get("VMS ID", "")).replace(".0", "")
            title = str(row.get("Job Title", ""))
            manager = str(row.get("Manager", "")) # NEW: Extract Manager
            desc = str(row.get("Description", ""))
            notes = str(row.get("FG Notes", ""))
            transcript = str(row.get("Spotlight Transcript", "")) # NEW: Extract Transcript

            # Build the dropdown string safely, ignoring empty values
            option_parts = [req_id, vms_id, title, manager]
            option_str = " - ".join([p.strip() for p in option_parts if p.strip()])
            
            if option_str and option_str != "-":
                options.append(option_str)
                # Added transcript to the mapping dictionary
                req_mapping[option_str] = {"desc": desc, "notes": notes, "transcript": transcript}

    if "fm_jd" not in st.session_state:
        st.session_state.fm_jd = ""
    if "fm_notes" not in st.session_state:
        st.session_state.fm_notes = ""
    if "fm_trans" not in st.session_state:
        st.session_state.fm_trans = ""

    def update_jd_text():
        if "fm_req_selector" not in st.session_state:
            return

        selected = st.session_state.fm_req_selector
        if selected != "None" and selected in req_mapping:
            st.session_state.fm_jd = req_mapping[selected].get("desc", "").strip()
            st.session_state.fm_notes = req_mapping[selected].get("notes", "").strip()
            st.session_state.fm_trans = req_mapping[selected].get("transcript", "").strip()
        else:
            # Clear the text boxes if "None" is selected
            st.session_state.fm_jd = ""
            st.session_state.fm_notes = ""
            st.session_state.fm_trans = ""
            
    selected_req = st.selectbox(
        "🔍 Select a Fannie Mae Requisition (Type to search):",
        options=options,
        key="fm_req_selector",
        on_change=update_jd_text
    )

    Job_Description = st.text_area(
        "1. Official Job Description (Paste generic JD here):", 
        height=150, 
        key="fm_jd"
    )

    Chat_Notes = st.text_area(
        "2. Fieldglass Chat Notes (Optional - Overrides JD):", 
        height=100, 
        key="fm_notes"
    )

    Call_Transcript = st.text_area(
        "3. Spotlight Call Transcript (Optional - Highest Priority):", 
        height=100, 
        key="fm_trans"
    )

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="fm_res")

    if st.button("🚀 Generate Fannie Mae Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line. (e.g., If the resume says "DUPONT, DELAWARE TCS", you must output exactly "DUPONT").
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title (e.g., remove '- Contract', '(Contract)', or '- Consultant').
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title()
                    mapping = {
                        "FullName": name, "Location": Current_Location_City_ST, "Remote": Remote_or_Onsite,
                        "Certifications": data.get("Certifications", ""),
                        "FormerFM": Former_FM, "Links": LinkedIn_GitHub_Portfolio_Link,
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY": "", "SKILL1": "", "YEARS1": "", "SKILL2": "", "YEARS2": "",
                        "SKILL3": "", "YEARS3": "", "SKILL4": "", "YEARS4": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The summary should read like an experienced recruiter verbally explaining to the hiring manager why this candidate deserves an interview. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed, optimized for both human reading and ATS exact-match parsing.

                            ========================
                            JD vs. MANAGER INPUT (CRITICAL)
                            ========================

                            You are receiving up to three sources of information:

                            1. Official Job Description
                            2. Spotlight Call Transcript
                            3. Fieldglass Chat Notes

                            OFFICIAL JOB DESCRIPTION:
                            Primary source for ATS keyword coverage and phrase matching.

                            SPOTLIGHT CALL + CHAT NOTES:
                            Primary source for hiring manager priorities and emphasis.

                            If manager input contradicts or removes a requirement from the Job Description, follow the manager.

                            Preserve strong coverage of the Official Job Description because automated shortlisting is assumed to rely primarily on Job Description terminology.

                            Never highlight technical skills or responsibilities that the manager explicitly negated.

                            ========================
                            ATS COVERAGE RULE (CRITICAL)
                            ========================

                            Assume the Fieldglass AI shortlisting engine evaluates candidates primarily against the Official Job Description.

                            Therefore:
                            - The Official Job Description is the primary source for keyword coverage and phrase matching.
                            - Spotlight Call transcripts and Chat Notes should influence emphasis and prioritization, but should not reduce coverage of important Job Description concepts.
                            - Preserve exact terminology from the Job Description whenever supported by the resume.
                            - Prefer exact Job Description terminology over synonyms when both are supported by the resume.
                            - Exact phrase overlap is preferred whenever it does not distort the truth.
                            - When multiple resume experiences are equally relevant, prefer the one that produces the greatest overlap with Job Description terminology and concepts.

                            ========================
                            THE NARRATIVE BLUEPRINT (4 Sentences Max)
                            ========================
                            Follow this exact structure for the SUMMARY. Every sentence MUST sell the candidate's fit for the role:
                            - Sentence 1: The Anchor (Authority). Who are they, what is their TOTAL progressive professional experience, and what is their dominant expertise that solves the PRIMARY technical "must-have" of the Job Description? (Calculate their total overall years strictly rounding DOWN to the nearest whole year formatted as 'X+ years'. Use their FIRST NAME only. You MUST use the exact job title requested in the JD if the candidate's history supports it. Avoid generic fluff).
                            - Sentence 2: The Alignment (The Hook). You MUST explicitly name the current employer unless an earlier role provides substantially stronger evidence for the target role. Highlight the most relevant accomplishment or responsibility that demonstrates success performing the core duties of this role. Show why this experience provides strong evidence they can perform the responsibilities of the target position.
                            - Sentence 3: The Execution & Impact (The Proof). Explain how they executed the work and why it mattered. Highlight the scale, complexity, operational ownership, or business impact. Weave tools and methodologies naturally into the sentence. DO NOT write a comma-separated list of tools. DO NOT repeat verbs from Sentence 2.
                            - Sentence 4: The Closer (The ROI). Based on their past execution, what specific value will they deliver on Day 1 in THIS new role? (Use a strong, direct structure like: "[First Name]'s success in [X] makes them an immediate asset for driving this team's [specific technical goal/initiative]." Do NOT mention the physical location/city).

                            CRITICAL ATS HACK: Across these 4 sentences, you MUST seamlessly embed 2-3 exact phrases from the Job Description (including soft skills like "changing priorities" or "system analysis") to maximize the Fieldglass match score. Do not force them if they ruin the sentence flow, but prioritize exact phrase matching where supported by the resume.

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - TECH MATCHING: Strictly align the tools you highlight with the JD. If the JD asks for AWS, highlight AWS. Do not highlight competing tech (like Azure or GCP) just because it's prominent in the resume, unless it is their only experience.
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - SHOW, DON'T TELL. 
                              🔴 Bad: "John's background in AWS makes him a great fit for this role."
                              🟢 Good: "Because John spent the last three years building highly available data lakes in AWS, he can immediately step in to optimize your current infrastructure."
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely," "exceptionally well-prepared," "fits well," "aligns with," "enterprise-grade platforms."
                            - Do NOT use transition crutches: "Additionally," "Furthermore," "Moreover."
                            - Do NOT repeat or restate the job description.
                            - Prefer clear business language over consultant buzzwords.
                            - Favor evidence and ownership over dramatic wording.
                            - Use verbs that naturally fit the work performed. Do not force "leadership verbs."

                            ========================
                            EXAMPLE OF A PERFECT SUMMARY
                            ========================
                            "Sarah is a Senior Data Engineer with 7+ years of experience architecting cloud-native data migrations within heavily regulated financial environments. Most recently at Capital One, she led the end-to-end migration of a legacy on-prem data warehouse to AWS, directly mirroring the scale and compliance rigor required for this team's current cloud initiative. By engineering automated ETL pipelines with Python, PySpark, and Apache Airflow, she processed 5TB of daily transaction data and reduced reporting latency by 40%. Sarah's success in navigating complex data governance structures makes her an immediate asset for driving this team's AWS migration goals."

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth for experience.
                            - The technical interview Q&A is SECONDARY and should only be used to:
                              • clarify depth
                              • add supporting detail
                              • reinforce experience already shown in the resume

                            DO NOT:
                            - Base the summary primarily on Q&A responses.
                            - Introduce tools, systems, or concepts that are only mentioned in Q&A but not supported by the resume.
                            - Overweight theoretical or idealized answers from Q&A.

                            If there is any conflict: 👉 prioritize the resume over Q&A.
                            The summary should reflect what the candidate has DONE, not just what they can describe.

                            ========================
                            EXPERIENCE PRIORITY RULE (CRITICAL)
                            ========================

                            Prefer accomplishments explicitly demonstrated in experience bullets over technologies listed only in the skills section.

                            Favor things the candidate built, created, supported, deployed, maintained, optimized, or owned rather than technologies merely listed in a skills inventory.

                            Prefer recent and recurring responsibilities over isolated projects whenever both are supported by the resume.

                            When multiple experiences are relevant, prefer the experience that overlaps with the largest number of Job Description requirements and responsibilities.

                            ========================
                            OPERATIONAL OWNERSHIP RULE
                            ========================

                            Prefer recurring responsibilities and direct ownership over one-time projects or transformation initiatives.

                            Unless the role is explicitly implementation-focused, prioritize evidence of day-to-day execution, production support, issue resolution, audits, compliance, platform ownership, and measurable business outcomes over project participation or consulting activities.

                            ========================
                            SKILLS SECTION
                            ========================
                            - EXACTLY 4 items
                            - Prioritize the specific "Must-Have" technologies AND methodologies (e.g., Agile, Data Governance) requested in the JD.
                            - Highly relevant + keyword-rich
                            - Use only tools explicitly mentioned in resume/Q&A/notes

                            Format:
                            "Skill Area (Tool1, Tool2, Tool3, Tool4)"

                            Years format:
                            "X+ years, current" OR "X+ years, 2026"

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME. If the JD asks for "Redshift" and the resume only says "AWS", you MUST NOT write "Redshift" under any circumstances. 
                            3. DO NOT INFLATE TO MATCH THE JD. If the candidate lacks a requested skill, omit it completely. It is better to have an incomplete match than a fabricated one.
                            4. FACT AUDIT: Before outputting the final JSON, you must verify every single tool mentioned in your SUMMARY and SKILLS. If a tool exists in the JD but not in the Resume/Q&A, you must delete it from your output.
                            5. DO NOT ELEVATE MANAGER LANGUAGE INTO CANDIDATE EXPERIENCE. If the Job Description, Chat Notes, or Spotlight Call mention technologies, responsibilities, or business goals, you may reference them as objectives of the target role, but you MUST NOT imply that the candidate performed those activities unless supported by the resume or Q&A.

                            ========================
                            YEARS ACCURACY RULES (REALISTIC RECRUITER MODE)
                            ========================
                            - Use July 2026 as the current date.
                            - STRICT MATH (DATE-DRIVEN ONLY): Calculate years strictly based on the earliest chronological date provided in the 'Professional Experience' or 'Work History' section. You MUST completely IGNORE any self-reported years of experience in the candidate's summary blurb (e.g., if their summary claims "12+ years" but their listed jobs only go back to 2019, you must calculate from 2019). Round DOWN to the nearest whole year and use the exact format "X+ years". Do not use phrases like "nearly X years". (e.g., If the job history calculates to 7 years and 10 months, output "7+ years". NEVER round up to 8+). 
                            - Foundational skills (e.g., Python, SQL, general engineering) should get their maximum calculated years.
                            - Advanced/Specialized tools (e.g., SageMaker, Kubernetes, Cloud Architecture) should realistically be calculated at 1-2 years less than their maximum total experience unless the resume explicitly proves Day 1 usage. 
                            - Default to "current" if they are still working in this general technical field. Do NOT use past years (e.g., "2022"). Always bridge past experience to their current tenure.

                            ========================
                            🧠 FINAL POLISH CHECKLIST (ONE PASS ONLY)
                            ========================
                            Before outputting the JSON, evaluate your drafted SUMMARY and SKILLS against these 7 checks:
                            1. SIMPLIFIED: No run-on sentences or unnecessary filler.
                            2. HUMAN TONE: No generic claims like "highly experienced."
                            3. TOOL DENSITY: Maximum 3 tools per sentence.
                            4. NO REPETITION: Do not repeat verbs or concepts across sentences.
                            5. THE PITCH: Ensure a natural, confident recruiter tone.
                            6. CURRENT JOB: Did you explicitly name their most recent employer in Sentence 2?
                            7. ACCURATE MATH: Did you strictly round down their years of experience and use the 'X+ years' format to prevent ATS parser flags?
                            8. OWNERSHIP TEST: Did I emphasize what the candidate repeatedly owned and executed, rather than simply the most impressive-sounding project?
                            9. ATS COVERAGE TEST: Did I naturally incorporate the most important technologies, responsibilities, and terminology from the Official Job Description that are supported by the resume?
                            10. EVIDENCE TEST: Would every skill, tool, accomplishment, and responsibility in the summary survive a side-by-side comparison against the actual resume?
                            11. JD OVERLAP TEST: Did I emphasize the experiences that provide the strongest overlap with the Official Job Description rather than simply the most impressive accomplishments?

                            Output only the final, polished JSON.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY": "4 sentence summary following the blueprint",
                              "SKILL1": "Skill Area (tools)",
                              "YEARS1": "X+ years, current",
                              "SKILL2": "Skill Area (tools)",
                              "YEARS2": "X+ years, current",
                              "SKILL3": "Skill Area (tools)",
                              "YEARS3": "X+ years, current",
                              "SKILL4": "Skill Area (tools)",
                              "YEARS4": "X+ years, current"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Official Job Description:
                            {Job_Description}

                            Fieldglass Chat Notes:
                            {Chat_Notes}

                            Spotlight Call Transcript:
                            {Call_Transcript}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY"] = summary_data.get("SUMMARY", "")
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["YEARS1"] = summary_data.get("YEARS1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["YEARS2"] = summary_data.get("YEARS2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["YEARS3"] = summary_data.get("YEARS3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")
                            mapping["YEARS4"] = summary_data.get("YEARS4", "")
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_Fannie_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- 🟣 FREDDIE MAC HELPER FUNCTIONS 🟣 ---
# ====================================================================

FREDDIE_VNDLY_SKILLS_FILENAME = "freddie_vndly_skills_curated.json"


@st.cache_data(show_spinner=False)
def fred_load_vndly_skill_catalog(path=FREDDIE_VNDLY_SKILLS_FILENAME):
    """
    Load the curated VNDLY skill catalogue bundled with the Streamlit app.
    """
    if not os.path.exists(path):
        return []

    try:
        with open(path, "r", encoding="utf-8") as f:
            payload = json.load(f)

        raw_skills = payload.get("skills", payload)

        if not isinstance(raw_skills, list):
            return []

        cleaned = []

        for item in raw_skills:
            if not isinstance(item, dict):
                continue

            skill_name = str(
                item.get("skill_name")
                or item.get("value")
                or item.get("label")
                or ""
            ).strip()

            if not skill_name:
                continue

            cleaned.append(
                {
                    "id": item.get("id"),
                    "skill_name": skill_name,
                    "catalog_tier": str(
                        item.get("catalog_tier", "canonical_preferred")
                    ).strip() or "canonical_preferred",
                }
            )

        return cleaned

    except Exception:
        return []


def fred_normalize_skill_name(value):
    """Conservative normalization used only for exact catalogue lookups."""
    return re.sub(
        r"\s+",
        " ",
        str(value or "").replace("\u00a0", " "),
    ).strip().casefold()


def fred_catalog_lookup(catalog):
    """Build a normalized-name -> exact VNDLY record lookup."""
    lookup = {}

    for item in catalog or []:
        if not isinstance(item, dict):
            continue

        name = str(item.get("skill_name", "")).strip()
        if not name:
            continue

        lookup[fred_normalize_skill_name(name)] = item

    return lookup


def fred_sanitize_recommended_skills(raw_items, catalog):
    """
    Force Gemini recommendations back onto exact catalogue values.
    """
    if not isinstance(raw_items, list):
        return []

    lookup = fred_catalog_lookup(catalog)
    cleaned = []
    seen = set()

    for item in raw_items:
        if isinstance(item, dict):
            proposed = (
                item.get("skill_name")
                or item.get("value")
                or item.get("label")
                or ""
            )
            evidence = str(item.get("evidence", "") or "").strip()
            reason = str(item.get("reason", "") or "").strip()
        else:
            proposed = str(item)
            evidence = ""
            reason = ""

        exact = lookup.get(fred_normalize_skill_name(proposed))

        if not exact:
            continue

        key = fred_normalize_skill_name(exact["skill_name"])
        if key in seen:
            continue

        seen.add(key)

        cleaned.append(
            {
                "id": exact.get("id"),
                "skill_name": exact["skill_name"],
                "catalog_tier": exact.get(
                    "catalog_tier",
                    "canonical_preferred",
                ),
                "evidence": evidence,
                "reason": reason,
            }
        )

    return cleaned


def fred_dedupe_skill_groups(*groups):
    """Remove duplicates while preserving priority order."""
    seen = set()
    output = []

    for group in groups:
        clean_group = []

        for item in group or []:
            name = str(item.get("skill_name", "")).strip()
            key = fred_normalize_skill_name(name)

            if not name or key in seen:
                continue

            seen.add(key)
            clean_group.append(item)

        output.append(clean_group)

    return output


def fred_resolve_structured_vndly_skills(raw_field, catalog):
    """
    Resolve a raw Google-Sheet Must/Nice field back to the exact VNDLY skill
    values that produced it.

    The sheet stores multiple selected skills as one comma-separated string, but
    some individual VNDLY skills also contain commas.  This uses the known exact
    catalogue and a longest-match dynamic-programming pass so we do not mistake
    a JD keyword such as "Java" for an exact requisition-selected skill.
    """
    raw = re.sub(r"\s+", " ", str(raw_field or "").replace("\u00a0", " ")).strip()

    if not raw:
        return []

    candidates = []

    for item in catalog or []:
        if not isinstance(item, dict):
            continue

        name = re.sub(
            r"\s+",
            " ",
            str(item.get("skill_name", "") or "").replace("\u00a0", " "),
        ).strip()

        if name:
            candidates.append((name, item))

    # Prefer the longest exact catalogue value at any ambiguous position.
    candidates.sort(key=lambda pair: len(pair[0]), reverse=True)

    raw_fold = raw.casefold()
    memo = {}

    def skip_separator(pos):
        while pos < len(raw) and raw[pos].isspace():
            pos += 1

        if pos < len(raw) and raw[pos] in {",", ";", "|"}:
            pos += 1
            while pos < len(raw) and raw[pos].isspace():
                pos += 1

        return pos

    def solve(pos):
        while pos < len(raw) and raw[pos].isspace():
            pos += 1

        if pos >= len(raw):
            return []

        if pos in memo:
            return memo[pos]

        for name, item in candidates:
            name_fold = name.casefold()

            if not raw_fold.startswith(name_fold, pos):
                continue

            end = pos + len(name)

            # Exact skill must end at the raw-field end or at a separator.
            check = end
            while check < len(raw) and raw[check].isspace():
                check += 1

            if check < len(raw) and raw[check] not in {",", ";", "|"}:
                continue

            next_pos = skip_separator(check)

            remainder = solve(next_pos)

            if remainder is not None:
                result = [item] + remainder
                memo[pos] = result
                return result

        memo[pos] = None
        return None

    resolved = solve(0)

    if resolved is not None:
        return resolved

    # Conservative fallback: only accept a whole-field exact catalogue match.
    whole = fred_catalog_lookup(catalog).get(
        fred_normalize_skill_name(raw)
    )

    return [whole] if whole else []


def fred_filter_exact_req_recommendations(raw_items, allowed_items, catalog):
    """
    Validate Gemini output against the deterministic set of skills actually
    selected on this requisition's structured Must/Nice field.
    """
    sanitized = fred_sanitize_recommended_skills(raw_items, catalog)

    allowed = {
        fred_normalize_skill_name(item.get("skill_name", ""))
        for item in (allowed_items or [])
        if str(item.get("skill_name", "")).strip()
    }

    if not allowed:
        return []

    return [
        item
        for item in sanitized
        if fred_normalize_skill_name(item.get("skill_name", "")) in allowed
    ]


FREDDIE_VNDLY_SKILL_CAP = 8


def fred_extract_work_history_text(raw_resume_text):
    """
    Return the portion of the resume most likely to contain dated work history.
    Used only as a conservative evidence-strength check for VNDLY skill ranking.
    """
    text = str(raw_resume_text or "")
    upper = text.upper()

    start_markers = [
        "PROFESSIONAL EXPERIENCE",
        "WORK EXPERIENCE",
        "WORK HISTORY",
        "EMPLOYMENT HISTORY",
        "EXPERIENCE",
    ]

    start = None
    for marker in start_markers:
        idx = upper.find(marker)
        if idx != -1 and (start is None or idx < start):
            start = idx

    if start is None:
        return text

    work = text[start:]

    # Education/certification sections usually mark the end of work history.
    upper_work = work.upper()
    end_candidates = []
    for marker in [
        "\nEDUCATION",
        "\nACADEMIC",
        "\nCERTIFICATIONS",
        "\nCERTIFICATION",
    ]:
        idx = upper_work.find(marker)
        if idx > 0:
            end_candidates.append(idx)

    if end_candidates:
        work = work[:min(end_candidates)]

    return work


def fred_skill_work_history_evidence(skill_name, raw_resume_text):
    """
    Conservative deterministic check for whether a VNDLY skill is evidenced in
    the dated work-history portion of the resume rather than only in a summary
    or top-level skills inventory.

    Returns:
      2 = clear dated work-history evidence
      1 = only broader/semantic domain evidence
      0 = no meaningful dated work-history evidence found
    """
    work = fred_extract_work_history_text(raw_resume_text)
    work_fold = work.casefold()
    skill = str(skill_name or "").strip()
    key = fred_normalize_skill_name(skill)

    if not skill:
        return 0

    alias_map = {
        "java": ["java"],
        "j2ee": ["j2ee", "java ee", "jee"],
        "java database connectivity (jdbc)": ["jdbc", "java database connectivity"],
        "java servlets": ["servlet", "servlets"],
        "javaserver pages (jsp)": ["jsp", "jsps", "javaserver pages"],
        "spring framework": ["spring framework", "spring "],
        "spring boot": ["spring boot"],
        "relational database": [
            "relational database", "oracle", "db2", "sql server", "mysql",
            "postgresql", "sybase", "jdbc", "sql queries", "sql statements"
        ],
        "relational databases": [
            "relational database", "oracle", "db2", "sql server", "mysql",
            "postgresql", "sybase", "jdbc", "sql queries", "sql statements"
        ],
        "junit testing": ["junit"],
        "mockito": ["mockito"],
        "git": ["git", "gitlab", "github"],
        "jenkins (software)": ["jenkins"],
        "docker (software)": ["docker"],
        "maven": ["maven"],
        "gradle": ["gradle"],
        "eclipse ide": ["eclipse"],
        "python": ["python"],
        "financial services": [
            "bank", "banking", "financial", "investment", "portfolio",
            "securit", "trading", "trade ", "retirement", "securities",
            "asset", "capital markets"
        ],
    }

    aliases = alias_map.get(key)

    if aliases:
        if any(alias.casefold() in work_fold for alias in aliases):
            return 2
        return 0

    # Generic fallback: strip catalogue parentheticals and test the meaningful
    # core phrase in the dated work history.
    core = re.sub(r"\([^)]*\)", "", skill).strip().casefold()

    if core and len(core) >= 4 and core in work_fold:
        return 2

    # Domain-like catalogue values can be supported semantically even if the
    # exact catalogue phrase is not present.
    domain_terms = {
        "mortgage": ["mortgage", "appraisal", "loan", "uad"],
        "capital markets": ["capital markets", "securities", "trading", "trade "],
        "banking": ["bank", "banking"],
        "insurance": ["insurance"],
    }

    for domain_key, terms in domain_terms.items():
        if domain_key in key and any(term in work_fold for term in terms):
            return 1

    return 0


def fred_vndly_skill_family(skill_name):
    """High-confidence skill families used only for Top-8 redundancy control."""
    key = fred_normalize_skill_name(skill_name)

    families = {
        "spring": {
            "spring",
            "java spring",
            "spring framework",
            "spring boot",
            "spring mvc",
            "spring web",
            "spring cloud",
        },
        "relational_database": {
            "relational database",
            "relational databases",
            "oracle",
            "db2",
            "ibm db2",
            "microsoft sql server",
            "mysql",
            "postgresql",
            "sybase",
        },
        "testing": {
            "junit testing",
            "mockito",
            "unit testing",
        },
        "source_control": {
            "git",
            "github",
            "gitlab",
        },
    }

    for family_name, members in families.items():
        if key in members:
            return family_name

    return ""


def fred_skill_is_explicitly_named_in_job(skill_name, structured_req):
    """
    Conservative check for whether a catalogue skill is explicitly named in
    current formal requisition intelligence.
    """
    req_text = json.dumps(
        structured_req or {},
        ensure_ascii=False,
    ).casefold()

    key = fred_normalize_skill_name(skill_name)

    aliases = {
        "java database connectivity (jdbc)": [
            "jdbc",
            "java database connectivity",
        ],
        "javaserver pages (jsp)": [
            "jsp",
            "javaserver pages",
            "java server pages",
        ],
        "java servlets": ["servlet", "servlets"],
        "junit testing": ["junit"],
        "mockito": ["mockito"],
        "j2ee": ["j2ee", "j2ee technologies", "java ee"],
        "java": ["java"],
        "git": ["git"],
        "eclipse ide": ["eclipse"],
        "jenkins (software)": ["jenkins"],
        "docker (software)": ["docker"],
        "gradle": ["gradle"],
        "db2": ["db2"],
        "ibm db2": ["ibm db2", "db2"],
        "microsoft sql server": ["sql server"],
        "spring framework": ["spring framework"],
        "relational database": [
            "relational database",
            "relational databases",
        ],
        "financial services": [
            "finance background",
            "financial services",
            "financial domain",
        ],
    }

    terms = aliases.get(key, [key])
    return any(term and term in req_text for term in terms)


def fred_skill_aliases_for_requirement(skill_name):
    """Aliases used to map catalogue skills back to Freddie's requirement text."""
    key = fred_normalize_skill_name(skill_name)
    aliases = {
        "j2ee": ["j2ee", "j2ee technologies", "java ee"],
        "java database connectivity (jdbc)": ["jdbc", "java database connectivity"],
        "javaserver pages (jsp)": ["jsp", "javaserver pages", "java server pages"],
        "java servlets": ["servlet", "servlets"],
        "junit testing": ["junit"],
        "mockito": ["mockito"],
        "spring framework": ["spring framework"],
        "spring boot": ["spring boot"],
        "relational database": ["relational database", "relational databases"],
        "financial services": ["finance background", "financial services", "financial domain"],
        "eclipse ide": ["eclipse"],
        "jenkins (software)": ["jenkins"],
        "docker (software)": ["docker"],
        "microsoft sql server": ["sql server"],
        "ibm db2": ["ibm db2", "db2"],
        "db2": ["db2"],
        "git": ["git"],
        "gradle": ["gradle"],
        "maven": ["maven"],
        "java": ["java"],
    }
    return aliases.get(key, [key])


def fred_skill_matches_text(skill_name, text_value):
    text_fold = str(text_value or "").casefold()
    return any(alias and alias in text_fold for alias in fred_skill_aliases_for_requirement(skill_name))


def fred_skill_requirement_group(skill_name, structured_req):
    """Return the 0-based must-have competency group that names this skill."""
    groups = (structured_req or {}).get("must_have_competency_groups", []) or []
    for group_index, group in enumerate(groups):
        group_text = " ".join(
            [
                str(group.get("name", "") or ""),
                " ".join(str(x) for x in (group.get("terms", []) or [])),
                str(group.get("reason", "") or ""),
            ]
        )
        if fred_skill_matches_text(skill_name, group_text):
            return group_index
    return 99


def fred_skill_requirement_term_position(skill_name, structured_req):
    """Position of the skill inside its highest-priority must-have group."""
    groups = (structured_req or {}).get("must_have_competency_groups", []) or []
    for group_index, group in enumerate(groups):
        terms = group.get("terms", []) or []
        for term_index, term in enumerate(terms):
            if fred_skill_matches_text(skill_name, str(term)):
                return group_index, term_index
        if fred_skill_matches_text(skill_name, str(group.get("name", "") or "")):
            return group_index, 99
    return 99, 99


def fred_skill_clause_strength(skill_name, structured_req):
    """
    Prefer skills named in focused/dedicated mandatory clauses over tools buried
    in a long multi-tool list. Higher is stronger.
    """
    clauses = []
    for key in ["explicit_must_have_requirements", "required_requirements", "domain_requirements"]:
        clauses.extend((structured_req or {}).get(key, []) or [])

    best = 0
    aliases = fred_skill_aliases_for_requirement(skill_name)

    for clause in clauses:
        clause_text = str(clause or "")
        fold = clause_text.casefold()
        if not any(alias and alias in fold for alias in aliases):
            continue

        # Dedicated requirements score higher than long tool bundles.
        separators = clause_text.count(",") + len(re.findall(r"\band\b", fold)) + clause_text.count(";")
        strength = max(1, 12 - separators)

        # Repeated mention across formal clauses is additional signal.
        occurrences = sum(fold.count(alias) for alias in aliases if alias)
        strength += min(3, occurrences)
        best = max(best, strength)

    return best


def fred_supplement_explicit_required_skills(
    ranked_skills,
    structured_req,
    catalog,
    raw_resume_text,
):
    """
    Deterministically add candidate-supported catalogue skills that are explicitly
    required by the current requisition, even when Gemini omitted them from its
    candidate pool. This protects high-signal requirements such as J2EE/JDBC.
    """
    output = [dict(item) for item in (ranked_skills or [])]
    seen = {
        fred_normalize_skill_name(item.get("skill_name", ""))
        for item in output
    }

    for cat_item in catalog or []:
        name = str(cat_item.get("skill_name", "") or "").strip()
        if not name:
            continue

        key = fred_normalize_skill_name(name)
        if key in seen:
            continue

        if cat_item.get("catalog_tier") not in {"canonical_preferred", "compound_review"}:
            continue

        if not fred_skill_is_explicitly_named_in_job(name, structured_req):
            continue

        evidence_strength = fred_skill_work_history_evidence(name, raw_resume_text)
        if evidence_strength <= 0:
            continue

        output.append(
            {
                "id": cat_item.get("id"),
                "skill_name": name,
                "catalog_tier": cat_item.get("catalog_tier", "canonical_preferred"),
                "source_category": "required_job_intelligence",
                "evidence": "Dated work-history evidence",
                "reason": "Explicit current Freddie requirement",
                "work_history_evidence_strength": evidence_strength,
            }
        )
        seen.add(key)

    return output


def fred_required_skill_sort_key(item, structured_req, original_index=0):
    category_priority = {
        "exact_structured_must": 0,
        "required_job_intelligence": 1,
        "exact_structured_nice": 2,
        "preferred_job_intelligence": 3,
        "additional_high_value": 4,
    }
    category = str(item.get("source_category", "") or "").strip().lower()
    group_index = fred_skill_requirement_group(item.get("skill_name", ""), structured_req)
    clause_strength = fred_skill_clause_strength(item.get("skill_name", ""), structured_req)
    evidence = int(item.get("work_history_evidence_strength", 0) or 0)
    return (
        category_priority.get(category, 99),
        -clause_strength,
        group_index,
        -evidence,
        original_index,
    )


def fred_apply_vndly_redundancy_control(
    ranked_skills,
    structured_req,
    cap=FREDDIE_VNDLY_SKILL_CAP,
):
    """
    Final Top-8 selection that balances requirement priority, distinct coverage,
    dated evidence, and redundancy control.
    """
    candidates = []
    seen = set()

    for idx, item in enumerate(ranked_skills or []):
        name = str(item.get("skill_name", "") or "").strip()
        key = fred_normalize_skill_name(name)
        if not name or key in seen:
            continue
        copy_item = dict(item)
        copy_item["_sort_key"] = fred_required_skill_sort_key(copy_item, structured_req, idx)
        candidates.append(copy_item)
        seen.add(key)

    if not candidates:
        return []

    # First reserve one strong candidate-supported skill for each must-have
    # competency group so later generic tools cannot crowd out a whole requirement area.
    selected = []
    selected_keys = set()
    family_counts = {}
    groups = (structured_req or {}).get("must_have_competency_groups", []) or []

    def can_add(item):
        name = str(item.get("skill_name", "") or "").strip()
        key = fred_normalize_skill_name(name)
        if not name or key in selected_keys:
            return False

        family = fred_vndly_skill_family(name)

        selected_names = {
            fred_normalize_skill_name(existing.get("skill_name", ""))
            for existing in selected
        }

        # J2EE already communicates the enterprise-Java umbrella. With only
        # eight slots, Servlets/JSP are normally redundant once J2EE is selected;
        # JDBC remains eligible because it adds a distinct database-connectivity
        # screening signal.
        if "j2ee" in selected_names and key in {
            "java servlets",
            "javaserver pages (jsp)",
        }:
            return False

        if key == "j2ee" and (
            "java servlets" in selected_names
            or "javaserver pages (jsp)" in selected_names
        ):
            # Prefer the umbrella J2EE requirement over its narrower UI/server
            # components when both are candidate-supported.
            pass

        family_caps = {
            "spring": 1,
            "relational_database": 1,
            "source_control": 1,
            "testing": 2,
        }

        if family:
            cap_for_family = family_caps.get(family, 1)
            if family_counts.get(family, 0) >= cap_for_family:
                return False

        return True

    def add_item(item):
        name = str(item.get("skill_name", "") or "").strip()
        key = fred_normalize_skill_name(name)
        selected.append(item)
        selected_keys.add(key)
        family = fred_vndly_skill_family(name)
        if family:
            family_counts[family] = family_counts.get(family, 0) + 1

    for group_index in range(min(len(groups), cap)):
        group_candidates = [
            item for item in candidates
            if fred_skill_requirement_group(item.get("skill_name", ""), structured_req) == group_index
            and can_add(item)
        ]
        if group_candidates:
            group_candidates.sort(key=lambda item: item["_sort_key"])
            add_item(group_candidates[0])

    # Fill the remaining slots with the strongest uncovered formal requirements.
    for item in sorted(candidates, key=lambda x: x["_sort_key"]):
        if len(selected) >= cap:
            break
        if can_add(item):
            add_item(item)

    # Reorder the chosen set by requirement priority for a clean recruiter-facing list.
    selected.sort(
        key=lambda item: (
            fred_skill_requirement_term_position(
                item.get("skill_name", ""),
                structured_req,
            )[0],
            fred_skill_requirement_term_position(
                item.get("skill_name", ""),
                structured_req,
            )[1],
            item["_sort_key"],
        )
    )
    for item in selected:
        item.pop("_sort_key", None)
    return selected[:cap]


def fred_rerank_vndly_skills_by_evidence(
    ranked_skills,
    raw_resume_text,
    structured_req=None,
    cap=FREDDIE_VNDLY_SKILL_CAP,
):
    """Requirement-aware evidence reranking for the VNDLY candidate pool."""
    rescored = []

    for original_index, item in enumerate(ranked_skills or []):
        evidence_strength = fred_skill_work_history_evidence(
            item.get("skill_name", ""),
            raw_resume_text,
        )
        copy_item = dict(item)
        copy_item["work_history_evidence_strength"] = evidence_strength
        rescored.append(
            (
                fred_required_skill_sort_key(copy_item, structured_req or {}, original_index),
                copy_item,
            )
        )

    rescored.sort(key=lambda row: row[0])
    return [row[1] for row in rescored[:cap]]


def fred_build_final_ranked_vndly_skills(
    raw_ranked_items,
    selected_must_items,
    selected_nice_items,
    catalog,
    cap=FREDDIE_VNDLY_SKILL_CAP,
):
    """
    Validate Gemini's single ranked recommendation list and enforce a hard
    GLOBAL cap across every VNDLY skill source.

    Priority is decided in the AI ranking prompt, then this function enforces:
    - exact catalogue values only;
    - structured Must/Nice claims must truly be attached to this requisition;
    - legacy/free-text values are allowed only for exact structured req skills;
    - no duplicates;
    - maximum `cap` skills total.
    """
    if not isinstance(raw_ranked_items, list):
        return []

    catalog_lookup = fred_catalog_lookup(catalog)

    allowed_must = {
        fred_normalize_skill_name(item.get("skill_name", ""))
        for item in (selected_must_items or [])
        if str(item.get("skill_name", "") or "").strip()
    }
    allowed_nice = {
        fred_normalize_skill_name(item.get("skill_name", ""))
        for item in (selected_nice_items or [])
        if str(item.get("skill_name", "") or "").strip()
    }

    valid_categories = {
        "exact_structured_must",
        "exact_structured_nice",
        "required_job_intelligence",
        "preferred_job_intelligence",
        "additional_high_value",
    }

    output = []
    seen = set()

    for raw_item in raw_ranked_items:
        # Validate the complete candidate pool here. The final hard cap is
        # applied after deterministic work-history evidence reranking.

        if not isinstance(raw_item, dict):
            continue

        proposed = (
            raw_item.get("skill_name")
            or raw_item.get("value")
            or raw_item.get("label")
            or ""
        )
        category = str(
            raw_item.get("source_category", "")
        ).strip().lower()

        if category not in valid_categories:
            continue

        exact = catalog_lookup.get(
            fred_normalize_skill_name(proposed)
        )

        if not exact:
            continue

        name = str(exact.get("skill_name", "") or "").strip()
        key = fred_normalize_skill_name(name)

        if not name or key in seen:
            continue

        tier = exact.get(
            "catalog_tier",
            "canonical_preferred",
        )

        # Exact structured categories must actually come from the exact req field.
        if category == "exact_structured_must":
            if key not in allowed_must:
                continue
        elif category == "exact_structured_nice":
            if key not in allowed_nice:
                continue
        else:
            # Legacy/free-text is never used for inferred JD/manager/additional
            # recommendations. It is reserved for exact req-attached tags.
            if tier not in {
                "canonical_preferred",
                "compound_review",
            }:
                continue

        output.append(
            {
                "id": exact.get("id"),
                "skill_name": name,
                "catalog_tier": tier,
                "source_category": category,
                "evidence": str(
                    raw_item.get("evidence", "") or ""
                ).strip(),
                "reason": str(
                    raw_item.get("reason", "") or ""
                ).strip(),
            }
        )
        seen.add(key)

    return output


def fred_split_ranked_vndly_skills(ranked_skills):
    """Split the globally-ranked list into UI/email groups without changing rank."""
    groups = {
        "exact_must_have_skills": [],
        "exact_nice_to_have_skills": [],
        "required_vndly_skills_from_job_intelligence": [],
        "preferred_vndly_skills_from_job_intelligence": [],
        "additional_vndly_skills": [],
    }

    category_to_key = {
        "exact_structured_must": "exact_must_have_skills",
        "exact_structured_nice": "exact_nice_to_have_skills",
        "required_job_intelligence": "required_vndly_skills_from_job_intelligence",
        "preferred_job_intelligence": "preferred_vndly_skills_from_job_intelligence",
        "additional_high_value": "additional_vndly_skills",
    }

    for item in ranked_skills or []:
        key = category_to_key.get(
            str(item.get("source_category", "")).strip().lower()
        )
        if key:
            groups[key].append(item)

    return groups


def fred_clean_vndly_submission_summary(summary, candidate_name):
    """Deterministic cleanup for the recruiter-facing VNDLY comments summary."""
    if not summary:
        return ""

    cleaned = re.sub(r"\s+", " ", str(summary)).strip()
    first_name = str(candidate_name or "").strip().split()[0] if candidate_name else ""
    last_name = str(candidate_name or "").strip().split()[-1] if candidate_name else ""

    # Remove honorific constructions the model should never use in this field.
    if last_name and first_name:
        cleaned = re.sub(
            rf"\b(?:Mr|Ms|Mrs|Dr)\.?\s+{re.escape(last_name)}\b",
            first_name,
            cleaned,
            flags=re.IGNORECASE,
        )
        cleaned = re.sub(
            rf"\b(?:Mr|Ms|Mrs|Dr)\.?\s+{re.escape(candidate_name)}\b",
            first_name,
            cleaned,
            flags=re.IGNORECASE,
        )

    cleaned = re.sub(r"\b(?:Mr|Ms|Mrs|Dr)\.\s+", "", cleaned, flags=re.IGNORECASE)

    # Strip a few recurring editorial tails without rewriting factual content.
    cleaned = re.sub(
        r",?\s+(?:which is|making (?:him|her|them))\s+(?:crucial|critical|ideal)\s+for\s+(?:this|the)\s+[^.]+\.",
        ".",
        cleaned,
        flags=re.IGNORECASE,
    )

    return cleaned.strip()


def fred_build_vndly_candidate_package(
    api_key,
    candidate_name,
    raw_resume_text,
    structured_req,
    vndly_context,
    vetting_for_ai,
    resume_summary,
    calculated_total_experience,
    candidate_positioning_title,
):
    """
    Create the Freddie-only VNDLY submission summary and skill selections.
    """
    catalog = fred_load_vndly_skill_catalog()

    if not catalog:
        raise RuntimeError(
            f"VNDLY skill catalogue '{FREDDIE_VNDLY_SKILLS_FILENAME}' "
            "was not found or could not be read."
        )

    catalog_for_ai = [
        {
            "id": item.get("id"),
            "skill_name": item.get("skill_name", ""),
            "catalog_tier": item.get(
                "catalog_tier",
                "canonical_preferred",
            ),
        }
        for item in catalog
    ]

    catalog_json = json.dumps(
        catalog_for_ai,
        ensure_ascii=False,
        separators=(",", ":"),
    )

    vndly_context = vndly_context or {}
    raw_must = str(vndly_context.get("must_have_skills", "") or "").strip()
    raw_nice = str(vndly_context.get("nice_to_have_skills", "") or "").strip()

    selected_must_items = fred_resolve_structured_vndly_skills(
        raw_must,
        catalog,
    )
    selected_nice_items = fred_resolve_structured_vndly_skills(
        raw_nice,
        catalog,
    )

    selected_must_json = json.dumps(
        [
            {
                "id": item.get("id"),
                "skill_name": item.get("skill_name", ""),
                "catalog_tier": item.get(
                    "catalog_tier",
                    "canonical_preferred",
                ),
            }
            for item in selected_must_items
        ],
        ensure_ascii=False,
        indent=2,
    )

    selected_nice_json = json.dumps(
        [
            {
                "id": item.get("id"),
                "skill_name": item.get("skill_name", ""),
                "catalog_tier": item.get(
                    "catalog_tier",
                    "canonical_preferred",
                ),
            }
            for item in selected_nice_items
        ],
        ensure_ascii=False,
        indent=2,
    )

    package_prompt = f"""
Return a valid JSON object ONLY.

You are preparing the VNDLY submission metadata for a Freddie Mac contingent
worker candidate. This is SEPARATE from the resume Summary.

CANDIDATE:
{candidate_name}

ALREADY-APPROVED RESUME SUMMARY:
{resume_summary}

RECOMMENDED CANDIDATE POSITIONING TITLE:
{candidate_positioning_title or "BLANK"}

CANDIDATE TITLE RULE:
- The title above was selected from the candidate's evidence + actual role
  function using a natural external-market title.
- If it is populated and you identify the candidate by title in the VNDLY
  Submission Summary, use that title.
- Do NOT substitute the structured VNDLY Job Title or rate-card title.
- The VNDLY Job Title may still describe the requisition in job metadata, but it
  is not automatically the candidate's professional identity.
- If the positioning title is blank, use a concise functional identity supported
  by the resume; do not default to an awkward VMS/rate-card title.

PYTHON-CALCULATED TOTAL PROFESSIONAL EXPERIENCE:
{calculated_total_experience}

CRITICAL EXPERIENCE RULE:
- If the Python-calculated total above is populated, that is the ONLY total
  career-experience figure you may use.
- Never use, repeat, paraphrase, or round a candidate-written estimate such as
  "close to 20 years" or "nearly 20 years" when it differs from the calculated
  value.
- Do not claim a requirement-specific minimum number of years unless the dated
  work history itself supports that minimum for the relevant technology/domain.
- A technology named only in the candidate's Professional Summary or skills
  inventory can establish familiarity, but by itself cannot prove a multi-year
  threshold.

STRUCTURED REQUISITION INTELLIGENCE:
{structured_req}

IMPORTANT JOB-INTELLIGENCE RULE:
- Formal requirements are NOT limited to structured VNDLY Must/Nice skill tags.
- The Official JD and authoritative manager/MSP/Spotlight guidance can establish
  Required or Preferred requirements even when the structured skill fields are blank.
- Use the existing structured requisition intelligence hierarchy exactly as supplied.
- A requirement that is formally mandatory in the JD or authoritative manager guidance
  belongs under REQUIRED_VNDLY_SKILLS_FROM_JOB_INTELLIGENCE when candidate evidence supports it.
- A requirement that is explicitly preferred/nice-to-have in the JD or authoritative
  manager guidance belongs under PREFERRED_VNDLY_SKILLS_FROM_JOB_INTELLIGENCE when supported.

STRUCTURED VNDLY MUST HAVE SKILLS FIELD — RAW TEXT:
{raw_must or "BLANK"}

DETERMINISTIC EXACT MUST-HAVE SKILLS SELECTED ON THIS REQUISITION:
{selected_must_json}

STRUCTURED VNDLY NICE TO HAVE SKILLS FIELD — RAW TEXT:
{raw_nice or "BLANK"}

DETERMINISTIC EXACT NICE-TO-HAVE SKILLS SELECTED ON THIS REQUISITION:
{selected_nice_json}

SUPPLIER VETTING RESPONSES:
{vetting_for_ai}

FULL ORIGINAL CANDIDATE RESUME:
{raw_resume_text}

EXACT VNDLY SKILL CATALOGUE:
{catalog_json}

======================================================================
TASK 1 — VNDLY SUBMISSION SUMMARY / COMMENTS BOX
======================================================================

Write the short candidate snapshot that a recruiter will paste into the VNDLY
submission Comments field. This is SEPARATE from the longer resume Summary.

PRIMARY GOALS — IN THIS ORDER:
1. Help the MSP reviewer quickly recognize that the candidate satisfies the
   highest-signal current requirements and deserves to be shortlisted.
2. Give the hiring manager a concise, evidence-based reason to interview the
   candidate if the submission reaches them.

MENTAL MODEL:
The MSP may be scanning many submissions quickly. Write a compact "why this
candidate" snapshot, NOT another executive resume summary and NOT marketing copy.
The best version should feel like a strong recruiter condensed the most useful
parts of the resume into the VNDLY comments box.

LENGTH / STRUCTURE:
- 3-4 compact sentences as ONE paragraph.
- Target roughly 80-115 words. Do not exceed 130 words.
- Sentence 1: natural candidate positioning title + deterministic total career
  experience + the 1-2 most important role-specific anchors.
- Sentence 2: strongest concrete recent/repeated evidence for the manager/MSP's
  highest-priority requirement(s). Name an employer only when it strengthens
  credibility or gives useful scale/context.
- Sentence 3: cover the next most important requirement cluster with concrete
  evidence (domain, platform, content execution, testing, operations, etc.).
- Optional Sentence 4: one final high-value differentiator only if it adds new
  signal. Do not write a generic sales closer.

REQUIREMENT PRIORITY:
- Start from the CURRENT substantive requirement hierarchy in STRUCTURED
  REQUISITION INTELLIGENCE, not from whichever technologies are most prominent
  on the candidate's resume.
- Highest priority: current manager/MSP/Spotlight selection guidance and why
  prior candidates failed, when available.
- Then formal Must Haves / required competencies / domain requirements.
- Then explicit preferred skills only when they materially strengthen this
  particular candidate.
- Do not spend scarce summary space on adjacent candidate strengths that are not
  important to the current requisition.
- Prefer recurring responsibilities and direct hands-on execution over oversight,
  strategy, or one-time projects unless the role specifically asks for those.

STYLE — VERY IMPORTANT:
- Write like an experienced recruiter sending a concise internal candidate note.
- Use the candidate's FIRST NAME naturally after the first sentence.
- NEVER use honorifics or formal references such as "Mr.", "Ms.", "Mrs.",
  "Dr.", "Mr. Zaman", "Ms. Fisher", "the candidate", etc.
- Preserve employer names exactly as they appear in the resume. Do not silently
  correct/rebrand company names from outside knowledge.
- Use clear business English, not consultant language or inflated adjectives.
- Prefer factual evidence over claims about fit.
- Keep the paragraph skimmable; avoid long comma-heavy inventory sentences.
- Do not simply paraphrase or duplicate the formatted resume Summary.

GOOD MODEL — communications example:
"Kristin is a senior communications specialist with 18+ years of experience
 developing business communications, change communications, and audience-focused
 content across technology and government environments. At CGI, she managed
 enterprise communication channels and repositories including SharePoint for a
 7,000+ employee technology organization and supported high-volume communications
 workstreams. She has independently created newsletters, presentations, web and
 digital content, implementation communications, and executive materials while
 translating complex business and technical information into clear,
 audience-specific messaging. Her hands-on content production, SharePoint
 experience, operational coordination, and technology-change communications
 support the need for a seasoned individual contributor who can independently
 execute communications deliverables."

The example demonstrates the desired DENSITY and recruiter tone. Do not copy its
wording or force its structure onto technical roles.

BANNED / AVOID:
- "strongly matching", "aligning perfectly", "perfectly aligns", "ideal fit",
  "great fit", "strong fit", "well suited", "well-suited"
- "positions him/her/them", "meets all requirements", "directly matches"
- "possesses extensive", "possesses solid expertise", "brings a wealth of"
- "crucial for this role", "critical for this role", "data-intensive role"
- "immediate asset", "Day 1", "hit the ground running"
- Honorifics: "Mr.", "Ms.", "Mrs.", "Dr."
- Generic closing judgments when the evidence itself is stronger

OTHER RULES:
- Never invent or infer candidate experience.
- Do not merge evidence from separate employers into one sentence in a way that
  falsely implies all tools/responsibilities occurred together.
- Do not discuss weaknesses, missing skills, compensation, work authorization,
  citizenship, sponsorship, onsite status, interview availability, submission
  process, VNDLY, MSP, shortlisting strategy, or AI matching strategy.
- AI/GenAI may be mentioned only when it is genuinely a job technology and the
  candidate has supported experience.

======================================================================
======================================================================
TASK 2 — FINAL VNDLY SKILL SHORTLIST (GLOBAL HARD CAP = 8)
======================================================================

Produce ONE ranked shortlist of the BEST VNDLY skills to select for this
candidate submission.

CRITICAL:
- Build a ranked CANDIDATE POOL of up to 16 skills. Python will apply the final
  hard cap of 8 after deterministic evidence-strength validation.
- This is still a HIGH-SIGNAL candidate pool, NOT a comprehensive inventory.
- It is completely acceptable to return fewer than 8 when fewer skills deserve
  a selection.
- Rank every item from highest to lowest submission value.
- IMPORTANT: A technical skill appearing only in the candidate's Professional
  Summary or top-level skills inventory MUST NOT outrank a current Required skill
  that is explicitly demonstrated inside dated work-history roles.
- For example, if Docker appears only in a summary but J2EE/JDBC are explicitly
  shown in dated roles and are current JD requirements, J2EE/JDBC must rank above
  Docker.

PRIORITY HIERARCHY:
1. Exact structured VNDLY Must-Have skill attached to this requisition,
   when the candidate genuinely satisfies the COMPLETE requirement.
2. Explicit current MUST-HAVE / REQUIRED skill from the JD or authoritative
   manager/MSP/Spotlight guidance.
3. Exact structured VNDLY Nice-To-Have skill attached to this requisition,
   when genuinely satisfied.
4. Explicit current PREFERRED / NICE-TO-HAVE skill from the JD or authoritative
   manager/MSP/Spotlight guidance.
5. Additional candidate-supported skill only when it adds unusually useful
   signal for THIS requisition.

SELECTION PRINCIPLES:
- Optimize for SIGNAL, not coverage.
- A skill that directly represents the core requirement is more valuable than
  several narrower synonyms/components of the same concept.
- Prefer the exact terminology Freddie is screening for.
- Prefer skills supported by dated, substantive work history.
- Prefer recent and recurring evidence over a skills-list mention.
- Prefer a broad exact requirement such as "Relational Database" over selecting
  Oracle + DB2 + SQL Server + MySQL individually, unless a specific database is
  itself a critical screening requirement.
- Prefer "Spring Framework" over also consuming another slot with "Spring Boot"
  when the formal requirement is Spring Framework, unless Spring Boot itself is
  separately important to the role.
- Do not spend scarce slots on generic soft skills (Problem Solving, Analytical
  Thinking, Critical Thinking, Writing, Communication, etc.) unless the role or
  manager makes that specific competency unusually central AND it is more
  valuable than a technical/domain requirement.
- Do not spend scarce slots on adjacent candidate strengths that are not central
  to the requisition.
- Do not use one skill as a substitute for a different required skill
  (for example Maven is not proof of Gradle).
- Avoid redundant parent/child/synonym selections. With only 8 final slots,
  coverage of distinct explicit requirements is usually more valuable than a
  second closely related variant. For example, when Spring Framework is already
  selected and JDBC/J2EE are explicit candidate-proven requirements, do not spend
  another slot on Spring Boot merely because it is a candidate strength.
- Candidate evidence must genuinely support every selection.

STRUCTURED-SKILL RULE:
- If source_category is `exact_structured_must`, skill_name MUST come from:
  {selected_must_json}
- If source_category is `exact_structured_nice`, skill_name MUST come from:
  {selected_nice_json}
- Long/legacy/free-text catalogue values are allowed ONLY in those two exact
  structured categories.
- For all other categories, use clean `canonical_preferred` values first and
  `compound_review` only when materially more precise. Never use
  `legacy_free_text`.

SOURCE CATEGORY must be exactly one of:
- exact_structured_must
- exact_structured_nice
- required_job_intelligence
- preferred_job_intelligence
- additional_high_value

BEFORE RETURNING:
Ask yourself: "Which candidate-supported skills deserve to compete for the
final 8 VNDLY selections?" Include enough high-quality Required alternatives for
Python to make the final evidence-strength ranking. Do not omit a clearly
work-history-evidenced Required skill merely to include a weaker summary-only
Required skill.

======================================================================
OUTPUT
======================================================================

Return ONLY:

{{
  "VNDLY_SUMMARY": "",
  "FINAL_RECOMMENDED_VNDLY_SKILLS": [
    {{
      "skill_name": "",
      "source_category": "",
      "evidence": "",
      "reason": ""
    }}
  ]
}}
"""

    raw_package = fred_generate_json(
        api_key,
        package_prompt,
    )

    vndly_summary = fred_clean_vndly_submission_summary(
        raw_package.get("VNDLY_SUMMARY", ""),
        candidate_name,
    )

    if not vndly_summary:
        raise RuntimeError(
            "Gemini returned no VNDLY submission summary."
        )

    ranked_skill_pool = fred_build_final_ranked_vndly_skills(
        raw_package.get(
            "FINAL_RECOMMENDED_VNDLY_SKILLS",
            [],
        ),
        selected_must_items,
        selected_nice_items,
        catalog,
        cap=16,
    )

    ranked_skill_pool = fred_supplement_explicit_required_skills(
        ranked_skill_pool,
        structured_req,
        catalog,
        raw_resume_text,
    )

    evidence_ranked_skills = fred_rerank_vndly_skills_by_evidence(
        ranked_skill_pool,
        raw_resume_text,
        structured_req=structured_req,
        cap=32,
    )

    ranked_skills = fred_apply_vndly_redundancy_control(
        evidence_ranked_skills,
        structured_req,
        cap=FREDDIE_VNDLY_SKILL_CAP,
    )

    grouped_skills = fred_split_ranked_vndly_skills(
        ranked_skills
    )

    return {
        "candidate_positioning_title": str(
            candidate_positioning_title or ""
        ).strip(),
        "vndly_summary": vndly_summary,
        "final_recommended_vndly_skills": ranked_skills,
        **grouped_skills,
    }



def fred_skill_names(items):
    return [
        str(item.get("skill_name", "")).strip()
        for item in (items or [])
        if str(item.get("skill_name", "")).strip()
    ]


def fred_build_submission_email_body(
    candidate_name,
    vndly_context,
    package,
):
    """Plain-text ready-to-use Freddie submission email."""
    vndly_context = vndly_context or {}
    package = package or {}

    job_id = str(vndly_context.get("job_id", "") or "").strip()
    job_title = str(vndly_context.get("job_title", "") or "").strip()
    manager = str(vndly_context.get("resource_manager", "") or "").strip()

    ranked_skills = package.get(
        "final_recommended_vndly_skills",
        [],
    ) or []

    source_labels = {
        "exact_structured_must": "Exact Structured Must",
        "exact_structured_nice": "Exact Structured Nice",
        "required_job_intelligence": "Required",
        "preferred_job_intelligence": "Preferred",
        "additional_high_value": "Additional",
    }

    if ranked_skills:
        skill_lines = []
        for index, item in enumerate(ranked_skills, start=1):
            name = str(item.get("skill_name", "") or "").strip()
            category = str(
                item.get("source_category", "") or ""
            ).strip().lower()
            label = source_labels.get(category, "Recommended")
            if name:
                skill_lines.append(
                    f"{index}. {name} [{label}]"
                )
        skills_text = "\n".join(skill_lines)
    else:
        skills_text = "None recommended"

    lines = [
        f"Candidate: {candidate_name}",
        f"Freddie Job: {job_id}" + (f" - {job_title}" if job_title else ""),
        f"Resource Manager: {manager or 'N/A'}",
        "",
        "VNDLY SUBMISSION SUMMARY",
        str(package.get("vndly_summary", "") or "").strip(),
        "",
        f"RECOMMENDED VNDLY SKILLS - TOP {FREDDIE_VNDLY_SKILL_CAP} MAX",
        skills_text,
    ]

    return "\n".join(lines).strip()


def fred_send_submission_email(
    candidate_name,
    vndly_context,
    package,
):
    """
    Send the Freddie submission metadata through Gmail SMTP.
    Credentials live only in Streamlit Secrets.
    """
    sender = str(
        st.secrets.get(
            "FREDDIE_EMAIL_SENDER",
            "kennyconv@gmail.com",
        )
    ).strip()
    recipient = str(
        st.secrets.get(
            "FREDDIE_EMAIL_RECIPIENT",
            "kkerrigan@conv.com",
        )
    ).strip()
    app_password = str(
        st.secrets.get(
            "FREDDIE_EMAIL_APP_PASSWORD",
            "",
        )
    ).replace(" ", "").strip()

    if not sender or not recipient or not app_password:
        raise RuntimeError(
            "Freddie email Secrets are incomplete. Expected "
            "FREDDIE_EMAIL_SENDER, FREDDIE_EMAIL_RECIPIENT, and "
            "FREDDIE_EMAIL_APP_PASSWORD."
        )

    job_id = str((vndly_context or {}).get("job_id", "") or "").strip()
    job_title = str((vndly_context or {}).get("job_title", "") or "").strip()

    subject_parts = [
        "Freddie VNDLY Submission",
        job_id,
        candidate_name,
    ]
    subject = " - ".join(
        part for part in subject_parts if str(part).strip()
    )

    if job_title:
        subject += f" ({job_title})"

    msg = EmailMessage()
    msg["From"] = sender
    msg["To"] = recipient
    msg["Subject"] = subject
    msg.set_content(
        fred_build_submission_email_body(
            candidate_name,
            vndly_context,
            package,
        )
    )

    with smtplib.SMTP_SSL(
        "smtp.gmail.com",
        465,
        timeout=30,
    ) as smtp:
        smtp.login(sender, app_password)
        smtp.send_message(msg)


def fred_generation_email_key(
    candidate_name,
    job_id,
    raw_resume_text,
    vndly_summary,
):
    """
    Stable session key preventing duplicate sends for the same generated package.
    """
    payload = "||".join(
        [
            str(candidate_name or ""),
            str(job_id or ""),
            str(raw_resume_text or ""),
            str(vndly_summary or ""),
        ]
    )
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def fred_generate_json(api_key, prompt, max_attempts=6):
    """
    Freddie-specific Gemini JSON helper with the same retry philosophy
    already used throughout the app.
    """
    client = genai.Client(api_key=api_key)

    models_to_try = [
        "gemini-2.5-flash",
        "gemini-3-flash-preview",
        "gemini-3.1-flash-lite-preview",
        "gemini-pro-latest",
    ]

    last_error = None

    for attempt in range(max_attempts):
        if attempt == 0:
            current_model = models_to_try[0]
        elif attempt in [1, 2]:
            current_model = models_to_try[1]
        elif attempt in [3, 4]:
            current_model = models_to_try[2]
        else:
            current_model = models_to_try[3]

        try:
            response = client.models.generate_content(
                model=current_model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    response_mime_type="application/json"
                ),
            )

            data = repair_and_load_json(response.text)

            if data:
                return data

        except Exception as e:
            last_error = e

            if "503" in str(e) and attempt < max_attempts - 1:
                time.sleep(2 ** ((attempt % 2) + 1))
                continue

            if attempt < max_attempts - 1:
                time.sleep(1)
                continue

    raise RuntimeError(
        f"Gemini could not return usable Freddie Mac JSON after retries. "
        f"Last error: {last_error}"
    )


def fred_extract_vetting_block(job_description):
    """
    Deterministic first-pass locator for Freddie's Supplier Vetting Questions.
    This does NOT decide what constitutes a question. It only gives Gemini
    a high-signal excerpt while Gemini also receives the complete JD.
    """
    if not job_description:
        return ""

    text = str(job_description).replace("\r\n", "\n").replace("\r", "\n")

    heading_pattern = re.compile(
        r"(?im)^[ \t]*"
        r"(?:required[ \t]+)?"
        r"supplier[ \t]+vetting[ \t]+questions?"
        r"[ \t]*(?:[:\-–—])[ \t]*(.*)?$"
    )

    match = heading_pattern.search(text)

    if not match:
        return ""

    start = match.start()

    # Common Freddie headings that usually terminate the vetting block.
    stop_pattern = re.compile(
        r"(?im)^[ \t]*("
        r"job description|"
        r"position overview|"
        r"position description|"
        r"responsibilities|"
        r"key responsibilities|"
        r"required qualifications|"
        r"preferred qualifications|"
        r"qualifications|"
        r"must have|"
        r"must-have|"
        r"nice to have|"
        r"nice-to-have|"
        r"keys to success|"
        r"key success in the role"
        r")[ \t]*:?[ \t]*$"
    )

    remainder = text[match.end():]
    stop_match = stop_pattern.search(remainder)

    if stop_match:
        end = match.end() + stop_match.start()
    else:
        # Safety cap. Gemini still receives the full JD separately.
        end = min(len(text), start + 5000)

    return text[start:end].strip()


def fred_normalize_question_list(raw_questions):
    """
    Sanitizes Gemini's vetting-question output without rewriting its substance.
    """
    if not isinstance(raw_questions, list):
        return []

    cleaned = []

    for item in raw_questions:
        if isinstance(item, dict):
            question = (
                item.get("question")
                or item.get("text")
                or item.get("instruction")
                or ""
            )
        else:
            question = str(item)

        question = str(question).strip()

        # Strip duplicated numbering only. Do NOT rewrite the wording.
        question = re.sub(
            r"^\s*(?:question\s*)?\d+\s*[\.\)\:\-]\s*",
            "",
            question,
            flags=re.IGNORECASE,
        ).strip()

        if question and question.lower() not in {
            "n/a",
            "na",
            "none",
            "not applicable",
            "no",
        }:
            cleaned.append(question)

    return cleaned


def fred_vetting_response_required(question):
    """
    Determine whether an official Freddie Supplier Vetting item requires
    a candidate response.

    Conservative rule:
    - Explicitly optional / not-required items do NOT require a response.
    - Everything else defaults to REQUIRED.
    """

    text = str(question or "").strip().lower()

    if not text:
        return False

    optional_phrases = [
        "not required",
        "not mandatory",
        "optional",
        "encouraged to submit",
        "encouraged but not required",
        "may submit",
    ]

    if any(phrase in text for phrase in optional_phrases):
        return False

    return True


def fred_source_has_substantive_selection_intelligence(source_text):
    """
    Freddie-only deterministic safeguard.

    Returns True only when MSP/VNDLY notes or Spotlight transcript appear to
    contain substantive candidate-profile / selection guidance rather than
    routine staffing operations such as scheduling, offers, OOO status,
    backfills, reposting, or "JD updated" notices.
    """
    text = str(source_text or "").strip().lower()

    if not text:
        return False

    substantive_patterns = [
        r"\bmanager\b.{0,120}\b(?:need|needs|require|requires|required|want|wants|looking|focus|priority|critical|emphas)",
        r"\b(?:must|need|needs|required|requires|requirement|critical|priority)\b.{0,120}\b(?:candidate|experience|background|skill|technical|functional|business|domain|java|python|angular|sql|mortgage|finance|operations|frontend|backend|full[- ]?stack)",
        r"\b(?:candidate|candidates)\b.{0,120}\b(?:lacked|lack|missing|too|weak|stronger|depth|explain|demonstrate|show|not vetted|move forward|advance)",
        r"\b(?:interview|selection|shortlist)\b.{0,120}\b(?:feedback|lacked|missing|need|needs|require|stronger|depth|explain|demonstrate)",
        r"\b(?:not the focus|no longer required|not required|deemphasized|de-emphasized|instead of|rather than)\b",
        r"\b\d{1,3}\s*%\b.{0,80}\b(?:ui|frontend|front-end|backend|back-end|technical|functional|business)",
    ]

    return any(re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
               for pattern in substantive_patterns)


def fred_cap_skill_years(value, total_experience):
    """
    Freddie-only hard ceiling: no skill/domain tenure may exceed the
    deterministic total professional experience.
    """
    value_text = str(value or "").strip()
    total_text = str(total_experience or "").strip()

    if not value_text or not total_text:
        return value_text

    value_match = re.search(r"(\d+)\s*\+\s*years", value_text, flags=re.IGNORECASE)
    total_match = re.search(r"(\d+)\s*\+\s*years", total_text, flags=re.IGNORECASE)

    if not value_match or not total_match:
        return value_text

    skill_years = int(value_match.group(1))
    total_years = int(total_match.group(1))

    if skill_years <= total_years:
        return value_text

    start, end = value_match.span(1)
    return value_text[:start] + str(total_years) + value_text[end:]


@st.cache_data(ttl=3600, show_spinner=False)
def fred_analyze_requisition(
    _api_key,
    job_description,
    msp_notes="",
    spotlight_transcript="",
    vndly_context=None,
):
    """
    Analyze the Freddie requisition ONCE and return structured job intelligence.

    Freddie-only design:
    - Official VNDLY JD + structured VNDLY fields establish formal requirements.
    - Dated Spotlight / MSP notes establish current manager emphasis and changes.
    - Internal Notes provide recruiter context, but cannot manufacture requirements.
    - Operational staffing noise is deliberately separated from selection intelligence.

    Cached by the JD/notes/transcript/context so Streamlit does not re-call Gemini
    whenever a recruiter types into another field.
    """
    vndly_context = vndly_context or {}

    if not job_description or not str(job_description).strip():
        return {
            "target_title": "",
            "explicit_must_have_requirements": [],
            "must_have_competency_groups": [],
            "required_requirements": [],
            "preferred_requirements": [],
            "domain_requirements": [],
            "current_manager_priorities": [],
            "selection_feedback": [],
            "role_clarifications": [],
            "role_evolution": [],
            "deemphasized_or_negated_requirements": [],
            "supplier_vetting_questions": [],
        }

    vetting_hint = fred_extract_vetting_block(job_description)

    # Only the role-relevant structured VNDLY fields are sent to Gemini here.
    # Competition/admin fields remain available in the app but do not become
    # candidate requirements.
    matching_context = {
        "job_id": vndly_context.get("job_id", ""),
        "job_title": vndly_context.get("job_title", ""),
        "resource_manager": vndly_context.get("resource_manager", ""),
        "organization_unit": vndly_context.get("organization_unit", ""),
        "cost_center": vndly_context.get("cost_center", ""),
        "workday_cost_center": vndly_context.get("workday_cost_center", ""),
        "job_category": vndly_context.get("job_category", ""),
        "min_experience": vndly_context.get("min_experience", ""),
        "max_experience": vndly_context.get("max_experience", ""),
        "must_have_skills": vndly_context.get("must_have_skills", ""),
        "nice_to_have_skills": vndly_context.get("nice_to_have_skills", ""),
        "work_site_name": vndly_context.get("work_site_name", ""),
        "city": vndly_context.get("city", ""),
        "state": vndly_context.get("state", ""),
        "time_type": vndly_context.get("time_type", ""),
        "mnpi": vndly_context.get("mnpi", ""),
        "internal_notes": vndly_context.get("internal_notes", ""),
    }

    matching_context_json = json.dumps(
        matching_context,
        indent=2,
        ensure_ascii=False,
        default=str,
    )

    prompt = f"""
Return a valid JSON object ONLY.

You are analyzing a Freddie Mac contingent-worker requisition from Workday VNDLY.

Your job in this step is NOT to evaluate a candidate and NOT to write resume
content. Your only job is to convert the requisition into accurate structured
job intelligence that will later be used for candidate matching.

======================================================================
SOURCE HIERARCHY
======================================================================

You are receiving FOUR kinds of role intelligence:

1. OFFICIAL VNDLY JOB DESCRIPTION
2. STRUCTURED VNDLY JOB FIELDS
3. DATED MSP / VNDLY NOTES and SPOTLIGHT CALL TRANSCRIPTS
4. INTERNAL RECRUITER NOTES, if present

Use them differently.

FORMAL REQUIREMENTS:
- The Official VNDLY Job Description and populated structured VNDLY
  Must Have Skills / Nice To Have Skills / experience fields establish formal
  job requirements.
- Structured "Must Have Skills" are formal requirements when populated.
- Structured "Nice To Have Skills" are preferred requirements when populated.
- Do NOT promote a Nice To Have into a Must Have.
- First look in the JD for Freddie's explicitly labeled "Must Have
  Qualifications", "Must Have", "Must-Have", or equivalent section.
- Treat explicitly labeled Must Haves as primary formal matching criteria.
- Classify requirements at the INDIVIDUAL CLAUSE level, even when several
  requirements appear inside one paragraph.
- Wording such as "preferred", "nice to have", "a plus", "desired", or
  "preferred but not required" belongs under preferred_requirements even if
  physically located inside a Must Have paragraph.
- Only clearly mandatory items may appear in explicit_must_have_requirements.
- Scan the ENTIRE Official VNDLY Job Description for explicit mandatory language,
  not only the formally labeled Must Have section. Any statement anywhere in the
  JD using language such as "must have", "required", "minimum", "requires",
  "need(s)", or an explicit minimum-years requirement must be retained as a
  formal requirement when it clearly applies to the candidate profile.
- Example: if Responsibilities says "Must have 8+ years of experience", include
  that exact requirement in explicit_must_have_requirements even if it appears
  outside the labeled Must Have Qualifications section.
- If there is no explicit Must Have section, derive core required competencies
  from Required Qualifications and central recurring responsibilities.

ACTUAL HIRING-MANAGER / SELECTION EMPHASIS:
- Spotlight calls and role-relevant MSP/VNDLY clarifications can reveal which
  formal requirements matter most in practice.
- PROVENANCE GATE: current_manager_priorities, selection_feedback, and
  role_clarifications may ONLY contain information explicitly supported by the
  DATED MSP / VNDLY NOTES or DATED SPOTLIGHT CALL TRANSCRIPT.
- NEVER populate those three fields by copying, paraphrasing, or inferring from
  the Official VNDLY Job Description or structured VNDLY fields alone.
- If the notes/transcript contain only operational updates and no substantive
  manager/selection guidance, return [] for all three fields.
- Give especially high weight to explicit feedback explaining WHY prior
  candidates were not shortlisted, did not advance, or failed interviews, and
  what future submissions must demonstrate.
- Examples of high-signal selection intelligence include:
  "too IT-PM heavy", "needs stronger Angular", "lacked technical depth",
  "must be able to explain technologies on the resume", "mortgage required",
  "business operations PM is the critical requirement", or "model building is
  not the focus".
- If the manager explicitly says a JD item is not needed, outdated, optional,
  incorrectly stated, or no longer relevant, place it under
  deemphasized_or_negated_requirements.
- Do NOT silently erase formal JD requirements. Keep formal requirements and
  current manager priorities as separate concepts.

======================================================================
RECENCY / ROLE-EVOLUTION RULE — CRITICAL
======================================================================

The MSP/VNDLY Notes and Spotlight Call Transcript may contain multiple dated
entries over the lifecycle of one requisition.

- Treat the date attached to each entry as meaningful.
- When two entries conflict, a NEWER explicit hiring-manager clarification or
  selection-feedback statement normally represents the current direction.
- Do NOT let an older Spotlight statement override a newer manager/MSP
  clarification.
- When a newer update merely reports scheduling or staffing activity and does
  NOT change the candidate profile, it must NOT override older substantive role
  guidance.
- Capture only SUBSTANTIVE candidate-profile changes over time in role_evolution.
- HARD PROVENANCE GATE: role_evolution may ONLY be populated when the DATED
  MSP / VNDLY NOTES or DATED SPOTLIGHT CALL TRANSCRIPT explicitly states WHAT
  changed in the desired candidate profile.
- Do NOT create a role_evolution item merely because a note says the job was
  updated, refreshed, reposted, re-released, backfilled, replaced, or that a
  prior offer was rejected. Those are operational/history facts unless the
  source itself states the substantive change.
- NEVER infer that an updated/reposted/backfill JD "reflects the latest candidate
  profile" or similar. If the source does not name the changed skill, weighting,
  scope, domain, seniority, responsibility, or manager expectation, return [].
- A role_evolution item is appropriate only when the source identifies a
  meaningful change in required skills, weighting, scope, domain background,
  seniority/experience, responsibilities, or manager expectations.
- Example of VALID evolution: "Manager clarified the role is 40% UI / 60%
  backend and now requires stronger Angular."
- Example of INVALID evolution: "Backfill release; JD has been updated" when no
  substantive change is described.
- current_manager_priorities must represent what appears to matter NOW after
  considering chronology.
- role_clarifications should capture concise current clarifications such as
  "40% UI / 60% backend" or "business operations PM, not IT PM".
- selection_feedback should capture evidence about why prior submissions or
  interviewees did or did not advance.

======================================================================
OPERATIONAL-NOISE FILTER — CRITICAL
======================================================================

MSP/VNDLY notes mix job-selection intelligence with staffing-process updates.

Do NOT turn any of the following into candidate requirements unless the wording
explicitly changes or clarifies the desired candidate profile:

- manager OOO / availability
- interview scheduling or placeholders
- debrief scheduling
- shortlisting deadlines by themselves
- offer status / offer processing
- candidate counts
- positions filled / slots added
- keeping candidates warm
- supplier administration
- vendor distribution timing
- role hold / closure risk
- routine follow-up status
- billing/rate administration

A sentence such as "interviews scheduled Monday" is operational noise.

A sentence such as "manager rejected candidates because they were too
backend-heavy and now requires stronger Angular" is HIGH-SIGNAL selection
intelligence.

======================================================================
INTERNAL NOTES RULE
======================================================================

Internal Notes are recruiter context, not an authoritative Freddie source.

Allowed:
- use them to resolve ambiguity,
- surface known recruiter context,
- improve emphasis when consistent with official/manager evidence.

Not allowed:
- manufacture a formal requirement,
- override an explicit hiring-manager statement with speculation,
- assign an unsupported skill to the role merely because an internal note
  speculates about it.

======================================================================
STRUCTURED VNDLY CONTEXT RULE
======================================================================

The structured VNDLY fields may help interpret the role.

- Job Title, Min/Max Experience, Must Have Skills, Nice To Have Skills are
  direct matching inputs.
- Organization Unit, Cost Center, Workday Cost Center, Job Category, Work Site,
  Time Type, and MNPI are CONTEXTUAL evidence only.
- Contextual metadata may help interpret an ambiguous JD, but it cannot create
  a candidate requirement absent from formal or manager guidance.
- Resource Manager is identification/context only.

======================================================================
MUST-HAVE COMPETENCY GROUPS
======================================================================

Create up to FOUR logical competency groups representing the highest-signal
current Must Have requirements.

Examples:
"Application Development (Python, Java, Spring Boot)"
"Data & Analytics (SQL, Snowflake, JSON)"
"ServiceNow Testing & ATF"
"Mortgage / Financial Services"

Rules:
- Group closely related Freddie requirements together.
- Preserve Freddie's exact terminology wherever practical.
- Do not invent a technology or competency.
- These are JOB requirements only. Do not evaluate a candidate in this step.
- A group may contain technologies, methodologies, domain expertise, or
  operational competencies.
- Rank groups using the current requirement hierarchy after applying substantive
  dated manager clarifications.

Return each group as:
{{
    "name": "",
    "terms": [],
    "reason": ""
}}

======================================================================
ROLE EVOLUTION FORMAT
======================================================================

Only include material changes or clarifications.

Each role_evolution item must be:
{{
    "date": "YYYY-MM-DD or best available date",
    "change": "Concise description of the substantive change/clarification"
}}

Do NOT add routine operational updates to role_evolution.

======================================================================
SUPPLIER VETTING QUESTIONS — ZERO INVENTION
======================================================================

Extract ONLY supplier vetting questions/instructions actually supplied by
Freddie/MSP in the Official Job Description.

Rules:
- Preserve substantive wording and original order.
- Remove only leading numbering such as "1." or "Question 1:".
- Correcting obvious whitespace is allowed.
- Do NOT rewrite, improve, expand, split, or invent questions.
- If Freddie provides one assessment instruction rather than conventional
  questions, return that instruction as ONE item.
- If the JD says Supplier Vetting Questions are N/A / None, return [].
- Do NOT mistake interview information, candidate-template questions,
  qualifications, or ordinary JD bullets for Supplier Vetting Questions.
- The standard candidate-template questions at the top of Freddie JDs
  (Former CW, work authorization, location, etc.) are NOT supplier vetting
  questions.

A deterministic parser found this possible vetting block. It is ONLY a hint;
validate it against the full JD:

--- POSSIBLE VETTING BLOCK ---
{vetting_hint}
--- END POSSIBLE VETTING BLOCK ---

======================================================================
OUTPUT FORMAT
======================================================================

Return ONLY:

{{
  "target_title": "",
  "explicit_must_have_requirements": [""],
  "must_have_competency_groups": [
    {{
      "name": "",
      "terms": [""],
      "reason": ""
    }}
  ],
  "required_requirements": [""],
  "preferred_requirements": [""],
  "domain_requirements": [""],
  "current_manager_priorities": [""],
  "selection_feedback": [""],
  "role_clarifications": [""],
  "role_evolution": [
    {{
      "date": "",
      "change": ""
    }}
  ],
  "deemphasized_or_negated_requirements": [""],
  "supplier_vetting_questions": [""]
}}

======================================================================
INPUTS
======================================================================

STRUCTURED VNDLY JOB FIELDS:
{matching_context_json}

OFFICIAL VNDLY JOB DESCRIPTION:
{job_description}

DATED MSP / VNDLY NOTES:
{msp_notes}

DATED SPOTLIGHT CALL TRANSCRIPT:
{spotlight_transcript}
"""

    data = fred_generate_json(_api_key, prompt)

    required_keys = {
        "target_title",
        "explicit_must_have_requirements",
        "must_have_competency_groups",
        "required_requirements",
        "preferred_requirements",
        "domain_requirements",
        "current_manager_priorities",
        "selection_feedback",
        "role_clarifications",
        "role_evolution",
        "deemphasized_or_negated_requirements",
        "supplier_vetting_questions",
    }

    missing_keys = required_keys - set(data.keys())

    if missing_keys:
        raise RuntimeError(
            "Freddie requisition analysis returned incomplete JSON. "
            f"Missing fields: {', '.join(sorted(missing_keys))}"
        )

    data["supplier_vetting_questions"] = fred_normalize_question_list(
        data.get("supplier_vetting_questions", [])
    )

    for list_key in [
        "explicit_must_have_requirements",
        "required_requirements",
        "preferred_requirements",
        "domain_requirements",
        "current_manager_priorities",
        "selection_feedback",
        "role_clarifications",
        "deemphasized_or_negated_requirements",
    ]:
        value = data.get(list_key, [])
        data[list_key] = value if isinstance(value, list) else []

    role_evolution = data.get("role_evolution", [])
    if not isinstance(role_evolution, list):
        role_evolution = []

    normalized_evolution = []
    for item in role_evolution:
        if not isinstance(item, dict):
            continue

        change = str(item.get("change", "") or "").strip()
        change_date = str(item.get("date", "") or "").strip()

        if change:
            normalized_evolution.append(
                {
                    "date": change_date,
                    "change": change,
                }
            )

    data["role_evolution"] = normalized_evolution

    # ================================================================
    # FREDDIE-ONLY PROVENANCE SAFEGUARD
    # ================================================================
    # Manager-priority / selection / clarification / evolution fields are
    # allowed only when MSP notes or the Spotlight transcript actually contain
    # substantive candidate-profile guidance. Formal JD requirements remain in
    # their own formal-requirement buckets.
    substantive_manager_source = (
        fred_source_has_substantive_selection_intelligence(msp_notes)
        or fred_source_has_substantive_selection_intelligence(spotlight_transcript)
    )

    if not substantive_manager_source:
        data["current_manager_priorities"] = []
        data["selection_feedback"] = []
        data["role_clarifications"] = []
        data["role_evolution"] = []

    groups = data.get("must_have_competency_groups", [])
    if not isinstance(groups, list):
        groups = []

    normalized_groups = []

    for group in groups[:4]:
        if not isinstance(group, dict):
            continue

        name = str(group.get("name", "")).strip()
        terms = group.get("terms", [])

        if not isinstance(terms, list):
            terms = []

        terms = [str(x).strip() for x in terms if str(x).strip()]

        if name:
            normalized_groups.append(
                {
                    "name": name,
                    "terms": terms,
                    "reason": str(group.get("reason", "")).strip(),
                }
            )

    data["must_have_competency_groups"] = normalized_groups

    return data


def fred_standardize_dates(date_str):
    """
    Standardize Freddie Mac work-history dates while preserving source precision:
    MMM YYYY - MMM YYYY
    MMM YYYY - CURRENT
    YYYY - YYYY
    or YYYY - CURRENT when the source provides years only.

    Never invent a month for a year-only source date.

    Examples:
    2026-02 - Current      -> FEB 2026 - CURRENT
    2021-08 - 2026-02      -> AUG 2021 - FEB 2026
    07/2019 - 08/2021      -> JUL 2019 - AUG 2021
    February 2026 - Present -> FEB 2026 - CURRENT
    """
    if not date_str:
        return ""

    text = str(date_str)

    # Normalize line breaks, spacing, and dash characters.
    text = text.replace("\r", " ").replace("\n", " ")
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"\s+", " ", text).strip()

    month_names = {
        1: "JAN",
        2: "FEB",
        3: "MAR",
        4: "APR",
        5: "MAY",
        6: "JUN",
        7: "JUL",
        8: "AUG",
        9: "SEP",
        10: "OCT",
        11: "NOV",
        12: "DEC",
    }

    text_months = {
        "january": "JAN",
        "jan": "JAN",
        "february": "FEB",
        "feb": "FEB",
        "march": "MAR",
        "mar": "MAR",
        "april": "APR",
        "apr": "APR",
        "may": "MAY",
        "june": "JUN",
        "jun": "JUN",
        "july": "JUL",
        "jul": "JUL",
        "august": "AUG",
        "aug": "AUG",
        "september": "SEP",
        "sept": "SEP",
        "sep": "SEP",
        "october": "OCT",
        "oct": "OCT",
        "november": "NOV",
        "nov": "NOV",
        "december": "DEC",
        "dec": "DEC",
    }

    # Current / Present
    text = re.sub(
        r"\b(?:current|present)\b",
        "CURRENT",
        text,
        flags=re.IGNORECASE,
    )

    # YYYY-MM or YYYY/MM -> MMM YYYY
    def replace_year_month(match):
        year = int(match.group(1))
        month = int(match.group(2))
        return f"{month_names[month]} {year}"

    text = re.sub(
        r"\b((?:19|20)\d{2})[-/](0?[1-9]|1[0-2])\b",
        replace_year_month,
        text,
    )

    # MM/YYYY -> MMM YYYY
    def replace_month_year(match):
        month = int(match.group(1))
        year = int(match.group(2))
        return f"{month_names[month]} {year}"

    text = re.sub(
        r"\b(0?[1-9]|1[0-2])/((?:19|20)\d{2})\b",
        replace_month_year,
        text,
    )

    # Written month names -> three-letter uppercase months.
    month_pattern = (
        r"\b("
        r"January|Jan|February|Feb|March|Mar|April|Apr|May|"
        r"June|Jun|July|Jul|August|Aug|September|Sept|Sep|"
        r"October|Oct|November|Nov|December|Dec"
        r")\b"
    )

    text = re.sub(
        month_pattern,
        lambda m: text_months[m.group(1).lower()],
        text,
        flags=re.IGNORECASE,
    )

    # Normalize spacing around the range dash.
    text = re.sub(r"\s*-\s*", " - ", text)

    return text.strip()


def fred_calculate_total_experience(experience, current_date):
    """
    Calculate total professional experience deterministically from dated roles.

    Overlapping employment periods are merged so they are not double-counted.
    Returns the conservative completed-years format used in the Freddie
    summary, e.g. "7+ years".
    """

    month_lookup = {
        "JAN": 1,
        "FEB": 2,
        "MAR": 3,
        "APR": 4,
        "MAY": 5,
        "JUN": 6,
        "JUL": 7,
        "AUG": 8,
        "SEP": 9,
        "OCT": 10,
        "NOV": 11,
        "DEC": 12,
    }

    def parse_month_year(value):
        value = str(value).strip().upper()

        if value in {"CURRENT", "PRESENT"}:
            return current_date.year, current_date.month

        match = re.fullmatch(
            r"(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s+"
            r"((?:19|20)\d{2})",
            value,
        )

        if not match:
            return None

        month = month_lookup[match.group(1)]
        year = int(match.group(2))

        return year, month

    intervals = []

    for role in experience:
        dates = str(role.get("Dates", "") or "").strip()

        if not dates:
            continue

        match = re.fullmatch(
            r"\s*(.+?)\s+-\s+(.+?)\s*",
            dates,
        )

        if not match:
            continue

        start_parsed = parse_month_year(match.group(1))
        end_parsed = parse_month_year(match.group(2))

        if not start_parsed or not end_parsed:
            continue

        start_year, start_month = start_parsed
        end_year, end_month = end_parsed

        start_index = (start_year * 12) + (start_month - 1)
        end_index = (end_year * 12) + (end_month - 1)

        if end_index < start_index:
            continue

        intervals.append((start_index, end_index))

    if not intervals:
        return ""

    intervals.sort()

    merged = []

    for start, end in intervals:
        if not merged or start > merged[-1][1]:
            merged.append([start, end])
        else:
            merged[-1][1] = max(
                merged[-1][1],
                end,
            )

    total_months = sum(
        end - start
        for start, end in merged
    )

    completed_years = total_months // 12

    if completed_years < 1:
        return "<1 year"

    return f"{completed_years}+ years"


def fred_role_date_interval(dates, current_date):
    """
    Parse a Freddie-standardized role date range into conservative month indexes.

    Month/year dates use the stated month.
    Year-only dates never invent January for display; for calculations they use
    conservative bounds:
      - start year only -> December of that year
      - end year only   -> January of that year
    This avoids overstating tenure when month precision is unknown.
    """
    if not dates:
        return None

    text = str(dates).strip().upper()
    match = re.fullmatch(r"\s*(.+?)\s+-\s+(.+?)\s*", text)
    if not match:
        return None

    month_lookup = {
        "JAN": 1, "FEB": 2, "MAR": 3, "APR": 4, "MAY": 5, "JUN": 6,
        "JUL": 7, "AUG": 8, "SEP": 9, "OCT": 10, "NOV": 11, "DEC": 12,
    }

    def parse_part(value, is_start):
        value = value.strip().upper()

        if value in {"CURRENT", "PRESENT"}:
            return current_date.year, current_date.month, True

        m = re.fullmatch(
            r"(JAN|FEB|MAR|APR|MAY|JUN|JUL|AUG|SEP|OCT|NOV|DEC)\s+((?:19|20)\d{2})",
            value,
        )
        if m:
            return int(m.group(2)), month_lookup[m.group(1)], False

        m = re.fullmatch(r"((?:19|20)\d{2})", value)
        if m:
            year = int(m.group(1))
            # Conservative unknown-month assumptions for calculation only.
            return year, (12 if is_start else 1), False

        return None

    start = parse_part(match.group(1), True)
    end = parse_part(match.group(2), False)

    if not start or not end:
        return None

    sy, sm, _ = start
    ey, em, end_is_current = end

    start_idx = sy * 12 + (sm - 1)
    end_idx = ey * 12 + (em - 1)

    if end_idx < start_idx:
        return None

    return start_idx, end_idx, end_is_current, ey


def fred_skill_label_components(skill_label):
    """Return named supporting technologies from parentheses, if present."""
    label = str(skill_label or "").strip()
    if not label:
        return []

    components = []
    for part in re.findall(r"\(([^)]*)\)", label):
        for token in re.split(r"[,/&;+]", part):
            token = token.strip()
            if token:
                components.append(token)
    return components


def fred_role_text(role):
    """Flatten one structured role into searchable evidence text."""
    bullets = role.get("Bullets", []) or []
    if isinstance(bullets, str):
        bullets_text = bullets
    else:
        bullets_text = " ".join(str(x) for x in bullets)

    return " ".join(
        [
            str(role.get("Company", "") or ""),
            str(role.get("Title", "") or ""),
            bullets_text,
            str(role.get("Environment", "") or ""),
        ]
    ).casefold()


def fred_role_supports_skill_component(role, component):
    """Conservative alias check against one dated structured role."""
    role_text = fred_role_text(role)
    key = fred_normalize_skill_name(component)

    aliases = {
        "j2ee": ["j2ee", "java ee", "jee"],
        "microservices": ["microservice"],
        "spring boot": ["spring boot"],
        "spring framework": ["spring framework", "spring "],
        "java": ["java"],
        "oracle": ["oracle"],
        "db2": ["db2"],
        "sql server": ["sql server"],
        "microsoft sql server": ["sql server"],
        "sql": ["sql", "pl/sql"],
        "junit": ["junit"],
        "junit testing": ["junit"],
        "mockito": ["mockito"],
        "git": ["git", "gitlab", "github"],
        "maven": ["maven"],
        "jenkins": ["jenkins"],
        "gradle": ["gradle"],
        "jdbc": ["jdbc", "java database connectivity"],
        "jsp": ["jsp", "jsps", "javaserver pages", "java server pages"],
        "servlets": ["servlet"],
        "banking": ["bank", "banking"],
        "investment": ["investment", "portfolio", "asset", "securit", "trade ", "trading"],
        "securities": ["securit", "trade ", "trading", "investment"],
    }

    terms = aliases.get(key)
    if terms:
        return any(term in role_text for term in terms)

    core = re.sub(r"[^a-z0-9+#./ -]", "", key).strip()
    return bool(core and len(core) >= 4 and core in role_text)


def fred_role_supports_skill_core(role, skill_label):
    """
    Determine whether a role supports the CORE competency represented by a
    Freddie Skills-table row. Parenthetical tools are examples/evidence, not an
    AND checklist that must all coexist in the same job.
    """
    label = re.sub(r"\([^)]*\)", "", str(skill_label or "")).strip().casefold()
    role_text = fred_role_text(role)

    # High-confidence competency families used by the Freddie skills table.
    family_aliases = [
        (["financial", "banking", "investment", "securities"],
         ["bank", "banking", "financial", "investment", "portfolio", "asset", "securit", "trade ", "trading", "retirement"]),
        (["relational database", "database", "sql"],
         ["sql", "pl/sql", "oracle", "db2", "sql server", "mysql", "postgresql", "sybase", "jdbc", "database"]),
        (["java", "j2ee", "spring"],
         ["java", "j2ee", "spring", "servlet", "jsp", "jdbc", "microservice"]),
        (["testing", "test", "quality"],
         ["junit", "mockito", "unit test", "testing", "test case"]),
        (["devops", "build", "deployment", "version control", "ci/cd"],
         ["git", "gitlab", "github", "maven", "gradle", "jenkins", "docker", "ci/cd", "build", "deploy"]),
        (["sharepoint"], ["sharepoint"]),
        (["communications", "content"], ["communication", "content", "newsletter", "publication", "writing", "editing"]),
        (["change management", "change communications", "adoption"], ["change", "adoption", "readiness", "implementation communication"]),
    ]

    for label_terms, evidence_terms in family_aliases:
        if any(term in label for term in label_terms):
            return any(term in role_text for term in evidence_terms)

    # Generic fallback: any explicitly named parenthetical component supports
    # the row as an AREA competency; we do not require every component at once.
    components = fred_skill_label_components(skill_label)
    if components and any(
        fred_role_supports_skill_component(role, component)
        for component in components
    ):
        return True

    # Last-resort lexical support for a concise core label.
    core_words = [
        w for w in re.findall(r"[a-z0-9+#.]+", label)
        if len(w) >= 4 and w not in {"development", "experience", "background", "expertise"}
    ]
    return bool(core_words and any(word in role_text for word in core_words))


def fred_filter_skill_role_indexes_by_label(
    skill_label,
    experience,
    role_indexes,
):
    """
    Build the dated role set for a Freddie Skills-table row based on the CORE
    competency, not an artificial requirement that every parenthetical tool
    appear in the same job.

    Gemini's proposed role indexes remain useful for domain/context judgments,
    but Python also adds any dated role that clearly supports the core area.
    """
    proposed = fred_clean_skill_role_indexes(role_indexes, len(experience))
    inferred = []

    for idx, role in enumerate(experience, start=1):
        if fred_role_supports_skill_core(role, skill_label):
            inferred.append(idx)

    # Preserve Gemini's domain-aware choices and add deterministic evidence.
    combined = []
    for idx in proposed + inferred:
        if idx not in combined:
            combined.append(idx)

    return combined


def fred_calculate_skill_years_from_roles(
    experience,
    role_indexes,
    current_date,
):
    """
    Deterministically calculate one Freddie Skills-row YEARS value from only the
    dated roles that Gemini identified as supporting the ENTIRE skill label.
    Overlapping roles are merged so months are not double-counted.
    """
    if not isinstance(role_indexes, list):
        return ""

    clean_indexes = []
    for value in role_indexes:
        try:
            idx = int(value)
        except (TypeError, ValueError):
            continue
        if 1 <= idx <= len(experience) and idx not in clean_indexes:
            clean_indexes.append(idx)

    if not clean_indexes:
        return ""

    intervals = []
    latest_year = None
    used_current = False

    for idx in clean_indexes:
        role = experience[idx - 1]
        parsed = fred_role_date_interval(
            role.get("Dates", ""),
            current_date,
        )
        if not parsed:
            continue

        start_idx, end_idx, end_is_current, end_year = parsed
        intervals.append((start_idx, end_idx))

        if end_is_current:
            used_current = True

        if latest_year is None or end_year > latest_year:
            latest_year = end_year

    if not intervals:
        return ""

    intervals.sort()
    merged = [list(intervals[0])]

    for start_idx, end_idx in intervals[1:]:
        last = merged[-1]
        if start_idx <= last[1] + 1:
            last[1] = max(last[1], end_idx)
        else:
            merged.append([start_idx, end_idx])

    # Inclusive role months.
    total_months = sum((end - start + 1) for start, end in merged)
    completed_years = total_months // 12

    last_used = "current" if used_current else str(latest_year or "").strip()

    if completed_years < 1:
        return f"<1 year, {last_used}" if last_used else "<1 year"

    return (
        f"{completed_years}+ years, {last_used}"
        if last_used
        else f"{completed_years}+ years"
    )


def fred_clean_skill_role_indexes(value, max_roles):
    """Return unique valid 1-based role indexes from Gemini output."""
    if not isinstance(value, list):
        return []

    result = []
    for item in value:
        try:
            idx = int(item)
        except (TypeError, ValueError):
            continue
        if 1 <= idx <= max_roles and idx not in result:
            result.append(idx)
    return result


def fred_clean_summary(summary, candidate_name):
    """
    Final deterministic cleanup for Freddie Mac summaries.

    - Forces the candidate's name to use the same Title Case spelling
      extracted from the resume.
    - Removes/replaces recruiter-style phrases Gemini was explicitly
      instructed not to use.
    """
    if not summary:
        return ""

    cleaned = str(summary).strip()

    # ------------------------------------------------------------
    # 1. NORMALIZE CANDIDATE NAME CAPITALIZATION
    # ------------------------------------------------------------

    if candidate_name:
        # Replace any capitalization variation of the full candidate name
        # with the normalized Title Case version extracted from the resume.
        cleaned = re.sub(
            re.escape(candidate_name),
            candidate_name,
            cleaned,
            flags=re.IGNORECASE,
        )

    # ------------------------------------------------------------
    # 2. REMOVE / REPLACE BANNED RECRUITER PHRASES
    # ------------------------------------------------------------

    phrase_replacements = {
        "proven history in": "experience in",
        "proven history of": "experience with",
        "proven history": "experience",
    }

    for old_phrase, new_phrase in phrase_replacements.items():
        cleaned = re.sub(
            re.escape(old_phrase),
            new_phrase,
            cleaned,
            flags=re.IGNORECASE,
        )

    # Clean up any accidental double spaces.
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()

    return cleaned


def fred_safe_reorder_bullets(bullets, requested_order):
    """
    Reorder existing bullets ONLY.

    Gemini returns zero-based indexes. This function refuses to manufacture,
    delete, or rewrite any bullet. Invalid indexes are ignored and every
    unranked original bullet is appended in its original order.
    """
    if not isinstance(bullets, list):
        return bullets

    if not isinstance(requested_order, list):
        return bullets

    valid_order = []
    seen = set()

    for value in requested_order:
        try:
            idx = int(value)
        except (TypeError, ValueError):
            continue

        if 0 <= idx < len(bullets) and idx not in seen:
            valid_order.append(idx)
            seen.add(idx)

    if not valid_order:
        return bullets

    # Preserve every original bullet.
    valid_order.extend(
        idx for idx in range(len(bullets))
        if idx not in seen
    )

    return [bullets[idx] for idx in valid_order]


def fred_remove_table(table):
    tbl = table._element
    parent = tbl.getparent()

    if parent is not None:
        parent.remove(tbl)


def fred_find_table_containing(doc, text):
    target = str(text).lower()

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if target in cell.text.lower():
                    return table

    return None


def fred_get_reference_font(paragraph):
    """
    Capture the formatting already assigned to {{VETTING_QA}} in Word so
    generated Q&A inherits the template's typography.
    """
    if paragraph.runs:
        run = paragraph.runs[0]

        return {
            "name": run.font.name,
            "size": run.font.size,
            "color": (
                run.font.color.rgb
                if run.font.color and run.font.color.rgb
                else None
            ),
        }

    return {
        "name": None,
        "size": None,
        "color": None,
    }


def fred_apply_reference_font(run, reference, bold=None):
    if reference.get("name"):
        run.font.name = reference["name"]

    if reference.get("size"):
        run.font.size = reference["size"]

    if reference.get("color"):
        run.font.color.rgb = reference["color"]

    if bold is not None:
        run.bold = bold


def fred_insert_vetting_qa(cell, vetting_pairs):
    """
    Replace {{VETTING_QA}} with a dynamic number of bold questions and
    normal-weight candidate answers.
    """
    placeholder_paragraph = None

    for p in cell.paragraphs:
        if "{{vetting_qa}}" in p.text.lower():
            placeholder_paragraph = p
            break

    if placeholder_paragraph is None:
        return

    reference = fred_get_reference_font(placeholder_paragraph)

    # Clear the placeholder paragraph while keeping the paragraph itself.
    for run in list(placeholder_paragraph.runs):
        run._element.getparent().remove(run._element)

    placeholder_paragraph.paragraph_format.space_before = Pt(0)
    placeholder_paragraph.paragraph_format.space_after = Pt(2)

    for idx, pair in enumerate(vetting_pairs, start=1):
        question = str(pair.get("question", "")).strip()
        answer = str(pair.get("answer", "")).strip()

        if idx == 1:
            question_p = placeholder_paragraph
        else:
            question_p = cell.add_paragraph()

        question_p.paragraph_format.space_before = Pt(4 if idx > 1 else 0)
        question_p.paragraph_format.space_after = Pt(1)

        q_run = question_p.add_run(f"{idx}. {question}")
        fred_apply_reference_font(q_run, reference, bold=True)

        answer_p = cell.add_paragraph()
        answer_p.paragraph_format.left_indent = Pt(12)
        answer_p.paragraph_format.space_before = Pt(0)
        answer_p.paragraph_format.space_after = Pt(3)

        a_run = answer_p.add_run(answer)
        fred_apply_reference_font(a_run, reference, bold=False)


def process_freddie_word_doc(
    temp_path,
    mapping,
    vetting_pairs,
    out_path,
):
    """
    Freddie-specific Word processor.

    We intentionally keep this separate from process_word_doc() so the
    successful Fannie/other-client Q1-A1 behavior remains untouched.
    """
    doc = docx.Document(temp_path)
    expand_experience_placeholders(doc, mapping)
    ensure_experience_spacing(doc, mapping)

    # ================================================================
    # 1. SUPPLIER VETTING QUESTIONS
    # ================================================================

    vetting_table = fred_find_table_containing(doc, "{{VETTING_QA}}")

    if vetting_table is not None:
        if vetting_pairs:
            inserted = False

            for row in vetting_table.rows:
                for cell in row.cells:
                    if "{{vetting_qa}}" in cell.text.lower():
                        fred_insert_vetting_qa(cell, vetting_pairs)
                        inserted = True
                        break

                if inserted:
                    break
        else:
            # No official Freddie vetting questions = remove entire section.
            fred_remove_table(vetting_table)

    # ================================================================
    # 2. CERTIFICATIONS
    # ================================================================

    certifications = str(mapping.get("Certifications", "") or "").strip()

    if not certifications:
        cert_table = fred_find_table_containing(doc, "{{CERTIFICATIONS}}")

        if cert_table is not None:
            fred_remove_table(cert_table)

    # ================================================================
    # 3. EDUCATION
    # ================================================================

    has_any_education = any(
        str(mapping.get(f"School{i}", "") or "").strip()
        for i in range(1, 4)
    )

    education_table = fred_find_table_containing(doc, "{{School1}}")

    if education_table is not None:
        if not has_any_education:
            fred_remove_table(education_table)
        else:
            rows_to_delete = []

            for row in education_table.rows:
                row_text = " ".join(
                    cell.text for cell in row.cells
                ).lower()

                for i in range(1, 4):
                    if (
                        f"{{{{school{i}}}}}".lower() in row_text
                        and not str(mapping.get(f"School{i}", "") or "").strip()
                    ):
                        rows_to_delete.append(row)
                        break

            for row in rows_to_delete:
                try:
                    education_table._tbl.remove(row._tr)
                except Exception:
                    pass

    # ================================================================
    # 4. SKILLS — REMOVE UNUSED ROWS INSTEAD OF LEAVING BLANKS
    # ================================================================

    skills_table = fred_find_table_containing(doc, "{{SKILL1}}")

    if skills_table is not None:
        rows_to_delete = []

        for row in skills_table.rows:
            row_text = " ".join(
                cell.text for cell in row.cells
            ).lower()

            for i in range(1, 5):
                if (
                    f"{{{{skill{i}}}}}".lower() in row_text
                    and not str(mapping.get(f"SKILL{i}", "") or "").strip()
                ):
                    rows_to_delete.append(row)
                    break

        for row in rows_to_delete:
            try:
                skills_table._tbl.remove(row._tr)
            except Exception:
                pass

    # ================================================================
    # 5. REPLACE TAGS IN BODY PARAGRAPHS
    # ================================================================

    paras_to_remove = []

    kill_keys = [
        "Company",
        "Title",
        "Bullets",
        "Dates",
        "Summary",
        "Skill",
        "Years",
        "Certification",
        "School",
    ]

    for p in list(doc.paragraphs):
        for key, value in mapping.items():
            tag = f"{{{{{key}}}}}"

            if tag.lower() not in p.text.lower():
                continue

            if not value or not str(value).strip():
                if any(k.lower() in key.lower() for k in kill_keys):
                    if p not in paras_to_remove:
                        paras_to_remove.append(p)
                else:
                    replace_tag_safely(p, tag, "")
                continue

            if "Bullets" in key:
                for run in p.runs:
                    if "\n" in run.text:
                        run.text = run.text.replace("\n", "")

                lines = [
                    line.strip()
                    for line in str(value).split("\n")
                    if line.strip()
                ]

                if not lines:
                    if p not in paras_to_remove:
                        paras_to_remove.append(p)
                    continue

                replace_tag_safely(p, tag, lines[0])

                curr_p = p

                for line in lines[1:]:
                    curr_p = insert_bullet_line_after(
                        curr_p,
                        line,
                    )

                # Preserve explicit Environment / Technologies lines only,
                # following the same rule as your existing formatter.
                num_match = re.search(r"\d+", key)

                if num_match:
                    env_val = mapping.get(
                        f"Environment{num_match.group()}"
                    )

                    if env_val and str(env_val).strip():
                        env_p = curr_p.insert_paragraph_before("")
                        curr_p._p.addnext(env_p._p)

                        try:
                            env_p.style = doc.styles["Normal"]
                        except Exception:
                            pass

                        env_p.paragraph_format.left_indent = Pt(0)
                        env_p.paragraph_format.space_after = Pt(0)

                        clean_env = re.sub(
                            r"^Environment\s*:\s*",
                            "",
                            str(env_val).strip(),
                            flags=re.IGNORECASE,
                        )

                        b_run = env_p.add_run("Environment: ")
                        b_run.bold = True
                        b_run.font.name = "Times New Roman"
                        b_run.font.size = Pt(12)

                        n_run = env_p.add_run(clean_env)
                        n_run.bold = False
                        n_run.font.name = "Times New Roman"
                        n_run.font.size = Pt(12)

                        curr_p = env_p

                spacer = curr_p.insert_paragraph_before("")
                curr_p._p.addnext(spacer._p)

                try:
                    spacer.style = doc.styles["Normal"]
                except Exception:
                    pass

                spacer.paragraph_format.space_after = Pt(0)
                spacer.paragraph_format.space_before = Pt(0)

            else:
                replace_tag_safely(
                    p,
                    tag,
                    str(value).strip(),
                )

    # ================================================================
    # 6. REPLACE TAGS INSIDE TABLE CELLS
    # ================================================================

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in mapping.items():
                        tag = f"{{{{{key}}}}}"

                        if tag.lower() in p.text.lower():
                            replace_tag_safely(
                                p,
                                tag,
                                str(value) if value else "",
                            )

    # ================================================================
    # 7. DELETE EMPTY WORK-HISTORY PARAGRAPHS
    # ================================================================

    for p in paras_to_remove:
        try:
            delete_paragraph(p)
        except Exception:
            pass

    # Collapse extra blank lines created when optional Freddie sections were removed.
    collapse_extra_blank_paragraphs(doc)

    doc.save(out_path)

    return out_path


# ====================================================================
# --- 🟣 CLIENT APP 1.5: FREDDIE MAC 🟣 ---
# ====================================================================

def freddie_mac_app():
    TEMPLATE_FILENAME = "Freddie_Mac_Template.docx"

    st.title("Freddie Mac Precision Extractor")

    # ================================================================
    # API CONFIGURATION
    # ================================================================

    with st.sidebar:
        st.header("🔑 API Configuration")

        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input(
                "Gemini API Key",
                type="password",
                key="fred_api_key",
            )

            st.info(
                "Paste your Gemini API key here to run the tool."
            )

    # ================================================================
    # CANDIDATE INFORMATION
    # ================================================================

    st.header("📋 Candidate Information")

    # Initialize dependent fields before widgets are rendered.
    if "fred_site" not in st.session_state:
        st.session_state.fred_site = ""

    if "fred_elig" not in st.session_state:
        st.session_state.fred_elig = "N/A"

    def fred_sync_onsite():
        if st.session_state.get("fred_onsite") == "No":
            st.session_state.fred_site = "N/A"
        elif st.session_state.get("fred_site") == "N/A":
            st.session_state.fred_site = ""

    def fred_sync_former_cw():
        if st.session_state.get("fred_form") == "No":
            st.session_state.fred_elig = "N/A"
        elif st.session_state.get("fred_elig") == "N/A":
            st.session_state.fred_elig = "No"

    col1, col2 = st.columns(2)

    with col1:
        Current_Location = st.text_input(
            "Current Location:",
            key="fred_loc",
        )

        Auth_US = st.selectbox(
            "Do you currently possess unrestricted lawful authorization "
            "to work in the U.S indefinitely?",
            ["Yes", "No"],
            index=0,
            key="fred_auth",
        )

        with st.container(border=True):
            Available_Onsite = st.selectbox(
                "Available to be onsite (Yes / No):",
                ["Yes", "No"],
                index=0,
                key="fred_onsite",
                on_change=fred_sync_onsite,
            )

            Site = st.text_input(
                "If yes, at which locations:",
                key="fred_site",
            )

    with col2:
        Interview = st.text_input(
            "Interview Availability for next 7 days:",
            key="fred_int",
        )

        Sponsorship = st.selectbox(
            "Will you now or in the future require sponsorship for an "
            "immigration-related employment benefit?",
            ["Yes", "No"],
            index=1,
            key="fred_spon",
        )

        with st.container(border=True):
            Former_CW = st.selectbox(
                "Former CW/Employee?",
                ["Yes", "No"],
                index=1,
                key="fred_form",
                on_change=fred_sync_former_cw,
            )

            Eligibility = st.selectbox(
                "If Yes, Rehire Eligibility Check Completed?",
                ["Yes", "No", "N/A"],
                key="fred_elig",
            )

    # ================================================================
    # REQUISITION / JOB DESCRIPTION
    # ================================================================

    st.header("📝 Freddie Mac Requisition")

    # Freddie ONLY: private authenticated connection to the VNDLY/API
    # Google Sheet. Fannie/Deloitte/other client sheet logic remains untouched.
    SHEET_URL = (
        "https://docs.google.com/spreadsheets/d/"
        "1yMQmqYRl_wa6ynXZrORCMOFwHRPwC2GVFKCTGzMfAe4/"
        "edit?gid=0"
    )

    def load_freddie_reqs(url):
        try:
            conn = st.connection(
                "gsheets",
                type=GSheetsConnection,
            )

            df = conn.read(
                spreadsheet=url,
                worksheet="VNDLY Jobs",
                ttl=600,
            )

            return df.fillna("")

        except Exception as e:
            st.error(
                "❌ Unable to load the private Freddie Mac VNDLY sheet. "
                f"{str(e)}"
            )
            return pd.DataFrame()

    df_reqs = load_freddie_reqs(SHEET_URL)

    options = ["None"]
    req_mapping = {}

    if not df_reqs.empty:
        for idx, row in df_reqs.iterrows():
            req_id = str(row.get("Job ID", "")).replace(".0", "").strip()
            title = str(row.get("Job Title", "")).strip()
            manager = str(row.get("Resource Manager", "")).strip()
            status = str(row.get("Status", "")).strip()

            # MSP-only requisitions that were never distributed to Convergenz
            # remain in the VNDLY tracker but must not appear in the
            # Resume Formatter requisition dropdown.
            if status.lower() in {"don't have", "dont have"}:
                continue

            if not req_id:
                continue

            desc = str(row.get("Official Job Description", ""))
            notes = str(row.get("VNDLY / MSP Notes", ""))
            transcript = str(row.get("Spotlight Call Transcript", ""))
            internal_notes = str(row.get("Internal Notes", ""))

            # Clean VNDLY-native dropdown: Job ID - Job Title - Resource Manager.
            option_parts = [req_id, title, manager]
            option_str = " - ".join(
                [part.strip() for part in option_parts if part and part.strip()]
            )

            if option_str and option_str != "-":
                options.append(option_str)

                req_mapping[option_str] = {
                    "job_id": req_id,
                    "status": status,
                    "job_title": title,
                    "resource_manager": manager,
                    "job_posted_date": str(row.get("Job Posted Date", "")),
                    "distributed_to_convergenz": str(
                        row.get("Distributed to Convergenz", "")
                    ),
                    "distribution_lag_hours": row.get(
                        "Distribution Lag (Hours)",
                        "",
                    ),
                    "distribution_timing": str(
                        row.get("Distribution Timing", "")
                    ),
                    "open_positions": row.get("Open Positions", ""),
                    "total_positions": row.get("Total Positions", ""),
                    "pending_positions": row.get("Pending Positions", ""),
                    "max_submissions_per_vendor": row.get(
                        "Max Submissions / Vendor",
                        "",
                    ),
                    "total_supplier_candidates": row.get(
                        "Total Supplier Candidates",
                        "",
                    ),
                    "min_suppliers_submitted": row.get(
                        "Min Suppliers Submitted",
                        "",
                    ),
                    "closed_date": str(row.get("Closed Date", "")),
                    "reason_for_ending": str(row.get("Reason for Ending", "")),
                    "organization_unit": str(
                        row.get("Organization Unit", "")
                    ),
                    "cost_center": str(row.get("Cost Center", "")),
                    "workday_cost_center": str(
                        row.get("Workday Cost Center", "")
                    ),
                    "work_site_name": str(row.get("Work Site Name", "")),
                    "city": str(row.get("City", "")),
                    "state": str(row.get("State", "")),
                    "start_date": str(row.get("Start Date", "")),
                    "end_date": str(row.get("End Date", "")),
                    "min_experience": row.get("Min Experience", ""),
                    "max_experience": row.get("Max Experience", ""),
                    "must_have_skills": str(row.get("Must Have Skills", "")),
                    "nice_to_have_skills": str(
                        row.get("Nice To Have Skills", "")
                    ),
                    "rate": row.get("Rate", ""),
                    "max_bill_rate": row.get("Max Bill Rate", ""),
                    "created_on": str(row.get("Created On", "")),
                    "job_approved_date": str(
                        row.get("Job Approved Date", "")
                    ),
                    "status_date": str(row.get("Status Date", "")),
                    "job_category": str(row.get("Job Category", "")),
                    "time_type": str(row.get("Time Type", "")),
                    "single_sourced": str(row.get("Single Sourced", "")),
                    "intake_call_completed": str(
                        row.get("Intake Call Completed", "")
                    ),
                    "supplier_call_conducted": str(
                        row.get("Supplier Call Conducted", "")
                    ),
                    "mnpi": str(row.get("MNPI", "")),
                    "last_api_refresh": str(
                        row.get("Last API Refresh", "")
                    ),
                    "desc": desc,
                    "notes": notes,
                    "transcript": transcript,
                    "internal_notes": internal_notes,
                }

    if "fred_jd" not in st.session_state:
        st.session_state.fred_jd = ""

    if "fred_notes" not in st.session_state:
        st.session_state.fred_notes = ""

    if "fred_trans" not in st.session_state:
        st.session_state.fred_trans = ""

    if "fred_internal_notes" not in st.session_state:
        st.session_state.fred_internal_notes = ""

    def update_fred_jd_text():
        selected = st.session_state.get(
            "fred_req_selector",
            "None",
        )

        # Clear answers whenever the requisition changes so an answer
        # from Job A can NEVER accidentally survive into Job B.
        for key in list(st.session_state.keys()):
            if key.startswith("fred_vetting_answer_"):
                del st.session_state[key]

        if (
            selected != "None"
            and selected in req_mapping
        ):
            req = req_mapping[selected]

            st.session_state.fred_jd = (
                req.get("desc", "").strip()
            )

            st.session_state.fred_notes = (
                req.get("notes", "").strip()
            )

            st.session_state.fred_trans = (
                req.get("transcript", "").strip()
            )

            st.session_state.fred_internal_notes = (
                req.get("internal_notes", "").strip()
            )

        else:
            st.session_state.fred_jd = ""
            st.session_state.fred_notes = ""
            st.session_state.fred_trans = ""
            st.session_state.fred_internal_notes = ""

    selected_req = st.selectbox(
        "🔍 Select a Freddie Mac Requisition (Type to search):",
        options=options,
        key="fred_req_selector",
        on_change=update_fred_jd_text,
    )

    selected_vndly_context = (
        req_mapping.get(selected_req, {})
        if selected_req != "None"
        else {}
    )

    # Operational VNDLY metadata is useful to the recruiter, but most of it
    # should not become candidate claims. Keep it in a collapsed snapshot.
    if selected_vndly_context:
        with st.expander(
            "📊 VNDLY Requisition Snapshot",
            expanded=False,
        ):
            snapshot_lines = []

            def add_snapshot(label, value):
                value = str(value or "").strip()
                if value:
                    snapshot_lines.append(f"**{label}:** {value}")

            add_snapshot("Job ID", selected_vndly_context.get("job_id"))
            add_snapshot("Status", selected_vndly_context.get("status"))
            add_snapshot("Job Title", selected_vndly_context.get("job_title"))
            add_snapshot(
                "Resource Manager",
                selected_vndly_context.get("resource_manager"),
            )
            add_snapshot(
                "Organization Unit",
                selected_vndly_context.get("organization_unit"),
            )
            add_snapshot(
                "Cost Center",
                selected_vndly_context.get("cost_center"),
            )
            add_snapshot(
                "Work Site",
                selected_vndly_context.get("work_site_name"),
            )
            add_snapshot(
                "Location",
                ", ".join(
                    [
                        value
                        for value in [
                            str(selected_vndly_context.get("city", "")).strip(),
                            str(selected_vndly_context.get("state", "")).strip(),
                        ]
                        if value
                    ]
                ),
            )
            add_snapshot(
                "Positions",
                (
                    f"{selected_vndly_context.get('open_positions', '')} open / "
                    f"{selected_vndly_context.get('total_positions', '')} total"
                ).strip(),
            )
            add_snapshot(
                "Max Submissions / Vendor",
                selected_vndly_context.get("max_submissions_per_vendor"),
            )
            add_snapshot(
                "Total Supplier Candidates",
                selected_vndly_context.get("total_supplier_candidates"),
            )
            add_snapshot(
                "Distributed to Convergenz",
                selected_vndly_context.get("distributed_to_convergenz"),
            )
            add_snapshot(
                "End Date",
                selected_vndly_context.get("end_date"),
            )
            add_snapshot(
                "Max Bill Rate",
                selected_vndly_context.get("max_bill_rate"),
            )

            if snapshot_lines:
                st.markdown("\n\n".join(snapshot_lines))

    Job_Description = st.text_area(
        "1. Official VNDLY Job Description:",
        height=150,
        key="fred_jd",
    )

    MSP_Notes = st.text_area(
        "2. VNDLY / MSP Notes (Optional):",
        height=120,
        key="fred_notes",
    )

    Call_Transcript = st.text_area(
        "3. Spotlight Call Transcript "
        "(Optional - Hiring Manager Priority):",
        height=120,
        key="fred_trans",
    )

    # Internal Notes remain recruiter context. They are visible/editable so the
    # recruiter knows exactly what context is being supplied to the analyzer.
    Internal_Notes = st.text_area(
        "4. Internal Notes (Optional - Recruiter Context):",
        height=80,
        key="fred_internal_notes",
    )

    # Use the currently displayed Internal Notes value in the analysis context,
    # allowing a recruiter to add/adjust context without editing the source sheet.
    selected_vndly_context = dict(selected_vndly_context)
    selected_vndly_context["internal_notes"] = Internal_Notes

    # ================================================================
    # STRUCTURED REQUISITION ANALYSIS
    # ================================================================

    req_analysis = {
        "target_title": "",
        "explicit_must_have_requirements": [],
        "must_have_competency_groups": [],
        "required_requirements": [],
        "preferred_requirements": [],
        "domain_requirements": [],
        "current_manager_priorities": [],
        "selection_feedback": [],
        "role_clarifications": [],
        "role_evolution": [],
        "deemphasized_or_negated_requirements": [],
        "supplier_vetting_questions": [],
    }

    analysis_error = None

    if (
        API_KEY
        and Job_Description
        and Job_Description.strip()
    ):
        try:
            with st.spinner(
                "Analyzing Freddie Mac requirements "
                "and vetting questions..."
            ):
                req_analysis = fred_analyze_requisition(
                    API_KEY,
                    Job_Description,
                    MSP_Notes,
                    Call_Transcript,
                    selected_vndly_context,
                )

        except Exception as e:
            analysis_error = str(e)

            st.error(
                "❌ Could not analyze the Freddie Mac requisition. "
                f"{analysis_error}"
            )

    # Optional transparency for recruiters.
    if req_analysis.get("explicit_must_have_requirements"):
        with st.expander(
            "🎯 Freddie Requirement Analysis",
            expanded=False,
        ):
            st.markdown("**Explicit Must-Have Requirements:**")
            for requirement in req_analysis[
                "explicit_must_have_requirements"
            ]:
                st.markdown(f"- {requirement}")

            current_priorities = req_analysis.get(
                "current_manager_priorities",
                [],
            )
            if current_priorities:
                st.markdown("**Current Manager Priorities:**")
                for item in current_priorities:
                    st.markdown(f"- {item}")

            clarifications = req_analysis.get(
                "role_clarifications",
                [],
            )
            if clarifications:
                st.markdown("**Latest Role Clarifications:**")
                for item in clarifications:
                    st.markdown(f"- {item}")

            selection_feedback = req_analysis.get(
                "selection_feedback",
                [],
            )
            if selection_feedback:
                st.markdown("**Selection / Interview Feedback:**")
                for item in selection_feedback:
                    st.markdown(f"- {item}")

            evolution = req_analysis.get(
                "role_evolution",
                [],
            )
            if evolution:
                st.markdown("**Role Evolution:**")
                for item in evolution:
                    if isinstance(item, dict):
                        date_label = str(item.get("date", "") or "").strip()
                        change = str(item.get("change", "") or "").strip()
                        if change:
                            prefix = f"{date_label}: " if date_label else ""
                            st.markdown(f"- {prefix}{change}")

            groups = req_analysis.get(
                "must_have_competency_groups",
                [],
            )

            if groups:
                st.markdown("**Planned Skills Priorities:**")

                for idx, group in enumerate(
                    groups,
                    start=1,
                ):
                    terms = ", ".join(
                        group.get("terms", [])
                    )

                    if terms:
                        st.markdown(
                            f"{idx}. **{group.get('name', '')}** "
                            f"— {terms}"
                        )
                    else:
                        st.markdown(
                            f"{idx}. **{group.get('name', '')}**"
                        )

    # ================================================================
    # DYNAMIC OFFICIAL SUPPLIER VETTING QUESTIONS
    # ================================================================

    st.header("🎤 Supplier Vetting Questions")

    vetting_questions = req_analysis.get(
        "supplier_vetting_questions",
        [],
    )

    vetting_answers = []
    vetting_required_flags = []

    if analysis_error:
        st.warning(
            "Vetting questions cannot be safely determined until "
            "the requisition analysis succeeds."
        )

    elif not Job_Description.strip():
        st.info(
            "Select a Freddie Mac requisition above. "
            "Any official supplier vetting questions will appear here "
            "automatically."
        )

    elif not vetting_questions:
        st.success(
            "✅ No official Supplier Vetting Questions were identified "
            "for this requisition."
        )

    else:
        st.caption(
            "Questions below were extracted from Freddie Mac's VNDLY "
            "requisition. Enter the candidate's direct responses. "
            "The question wording is intentionally read-only."
        )

        safe_req_key = re.sub(
            r"[^A-Za-z0-9_\-]+",
            "_",
            selected_req or "manual"
        )
        
        for idx, question in enumerate(
            vetting_questions,
            start=1,
        ):
            response_required = fred_vetting_response_required(question)
            vetting_required_flags.append(response_required)

            if response_required:
                st.markdown(
                    f"**{idx}. {question}**"
                )
                response_label = f"Candidate Response {idx}"
                response_placeholder = (
                    "Paste the candidate's direct response here..."
                )
            else:
                st.markdown(
                    f"**{idx}. {question}**  \n"
                    "*Optional — Freddie/MSP does not require a response.*"
                )
                response_label = f"Candidate Response {idx} (Optional)"
                response_placeholder = (
                    "Optional — leave blank if nothing is being submitted."
                )

            answer = st.text_area(
                response_label,
                height=100,
                key=f"fred_vetting_answer_{safe_req_key}_{idx}",
                placeholder=response_placeholder,
            )

            vetting_answers.append(answer)

    # ================================================================
    # FILE UPLOAD
    # ================================================================

    st.header("📂 File Uploads")

    resume_file = st.file_uploader(
        "📤 Upload the candidate's resume...",
        type=["pdf", "docx", "doc"],
        key="fred_res",
    )

    # ================================================================
    # GENERATE SUBMISSION
    # ================================================================

    if st.button(
        "🚀 Generate Freddie Mac Submission",
        type="primary",
    ):

        # ------------------------------------------------------------
        # PRE-GENERATION VALIDATION
        # ------------------------------------------------------------

        validation_errors = []

        if not API_KEY:
            validation_errors.append(
                "Gemini API Key is missing."
            )

        if not Job_Description.strip():
            validation_errors.append(
                "A Freddie Mac requisition / Job Description "
                "must be selected."
            )

        if analysis_error:
            validation_errors.append(
                "The requisition analysis did not complete successfully."
            )

        if not Current_Location.strip():
            validation_errors.append(
                "Current Location is required."
            )

        if (
            Available_Onsite == "Yes"
            and (
                not Site.strip()
                or Site.strip().upper() == "N/A"
            )
        ):
            validation_errors.append(
                "Enter the location(s) where the candidate "
                "can be onsite."
            )

        if not Interview.strip():
            validation_errors.append(
                "Interview Availability for the next 7 days "
                "is required."
            )

        if (
            Former_CW == "Yes"
            and Eligibility == "N/A"
        ):
            validation_errors.append(
                "Rehire Eligibility cannot be N/A when the "
                "candidate is a former Freddie Mac CW/employee."
            )

        if not resume_file:
            validation_errors.append(
                "Upload the candidate's resume."
            )

        if not os.path.exists(TEMPLATE_FILENAME):
            validation_errors.append(
                f"The template '{TEMPLATE_FILENAME}' was not found."
            )

        if vetting_questions:
            missing_answers = [
                str(idx)
                for idx, (answer, response_required) in enumerate(
                    zip(
                        vetting_answers,
                        vetting_required_flags,
                    ),
                    start=1,
                )
                if response_required and not str(answer).strip()
            ]

            if missing_answers:
                validation_errors.append(
                    "Candidate responses are required for the mandatory "
                    "Supplier Vetting Question(s). Missing response(s): "
                    + ", ".join(missing_answers)
                )

        if validation_errors:
            for error in validation_errors:
                st.error(f"❌ {error}")

            return

        # ------------------------------------------------------------
        # PROCESS
        # ------------------------------------------------------------

        with st.spinner(
            "Building optimized Freddie Mac submission..."
        ):
            resume_path = None

            try:
                # ====================================================
                # SAVE / EXTRACT ORIGINAL RESUME
                # ====================================================

                safe_original_name = os.path.basename(
                    resume_file.name
                )

                resume_path = (
                    f"temp_freddie_{safe_original_name}"
                )

                with open(resume_path, "wb") as f:
                    f.write(
                        resume_file.getbuffer()
                    )

                raw_text = extract_text(resume_path)

                # ====================================================
                # PASS 1 — STRUCTURED RESUME EXTRACTION
                # ====================================================

                extraction_prompt = f"""
Return a valid JSON object ONLY.

You are extracting a candidate's ORIGINAL resume into a clean structured
format for a Freddie Mac VNDLY submission.

CRITICAL:
- This is an extraction step, NOT a tailoring step.
- Preserve the candidate's factual work history.
- Do NOT add JD terminology.
- Do NOT rewrite accomplishments to sound more relevant.
- Do NOT invent technologies, metrics, employers, dates, or responsibilities.

RULES:

1. NAME
- Pull from the resume header.
- Title Case.
- If completely missing/unreadable, use the filename if possible.

2. COMPANY
- Extract the primary END-CLIENT / employer name actually represented by the
  resume.
- Strip geographic locations that are merely appended to the company line.
- If a contracting vendor and end-client are clearly shown together, preserve
  the primary end-client used by the resume experience.
- Do not guess an end-client that is not stated.

3. EDUCATION
- Extract school and degree exactly enough to preserve meaning.

4. EXPERIENCE
Return EVERY dated professional role shown anywhere in the resume, including
condensed or abbreviated entries under headings such as "Earlier Experience",
"Additional Experience", or similar sections. Do not stop after the first seven
roles. Older abbreviated roles may have an empty Bullets array and blank
Environment; they must still be returned when Company/Title/Dates are present so
total career experience can be calculated correctly.

For each role return:
- Company
- Title
- Bullets as an ARRAY
- Environment
- Dates

TITLE:
- Remove only employment-type modifiers such as "(Contract)" or "- Contract"
  when they are merely suffixes.
- Do not upgrade or alter seniority.

BULLETS:
- Preserve every substantive original bullet.
- Light cleanup of malformed bullet characters/spacing is allowed.
- Keep bullets in ORIGINAL ORDER in this extraction response.
- Do NOT rewrite bullets to match Freddie's JD.

ENVIRONMENT:
- Populate ONLY when the original resume explicitly has an "Environment:"
  or "Technologies:" line.
- Never compile an Environment line from ordinary bullets.

DATES:
- Preserve the source dates faithfully enough for later date calculations.
- NEVER invent a month that the source resume does not provide.
- If a role is shown only as years (for example "2018 - 2020"), return only
  those years. Do NOT convert it to "January 2018 - January 2020", "01/2018",
  or any other assumed month.
- Mixed precision must also stay mixed (for example "2018 - March 2020").

5. CERTIFICATIONS
- Extract certifications into one comma-separated string.
- Preserve completion status exactly when the resume distinguishes completed
  certifications from credentials that are still "In Progress", "Pursuing",
  "Expected", "Pending", or otherwise not yet earned.
- NEVER present an in-progress or pending credential as completed.
- For an unfinished credential, append the status directly to the credential
  name using this format:
  "PMP — In Progress"
  "CBAP — In Progress"
- If the resume explicitly labels a group of certifications as "In Progress",
  apply "— In Progress" to every credential under that heading until the next
  clear section heading or status change.
- Do not invent an in-progress status for completed credentials.
- Do not remove a credential solely because it is in progress; preserve it with
  the explicit status.
- If none are found, return "".

CERTIFICATION TRUTH CHECK:
Before returning Certifications, compare every credential against the original
resume and confirm that no credential labeled In Progress/Pursuing/Pending has
been output as though it were completed.

OUTPUT:

{{
  "FullName": "",
  "Certifications": "",
  "Education": [
    {{
      "School": "",
      "Degree": ""
    }}
  ],
  "Experience": [
    {{
      "Company": "",
      "Title": "",
      "Bullets": [
        ""
      ],
      "Environment": "",
      "Dates": ""
    }}
  ]
}}

FILENAME:
{resume_file.name}

ORIGINAL RESUME:
{raw_text}
"""

                data = fred_generate_json(
                    API_KEY,
                    extraction_prompt,
                )

                name = str(
                    data.get(
                        "FullName",
                        "",
                    )
                ).strip().title()

                certifications = str(
                    data.get(
                        "Certifications",
                        "",
                    )
                ).strip()

                education = data.get(
                    "Education",
                    [],
                )

                if not isinstance(
                    education,
                    list,
                ):
                    education = []

                experience = data.get(
                    "Experience",
                    [],
                )

                if not isinstance(
                    experience,
                    list,
                ):
                    experience = []

                # Normalize ALL extracted dated roles first. Older abbreviated
                # "Earlier Experience" entries are retained here for accurate
                # Preserve every dated role for both total-career experience math
                # and the generated Freddie work-history section.
                normalized_experience = []

                for role in experience:
                    if not isinstance(
                        role,
                        dict,
                    ):
                        continue

                    bullets = role.get(
                        "Bullets",
                        [],
                    )

                    if isinstance(
                        bullets,
                        str,
                    ):
                        bullets = [
                            x.strip()
                            for x in bullets.split("\n")
                            if x.strip()
                        ]

                    if not isinstance(
                        bullets,
                        list,
                    ):
                        bullets = []

                    bullets = [
                        re.sub(
                            r"^[\s\-\•\*\d\.\)]+",
                            "",
                            str(x),
                        ).strip()
                        for x in bullets
                        if str(x).strip()
                    ]

                    normalized_experience.append(
                        {
                            "Company": str(
                                role.get(
                                    "Company",
                                    "",
                                )
                            ).strip(),
                            "Title": re.sub(
                                r"\s*\(.*$",
                                "",
                                str(
                                    role.get(
                                        "Title",
                                        "",
                                    )
                                ),
                            ).strip(),
                            "Bullets": bullets,
                            "Environment": str(
                                role.get(
                                    "Environment",
                                    "",
                                )
                            ).strip(),
                            "Dates": fred_standardize_dates(
                                str(
                                    role.get(
                                        "Dates",
                                        "",
                                    )
                                ).strip()
                            ),
                        }
                    )

                all_experience_for_years = normalized_experience
                experience = normalized_experience

                # ====================================================
                # OFFICIAL VETTING Q&A STRUCTURE
                # ====================================================

                vetting_pairs = [
                    {
                        "question": question,
                        "answer": answer.strip(),
                    }
                    for question, answer in zip(
                        vetting_questions,
                        vetting_answers,
                    )
                    if str(answer).strip()
                ]

                vetting_for_ai = "\n\n".join(
                    [
                        f"Q{idx}: {pair['question']}\n"
                        f"A{idx}: {pair['answer']}"
                        for idx, pair in enumerate(
                            vetting_pairs,
                            start=1,
                        )
                    ]
                )

                if not vetting_for_ai:
                    vetting_for_ai = (
                        "No candidate-supplied Supplier Vetting responses "
                        "were provided or required for this requisition."
                    )

                # ====================================================
                # PASS 2 — CANDIDATE MATCH / SUMMARY / SKILLS /
                #          SAFE BULLET PRIORITIZATION
                # ====================================================

                structured_req = json.dumps(
                    req_analysis,
                    indent=2,
                    ensure_ascii=False,
                )

                structured_exp = json.dumps(
                    experience,
                    indent=2,
                    ensure_ascii=False,
                )

                current_date_obj = date.today()
                current_date = current_date_obj.isoformat()

                calculated_total_experience = fred_calculate_total_experience(
                    all_experience_for_years,
                    current_date_obj,
                )

                tailoring_prompt = f"""
Return a valid JSON object ONLY.

You are an elite Senior Technical Recruiter preparing a Freddie Mac
contingent-worker submission through Workday VNDLY.

CURRENT DATE:
{current_date}

PYTHON-CALCULATED TOTAL PROFESSIONAL EXPERIENCE:
{calculated_total_experience}

The total professional experience value above was calculated deterministically
from the candidate's dated work history.

CRITICAL:
- If PYTHON-CALCULATED TOTAL PROFESSIONAL EXPERIENCE is populated, use that
  EXACT value when stating total career experience in the Summary.
- Do NOT independently recalculate it.
- Do NOT round it up.
- Do NOT change "7+ years" to "8+ years", "over 7 years", "nearly 8 years",
  or any other variation.
- If the calculated value is blank, do not invent a total-years figure.

There are THREE sequential audiences for this submission:

1. VNDLY / AUTOMATED MATCHING
   Make the candidate's genuine match to Freddie's highest-priority
   requirements easy for an automated matching system to detect.

2. MSP HUMAN SHORTLIST REVIEW
   Assume an MSP recruiter understands the requisition and staffing process
   but may NOT have the hiring manager's depth of technical knowledge.
   The MSP reviewer should be able to determine within approximately
   20-30 seconds:
   - Does this candidate clearly possess Freddie's most important Must Haves?
   - Is the required specialized experience genuinely demonstrated?
   - Is the required business/domain background present?
   - Is there concrete evidence behind the keywords?
   - Is there any obvious reason NOT to shortlist the candidate?

3. FREDDIE MAC HIRING MANAGER REVIEW
   If shortlisted, preserve enough specific technical and operational evidence
   for the hiring manager to immediately understand the candidate's depth and
   want to interview them.

DUAL-AUDIENCE WRITING PRINCIPLE:
- Do NOT dumb down or remove meaningful technical evidence.
- Instead, TRANSLATE specialized technical work into clear functional language
  that an MSP recruiter can understand, then support it with the exact
  technologies, methodologies, and domain terminology that matter to Freddie.
- The MSP recruiter should understand WHY the experience is relevant without
  needing to understand every underlying technology.
- The hiring manager should still see HOW the candidate performed the work and
  which technologies were actually used.
- Optimize for:
  clear match first,
  technical proof second,
  promotional language never.

======================================================================
HIERARCHY OF ROLE REQUIREMENTS
======================================================================

Use the ALREADY-ANALYZED structured requisition intelligence.

For EMPHASIS and PRESENT-DAY selection strategy, use this order:

1. Latest explicit hiring-manager correction / selection feedback that
   substantively changes or clarifies what future candidates must demonstrate
2. Current hiring-manager priorities / latest substantive Spotlight or MSP
   clarification
3. Freddie's formal Must Have Qualifications and populated structured VNDLY
   Must Have Skills
4. Detailed Required Qualifications and core recurring responsibilities
5. Preferred / Nice-to-Have qualifications

CRITICAL:
- "Latest" matters only when the newer information is substantive role or
  selection guidance. A newer scheduling/offer/OOO update does NOT override an
  older substantive manager requirement.
- Use selection_feedback and role_clarifications to understand why previous
  candidates failed and what must be made obvious in a new submission.
- Use role_evolution to understand how the target profile changed over time.
- Never highlight a requirement the manager explicitly negated.
- Formal Must Haves still require strong truthful coverage unless explicitly
  negated or clarified by authoritative manager guidance.
- Operational staffing noise must have ZERO influence on candidate claims.
- Internal recruiter notes may guide emphasis but may not manufacture a
  requirement or override explicit Freddie guidance.

The requisition has ALREADY been analyzed into structured job intelligence.
Do not independently reinvent its hierarchy.

STRUCTURED REQUISITION:
{structured_req}

======================================================================
VNDLY / MATCHING STRATEGY
======================================================================

When the candidate genuinely possesses a Freddie Must Have:

- Use Freddie's exact skill/technology/domain terminology when supported.
- Prefer the exact Freddie term over a synonym when both are truthful.
- Do NOT merely repeat keywords.
- Pair important terms with contextual evidence of how the candidate actually
  used them.
- Favor demonstrated skills in work-experience bullets over a standalone skills
  inventory.
- Favor recent, repeated, production, operational, and hands-on use.
- Domain requirements such as mortgage, financial services, appraisal,
  ServiceNow IRM/GRC, fixed income, etc. can be as important as technologies.

Do not engage in keyword stuffing.

======================================================================
CANDIDATE POSITIONING TITLE — NATURAL MARKET TITLE
======================================================================

Before writing the Summary, choose ONE concise professional title for the
candidate and return it as CANDIDATE_TITLE.

PURPOSE:
This is the natural functional title used to describe the candidate in the
submission narrative. It is NOT automatically Freddie's VNDLY Job Title.

CRITICAL DISTINCTION:
- Freddie's VNDLY Job Title may be a VMS/rate-card/taxonomy label selected for
  budgeting, job-family, or system purposes.
- Treat the VNDLY Job Title and STRUCTURED REQUISITION target_title as CLUES
  about the role, not as wording that must be copied into the candidate summary.
- Do NOT default to awkward system labels such as:
  "Developer-Java/J2EE Specialist",
  "Developer-Python Specialist",
  "Project Manager-Senior",
  or similar VMS/rate-card constructions.
- Preserve an official VNDLY title only when it is already a normal, natural
  market-facing professional title AND accurately describes the candidate.

HOW TO CHOOSE CANDIDATE_TITLE:
Use ALL relevant evidence together:
1. What function the current role actually needs, based on the JD plus current
   manager/MSP/Spotlight guidance.
2. What the candidate has actually done in dated work history.
3. The candidate's recent/repeated professional titles and responsibilities.
4. A more natural role title stated inside the body of the JD, when present.
5. The VNDLY Job Title only as supporting context, never as the default wording.

TITLE QUALITY RULES:
- Prefer a conventional external-market title a recruiter or hiring manager
  would naturally say aloud.
- Usually 2-5 words.
- Keep it functional and specific enough to communicate the candidate's lane.
- Use seniority only when BOTH the candidate's demonstrated career level and
  the target role support it.
- Do not inflate to Lead, Principal, Architect, Manager, etc. merely because the
  candidate once held that title.
- Do not downgrade an established senior candidate merely because the VNDLY
  rate-card title omits seniority.
- Avoid awkward VMS punctuation/order such as "Developer-Java/J2EE Specialist."
- Do not force the target role's function onto a candidate whose actual work
  materially differs from it.
- When the JD itself contains a natural role heading that is more descriptive
  than the VNDLY system title, strongly prefer that heading when the candidate
  genuinely supports it.

EXAMPLE:
VNDLY Job Title: "Developer-Java/J2EE Specialist"
JD body role heading: "Senior Java Developer"
Candidate: 18+ years of Java/Spring/J2EE development

Preferred CANDIDATE_TITLE:
"Senior Java Developer"

Not:
"Developer-Java/J2EE Specialist"

SUMMARY CONSISTENCY RULE:
- If Sentence 1 identifies the candidate by a professional title, use the exact
  CANDIDATE_TITLE you selected.
- Do not introduce a competing title elsewhere in the Summary.

======================================================================
SUMMARY — EXACTLY 4 SENTENCES
======================================================================

Write exactly FOUR sentences as ONE paragraph.

The Summary serves TWO human readers in sequence:

1. An MSP recruiter who knows Freddie's requirements but may not be deeply
   technical.
2. The Freddie Mac hiring manager who can evaluate the technical depth.

The first half of the Summary must make the shortlist decision easy.
The second half must provide the concrete proof that preserves hiring-manager
credibility.

----------------------------------------------------------------------
Sentence 1 — MSP MATCH ANCHOR
----------------------------------------------------------------------

- Use the candidate's FIRST NAME.
- Use the exact CANDIDATE_TITLE selected under the Natural Market Title rules
  above when identifying the candidate by title.
- Do NOT copy the VNDLY Job Title merely because it is present in the structured
  requisition.
- Do not relabel a Data Engineer as a Data Scientist solely because Freddie's
  requisition uses that title. Target-role terminology may be used only when
  the candidate's actual work clearly supports that functional identity.
- If PYTHON-CALCULATED TOTAL PROFESSIONAL EXPERIENCE above is populated, use
  that EXACT value for TOTAL CAREER EXPERIENCE only.
- Never independently calculate or alter the total-years figure.
- CRITICAL GRAMMAR RULE: The total-career experience figure must NOT be written
  in a way that implies the candidate has that same number of years with every
  skill named later in the sentence.
- Do NOT write:
  "7+ years of experience in Python, SQL, and Excel"
  unless EACH named skill independently supports 7+ years.
- Prefer:
  "7+ years of professional experience, including 4+ years of Python scripting
  and extensive SQL and Excel data analysis"
  when the individual skills have different tenures.
- Give the MSP reviewer an immediate checklist-level view of the candidate's
  match to Freddie's highest-priority requirements.
- Prioritize approximately 2-4 of the most important supported Must Have
  competencies rather than trying to mention everything.
- Keep this sentence understandable without requiring deep technical knowledge.

----------------------------------------------------------------------
Sentence 2 — SPECIALIZED MATCH IN PLAIN ENGLISH
----------------------------------------------------------------------

- Use the current/most-recent employer when it provides strong relevant proof.
- HOWEVER, if an earlier employer provides materially stronger evidence for
  Freddie's central need, use that stronger experience instead.
- Explain the candidate's most unusually relevant experience in clear
  functional language.
- Focus first on WHAT the candidate actually did and WHY it is relevant to
  Freddie's need.
- Do not require the MSP reviewer to interpret specialized technical jargon to
  understand the match.
- If a specialized technical concept is important, translate its practical
  meaning.

Example principle:

Instead of relying only on:
"parsed bounding boxes, labels, confidence scores, and metadata"

prefer a construction such as:
"worked directly with appraisal PDFs and property images, validating
information extracted by computer vision models against the source documents"

when that is supported by the resume.

The technical specifics can then appear in Sentence 3.

----------------------------------------------------------------------
Sentence 3 — TECHNICAL / OPERATIONAL PROOF
----------------------------------------------------------------------

- Now provide the concrete technical proof behind the plain-English match.
- Explain HOW the relevant work was performed.
- Naturally incorporate Freddie's most important exact technologies,
  methodologies, operational concepts, or domain terminology when supported.
- Favor demonstrated hands-on use over standalone skills-list mentions.
- Emphasize production ownership, validation, testing, troubleshooting,
  reconciliation, scale, measurable impact, or operational responsibility when
  relevant.
- Use technical terminology where it adds evidence rather than complexity.
- Do NOT create a dense keyword list.
- Normally use no more than 4 material named technologies/tools in this
  sentence. A fifth is allowed only when necessary to describe one coherent
  workflow accurately.

----------------------------------------------------------------------
Sentence 4 — DOMAIN + IMMEDIATE VALUE
----------------------------------------------------------------------

- Tie together the strongest remaining evidence that increases confidence in
  the shortlist decision.
- When Freddie requires or strongly values domain experience such as mortgage,
  financial services, appraisal, capital markets, fixed income, ServiceNow
  IRM/GRC, etc., make that domain connection explicit when genuinely supported.
- State the specific value the candidate's demonstrated experience indicates
  they can provide to THIS Freddie role.
- The conclusion should make the candidate's relevance obvious rather than
  merely praise the candidate.
- Keep it evidence based and grounded in prior execution.
- Do not say "meets the requirements", "great fit", "ideal candidate",
  "strong candidate", or similar evaluative conclusions.
- Do not mention city, onsite requirement, interview availability, rate,
  authorization, or sponsorship.

======================================================================
MSP TRANSLATION RULE
======================================================================

For specialized technical experience:

PLAIN-LANGUAGE RELEVANCE FIRST -> TECHNICAL PROOF SECOND.

The Summary should answer both:

"What does this experience mean for the work Freddie needs done?"

and

"What specific evidence proves the candidate can do it?"

Do not assume that an MSP reviewer understands why a technology is relevant
merely because the technology name appears.

Do not remove important technical terminology simply to make the Summary
easier to read.

Translate the relevance; preserve the evidence.

======================================================================
VETTING-ANSWER INTERACTION
======================================================================

If official Freddie Supplier Vetting Questions exist:

- Assume the MSP reviewer may read those responses before or alongside the
  Summary.
- Do NOT waste Summary space repeating detailed vetting answers.
- Use the Summary to make the CONCLUSION of the candidate's match immediately
  visible.
- Let the vetting answers provide additional depth and confirmation.
- Candidate-supplied vetting responses remain secondary evidence and may not
  override contradictory resume evidence.

======================================================================
STYLE
======================================================================

- Human recruiter voice.
- Direct, factual, concise.
- Evidence over praise.
- Clear enough for a knowledgeable staffing professional who is not a subject
  matter expert in the underlying technology.
- Technically credible enough for the Freddie hiring manager.
- Do NOT make the Summary sound simplistic or generic merely to accommodate
  the MSP reader.
- Do NOT use promotional or generic recruiter language such as:
  "highly accomplished"
  "highly experienced"
  "highly skilled"
  "strong background"
  "excellent fit"
  "ideal candidate"
  "perfect candidate"
  "strong candidate"
  "directly aligns"
  "directly aligning"
  "aligns with"
  "aligns perfectly"
  "positions them uniquely"
  "exceptionally well prepared"
  "proven history"
- Do not say that the candidate "meets Freddie Mac's requirements" or
  "directly aligns with Freddie Mac's core requirements." Demonstrate the
  alignment through specific evidence instead.
- Do not use "Additionally", "Furthermore", or "Moreover".
- Do not simply restate the JD.
- When stating total professional experience, use the conservative format
  "X+ years" based on completed years as of the CURRENT DATE above.
- TOTAL EXPERIENCE VS. SKILL TENURE:
  A total-career figure such as "7+ years" describes the candidate's overall
  professional experience only. Never allow sentence grammar to convert that
  number into claimed tenure for named skills whose evidence supports fewer
  years.

FINAL SUMMARY TEST:

Before returning the Summary, ask:

1. Could an MSP recruiter who understands the requisition but is not deeply
   technical explain in plain English WHY this candidate should be shortlisted?

2. Can that reviewer quickly identify Freddie's most important supported
   Must Haves without searching through technical prose?

3. Does the Summary provide concrete evidence behind those claims rather than
   merely saying the candidate is qualified?

4. Would the Freddie hiring manager still see meaningful technical,
   operational, and domain-specific depth?

5. Is every candidate claim defensible against the resume?

If any answer is no, revise the Summary before returning it.

======================================================================
HIGH-PRIORITY REQUIREMENT EVIDENCE STANDARD
======================================================================

The importance of a Freddie requirement does NOT lower the evidence threshold.
It raises it.

When a requirement is a Must Have, manager priority, or specifically emphasized
in MSP/VNDLY notes:

- Require direct resume evidence before presenting the candidate as possessing
  that competency.
- Do NOT upgrade related, adjacent, administrative, coordination, support,
  oversight, or exposure-level experience into hands-on ownership merely
  because the requirement is important to Freddie.
- Do NOT merge two separate resume facts into a stronger composite claim unless
  the resume explicitly connects them.

Examples:

If the resume says:
"Ensured published documents were stored in SharePoint libraries"

you may accurately describe:
- SharePoint exposure
- use of SharePoint libraries
- supporting document availability in SharePoint

but you may NOT automatically claim:
- SharePoint content management
- building SharePoint sites/pages
- publishing content in SharePoint
- administering SharePoint
- hands-on SharePoint development

unless the resume provides that evidence.

Likewise, if one bullet says the candidate created job aids and another says
documents were stored in SharePoint, do NOT combine those facts into:
"created and maintained job aids in SharePoint"
unless the resume explicitly connects the job aids to SharePoint.

For high-priority requirements, distinguish carefully between:

1. DIRECT HANDS-ON EVIDENCE
   The resume explicitly shows the candidate personally performed the required
   work.

2. RELATED / PARTIAL EVIDENCE
   The resume shows adjacent, supporting, coordinating, oversight, or
   exposure-level experience but does not prove the full requirement.

3. NO EVIDENCE
   The resume does not support the requirement.

Only DIRECT HANDS-ON EVIDENCE may be stated as full possession of the
requirement in SUMMARY or SKILLS.

RELATED / PARTIAL EVIDENCE may be described accurately at its demonstrated
level, but must never be upgraded into the full Freddie requirement.

When evidence is partial, it is better to omit the requirement from the Skills
table than to overstate the candidate.

======================================================================
SKILLS TABLE — EXACTLY 4 HIGH-SIGNAL ROWS WHEN SUPPORTED
======================================================================

The template has four Skills rows.

The Skills table serves as the MSP reviewer's FASTEST VISUAL CHECKLIST of the
candidate's match to Freddie's requisition.

PRIMARY STRATEGY:
- Start from Freddie's highest-priority supported Must Haves.
- Use the four rows to make the strongest genuine Freddie requirements
  immediately visible.
- Prefer Freddie's exact terminology when supported by the candidate.
- Write labels so a nontechnical MSP recruiter can compare them directly
  against Freddie's Must Have / Required Qualifications language.
- Keep labels short, specific, and easy to scan.
- Do not make the reviewer decode a vague umbrella category to determine
  whether a Must Have is present.
- Do not bury a critical Freddie Must Have among several less-important tools
  inside parentheses.
- A Skills row should communicate ONE clear shortlist signal.
- Technical detail in parentheses is useful only when it strengthens or proves
  that signal.
- The Skills table is not intended to summarize the candidate's entire career.
  It is a four-row shortlist checklist.

MSP-SCAN PRIORITY:

When multiple truthful Skills-row options exist, prefer the row that lets an
MSP reviewer most quickly answer:
"Does this candidate have the thing Freddie specifically asked for?"

For example, prefer a clear label such as:
"Computer Vision & Image Analytics (OpenCV, TensorFlow)"

over a broad label such as:
"Advanced Data Technologies (Python, OpenCV, TensorFlow, JSON)"

when Computer Vision / Image Analytics is what Freddie specifically requires.

Likewise, prefer:
"Mortgage / Financial Services Data"

over:
"Domain Expertise"

when Freddie specifically requires mortgage or financial-services experience.

IMPORTANT:
- This does NOT mean blindly copying Freddie's requirements.
- Every label must still be supported by candidate evidence.
- Do not manufacture a missing Must Have merely to create a visually perfect
  checklist.

CANDIDATE-EVIDENCE RULE:
- A row may contain ONLY skills, tools, methodologies, or domain expertise
  supported by the candidate's resume.
- Official candidate-supplied Freddie vetting answers may clarify depth of
  experience already grounded in the resume, but must not manufacture a skill.
- If a Must Have group is not supported, DO NOT fabricate it merely to fill a
  row.
- A related or adjacent activity does NOT count as full support for a more
  specific Freddie requirement.
- If Freddie requires hands-on execution, the resume must explicitly support
  hands-on execution before that requirement can appear as a Skills-row label.
- Do not infer full competency from a single incidental mention of a tool,
  platform, methodology, or domain.
- Move to the next relevant Required or Preferred competency the candidate
  genuinely possesses.
- If fewer than four defensible relevant groups exist, leave the unused
  SKILL/YEARS fields blank. The Word processor will remove those rows.

YEARS / LAST USED — STRICT EVIDENCE-BASED CALCULATION:

For every Skills-table row, you must ALSO return SKILL#_ROLE_INDEXES containing
the 1-based indexes from STRUCTURED ORIGINAL EXPERIENCE that support the ENTIRE
competency label as written.

ROLE-INDEX RULES:
- Return the dated roles that substantively support the CORE competency area.
- Parenthetical technologies are representative examples/evidence; they are NOT
  an AND checklist that must all appear in the same individual job.
- A Professional Summary or top-level skills inventory is NOT sufficient for a
  role index.
- Do not include a role merely because it is vaguely related.
- Example: for "Relational Databases & SQL (Oracle, DB2, SQL Server)", count
  dated roles doing meaningful relational-database/SQL work even if one role used
  Oracle and another used DB2. Do NOT require Oracle + DB2 + SQL Server in every
  single role.
- Example: for "Financial Services Background (Banking, Investment, Securities)",
  include dated financial-services roles across those subdomains; do not require
  all three words in every role.
- Python will merge overlapping role periods and round down.
- Avoid overstuffed rows that combine unrelated competency families. Keep each
  row coherent enough that one aggregate years figure is meaningful.

For every Skills-table row, calculate the displayed experience conservatively
from the dated work history.

CORE RULE:
- The YEARS value must be defensible for the ENTIRE competency label as written.
- Never use the tenure of one broad component to imply equal tenure with every
  named technology or specialty in the same row.

EXAMPLE OF WHAT NOT TO DO:
If the candidate has 7+ years of Python/Data Engineering but only began using
OpenCV and TensorFlow in 2026, DO NOT write:
"Python, Computer Vision & Image Analytics (OpenCV, TensorFlow) —
7+ years, current"

because that falsely implies 7+ years of Computer Vision/OpenCV/TensorFlow.

GROUPED-SKILL RULE:
- When a row contains multiple material named technologies, specialties, or
  domains, the displayed YEARS value must be literally defensible for the
  entire label.

- Do NOT solve a tenure mismatch merely by assigning the shortest tenure to a
  combined row if doing so would materially UNDERSTATE an important long-tenure
  Freddie Must Have.

- In particular, a high-priority Must Have such as Python, SQL, Java,
  ServiceNow, Snowflake, etc. should NOT be grouped with a much newer specialty
  when the combined label would cause the long-tenure Must Have itself to appear
  artificially inexperienced.

DOMAIN-GROUPING RULE:
- Apply the SAME tenure discipline to business/domain competencies as to
  technologies.
- Do not combine a narrower domain with a broader domain and assign the broader
  domain's longer tenure to the entire row.

EXAMPLES OF NARROWER VS. BROADER DOMAINS:
- Fixed Income vs. Financial Services
- Mortgage vs. Financial Services
- Appraisal vs. Mortgage
- IRM/GRC vs. ServiceNow
- Capital Markets vs. Banking
- Regulatory Reporting vs. General Data Analytics

If the candidate has:
- 4+ years of Regulatory Analytics
- but only 2+ years of Fixed Income

DO NOT write:
"Fixed Income & Regulatory Analytics" — "4+ years, current"

Instead either write:
"Financial & Regulatory Analytics" — "4+ years, current"

if that broader label is fully supported,

OR:
"Fixed Income & Capital Markets" — "2+ years, 2023"

if the narrower Freddie-relevant domain is the more important signal.

The label and YEARS value must describe the SAME defensible scope of
experience.

MANDATORY SPLIT RULE:
- If one material component has substantially longer tenure than another
  material component in the proposed row, and either component is important to
  Freddie's requirements, SPLIT them into separate Skills rows whenever there
  are available rows.

- Treat a difference of approximately 2 or more completed years as materially
  different unless the resume evidence clearly supports treating the items as
  one inseparable competency.

EXAMPLE:
Candidate has:
- Python: 7+ years, current
- Computer Vision/OpenCV/TensorFlow: <1 year, current

DO NOT write:
"Python & Computer Vision (OpenCV, TensorFlow)" — "<1 year, current"

because that incorrectly makes the candidate appear to have less than one year
of Python.

Instead prefer:
"Python & Data Engineering" — "7+ years, current"
"Computer Vision & Image Analytics (OpenCV, TensorFlow)" —
"<1 year, current"

Likewise, if:
- SQL = 7+ years
- Snowflake = <1 year

do not combine them into:
"SQL & Snowflake" — "<1 year"

when separate rows can represent the candidate more accurately.

ROW-ALLOCATION PRIORITY:
When deciding how to use the four available Skills rows, prioritize:

1. Freddie's explicit Must Have competencies that the candidate genuinely has.
2. High-signal specialized competencies central to the actual role.
3. Required domain expertise.
4. Relevant secondary / preferred competencies.

The objective is NOT to force Freddie's four planned competency groups directly
into four resume rows. The planned competency groups are guidance. The final
candidate Skills rows must be reorganized when necessary to represent the
candidate's actual experience accurately and advantageously.

Prefer:
- a truthful long-tenure Must Have row,
- plus a truthful short-tenure specialty row,

over one combined row whose YEARS value materially understates or overstates
either skill.

EXAMPLE:
If Python is supported for 7+ years but Computer Vision/OpenCV/TensorFlow only
since February 2026, prefer:
"Python & Data Engineering" — "7+ years, current"
and
"Computer Vision & Image Analytics (OpenCV, TensorFlow)" —
"<1 year, current"

rather than combining them under "7+ years".

CALCULATION RULES:
- Determine experience only from actual dated roles where that specific skill,
  technology, methodology, or domain competency is evidenced.
- Do NOT assign total career length to a skill merely because the candidate
  worked in a related occupation.
- Do NOT automatically subtract 1-2 years as a heuristic.
- Avoid double-counting overlapping employment periods.
- Calculate through the CURRENT DATE supplied above for roles marked Present
  or Current.
- Round DOWN to completed years rather than up.
- If clearly less than one completed year, use:
  "<1 year, current"
  or
  "<1 year, YYYY"
- Otherwise use:
  "X+ years, current"
  or
  "X+ years, YYYY"

YEAR LAST USED:
- "current" is allowed ONLY when that exact competency, or every material
  component represented by the row, is evidenced in the candidate's actual
  current role.
- Otherwise give the latest actual year in which the competency is evidenced.

FINAL AUDIT BEFORE RETURNING EACH ROW:

Ask BOTH questions:

1. "If a Freddie Mac MSP reviewer reads this Skill label and its Years value
   literally, would the candidate's dated resume support that interpretation?"

2. "Does this grouping accidentally make an important long-tenure Freddie
   requirement appear to have LESS experience merely because it was grouped
   with a newer specialty?"

3. "If this row combines two business/domain concepts, does the YEARS value
   genuinely apply to BOTH concepts as written?"

If not, broaden the label to a truthful shared category, narrow the YEARS value,
or split the concepts.

If either answer reveals a misleading result:
- split the competency,
- narrow the label,
- reduce the years when necessary,
- or leave the row blank.

Never solve an overstatement problem by creating a new material
understatement.

======================================================================
FINAL MSP / HIRING-MANAGER SKILLS CHECK
======================================================================

After the accuracy audit above, review the four Skills rows as a SET.

Ask:

1. If an MSP recruiter looked ONLY at these four rows and Freddie's Must Have
   requirements, would the candidate's strongest genuine matches be obvious?

2. Are Freddie's most important supported requirements named directly enough
   that the reviewer does not need to interpret a vague category?

3. Did any broad or impressive-sounding secondary competency consume a row
   that would be better used for a more important Freddie Must Have?

4. Does each row still contain enough specificity to be meaningful to the
   hiring manager?

5. Are the four rows complementary rather than redundant?

If a row is accurate but unnecessarily vague, rename it more directly using
Freddie's terminology when supported.

If a row is accurate but substantially less important than another supported
Freddie Must Have, replace it with the higher-priority supported competency.

Never sacrifice truthfulness or the strict YEARS rules merely to improve
scanability.

======================================================================
OFFICIAL FREDDIE VETTING ANSWERS
======================================================================

Freddie required that these responses come directly from the candidate.

They are SECONDARY supporting evidence.

Allowed:
- Clarify technical depth.
- Clarify how something already reflected in the work history was used.
- Clarify scope, approach, or operational context.

Not allowed:
- Use an answer to invent an employer/date.
- Calculate years solely from a candidate's unsupported statement.
- Convert a theoretical answer into performed experience.
- Override contradictory resume evidence.

OFFICIAL CANDIDATE-SUPPLIED VETTING Q&A:
{vetting_for_ai}

======================================================================
RESUME = PRIMARY SOURCE OF CANDIDATE TRUTH
======================================================================

ZERO-TOLERANCE RULES:

1. Never assign a JD technology to the candidate merely because Freddie asks
   for it.

2. Never infer a specific tool from a broader category.
   Example:
   Resume says AWS -> you may NOT invent Redshift.

3. Never upgrade exposure into ownership.

4. Never fabricate a metric, scale, accomplishment, mortgage domain,
   certification, platform, methodology, or responsibility.

5. Technologies that appear only in the JD/Spotlight/MSP notes describe the
   ROLE, not the candidate.

6. Every candidate claim in SUMMARY and SKILLS must survive a side-by-side
   audit against the resume, with only limited clarification from official
   candidate vetting responses.

======================================================================
EXPERIENCE PRIORITY / OPERATIONAL OWNERSHIP
======================================================================

Prefer:
- Things the candidate actually built, supported, maintained, tested,
  troubleshot, validated, deployed, operated, analyzed, or owned.
- Recurring direct responsibility.
- Recent relevant work.
- Production support and issue resolution when relevant.
- Operational ownership.
- Compliance/audit/validation responsibility when relevant.
- Evidence covering multiple Freddie requirements at once.

Avoid overweighting:
- One-time transformation projects when the target role is operational.
- Technologies listed only in a skills inventory.
- Generic management language.
- Impressive but irrelevant projects.

======================================================================
SAFE WORK-HISTORY BULLET PRIORITIZATION
======================================================================

You are NOT allowed to rewrite the work-history bullets.

Instead, for each role, rank the EXISTING ORIGINAL BULLETS by relevance to
this Freddie requisition.

Return zero-based bullet indexes.

Example:
If a role has 4 bullets and bullets 2, 0, 3, 1 are the best order:
[2, 0, 3, 1]

Rules:
- Do NOT omit bullets.
- Do NOT create bullets.
- Do NOT merge bullets.
- Do NOT rewrite bullets.
- Rank only the existing indexes.
- If the original order is already best, return the original indexes.
- If a role has no bullets, return [].

STRUCTURED ORIGINAL EXPERIENCE:
{structured_exp}

======================================================================
OUTPUT FORMAT
======================================================================

Return ONLY:

{{
  "CANDIDATE_TITLE": "",
  "SUMMARY": "",
  "SKILL1": "",
  "YEARS1": "",
  "SKILL1_ROLE_INDEXES": [],
  "SKILL2": "",
  "YEARS2": "",
  "SKILL2_ROLE_INDEXES": [],
  "SKILL3": "",
  "YEARS3": "",
  "SKILL3_ROLE_INDEXES": [],
  "SKILL4": "",
  "YEARS4": "",
  "SKILL4_ROLE_INDEXES": [],
  "BULLET_ORDER": {{
    "1": [],
    "2": [],
    "3": [],
    "4": [],
    "5": [],
    "6": [],
    "7": []
  }}
}}

======================================================================
FULL ORIGINAL RESUME
======================================================================

{raw_text}
"""

                summary_data = fred_generate_json(
                    API_KEY,
                    tailoring_prompt,
                )

                # ====================================================
                # VALIDATE CANDIDATE POSITIONING TITLE / SUMMARY / SKILLS
                # ====================================================

                candidate_positioning_title = str(
                    summary_data.get(
                        "CANDIDATE_TITLE",
                        "",
                    )
                ).strip()

                final_summary = str(
                    summary_data.get(
                        "SUMMARY",
                        "",
                    )
                ).strip()

                final_summary = fred_clean_summary(
                    final_summary,
                    name,
                )

                if not final_summary:
                    raise RuntimeError(
                        "The Freddie optimization step returned no "
                        "candidate Summary."
                    )

                # ====================================================
                # APPLY SAFE BULLET REORDERING
                # ====================================================

                bullet_orders = summary_data.get(
                    "BULLET_ORDER",
                    {},
                )

                if not isinstance(
                    bullet_orders,
                    dict,
                ):
                    bullet_orders = {}

                for idx, role in enumerate(
                    experience,
                    start=1,
                ):
                    original_bullets = role.get(
                        "Bullets",
                        [],
                    )

                    requested_order = bullet_orders.get(
                        str(idx),
                        [],
                    )

                    role["Bullets"] = (
                        fred_safe_reorder_bullets(
                            original_bullets,
                            requested_order,
                        )
                    )

                # ====================================================
                # FREDDIE-ONLY DETERMINISTIC SKILL YEARS
                # ====================================================
                # Gemini identifies only the dated roles that support the ENTIRE
                # competency label. Python performs the actual month math,
                # merges overlaps, rounds down, and determines year last used.
                for skill_num in range(1, 5):
                    years_key = f"YEARS{skill_num}"
                    role_key = f"SKILL{skill_num}_ROLE_INDEXES"

                    proposed_role_indexes = fred_clean_skill_role_indexes(
                        summary_data.get(role_key, []),
                        len(experience),
                    )

                    role_indexes = fred_filter_skill_role_indexes_by_label(
                        summary_data.get(f"SKILL{skill_num}", ""),
                        experience,
                        proposed_role_indexes,
                    )

                    deterministic_skill_years = (
                        fred_calculate_skill_years_from_roles(
                            experience,
                            role_indexes,
                            current_date_obj,
                        )
                    )

                    if deterministic_skill_years:
                        summary_data[years_key] = deterministic_skill_years
                    else:
                        # Conservative fallback: never leave a populated skill row
                        # with a blank years cell. If no role dates can be parsed,
                        # retain Gemini's estimate only after enforcing the total
                        # career ceiling.
                        summary_data[years_key] = fred_cap_skill_years(
                            summary_data.get(years_key, ""),
                            calculated_total_experience,
                        )

                # ====================================================
                # BUILD VNDLY SUBMISSION SUMMARY + RECOMMENDED SKILLS
                # ====================================================

                vndly_package = {
                    "candidate_positioning_title": candidate_positioning_title,
                    "vndly_summary": "",
                    "final_recommended_vndly_skills": [],
                    "exact_must_have_skills": [],
                    "exact_nice_to_have_skills": [],
                    "required_vndly_skills_from_job_intelligence": [],
                    "preferred_vndly_skills_from_job_intelligence": [],
                    "additional_vndly_skills": [],
                }
                vndly_package_error = None

                try:
                    with st.spinner(
                        "Building VNDLY submission summary and exact skill recommendations..."
                    ):
                        vndly_package = fred_build_vndly_candidate_package(
                            API_KEY,
                            name,
                            raw_text,
                            structured_req,
                            selected_vndly_context,
                            vetting_for_ai,
                            final_summary,
                            calculated_total_experience,
                            candidate_positioning_title,
                        )
                except Exception as package_error:
                    vndly_package_error = str(package_error)

                # ====================================================
                # BUILD WORD MAPPING
                # ====================================================

                mapping = {
                    "FullName": name,
                    "Location": Current_Location,
                    "Remote": Available_Onsite,
                    "FormerFM": Former_CW,
                    "FormerCW": Former_CW,
                    "Eligibility": Eligibility,
                    "Authorization": Auth_US,
                    "Sponsorship": Sponsorship,
                    "Site": Site,
                    "Interview": Interview,
                    "Certifications": certifications,
                    "SUMMARY": final_summary,
                    "SKILL1": str(
                        summary_data.get(
                            "SKILL1",
                            "",
                        )
                    ).strip(),
                    "YEARS1": str(
                        summary_data.get(
                            "YEARS1",
                            "",
                        )
                    ).strip(),
                    "SKILL2": str(
                        summary_data.get(
                            "SKILL2",
                            "",
                        )
                    ).strip(),
                    "YEARS2": str(
                        summary_data.get(
                            "YEARS2",
                            "",
                        )
                    ).strip(),
                    "SKILL3": str(
                        summary_data.get(
                            "SKILL3",
                            "",
                        )
                    ).strip(),
                    "YEARS3": str(
                        summary_data.get(
                            "YEARS3",
                            "",
                        )
                    ).strip(),
                    "SKILL4": str(
                        summary_data.get(
                            "SKILL4",
                            "",
                        )
                    ).strip(),
                    "YEARS4": str(
                        summary_data.get(
                            "YEARS4",
                            "",
                        )
                    ).strip(),
                }

                # Education
                for i in range(1, 4):
                    if i <= len(education):
                        edu_item = education[i - 1]

                        if not isinstance(
                            edu_item,
                            dict,
                        ):
                            edu_item = {}

                        mapping[f"School{i}"] = clean_school(
                            str(
                                edu_item.get(
                                    "School",
                                    "",
                                )
                            )
                        )

                        mapping[f"Degree{i}"] = str(
                            edu_item.get(
                                "Degree",
                                "",
                            )
                        ).strip()

                    else:
                        mapping[f"School{i}"] = ""
                        mapping[f"Degree{i}"] = ""

                # Work History
                for i in range(1, max(7, len(experience)) + 1):
                    if i <= len(experience):
                        role = experience[i - 1]

                        mapping[f"Company{i}"] = (
                            role.get(
                                "Company",
                                "",
                            )
                        )

                        mapping[f"Title{i}"] = (
                            role.get(
                                "Title",
                                "",
                            )
                        )

                        mapping[f"Bullets{i}"] = clean_bullets(
                            role.get(
                                "Bullets",
                                [],
                            )
                        )

                        mapping[f"Environment{i}"] = (
                            role.get(
                                "Environment",
                                "",
                            )
                        )

                        mapping[f"Dates{i}"] = (
                            role.get(
                                "Dates",
                                "",
                            )
                        )

                    else:
                        mapping[f"Company{i}"] = ""
                        mapping[f"Title{i}"] = ""
                        mapping[f"Bullets{i}"] = ""
                        mapping[f"Environment{i}"] = ""
                        mapping[f"Dates{i}"] = ""

                # ====================================================
                # GENERATE WORD DOCUMENT
                # ====================================================

                safe_name = (
                    re.sub(
                        r"[^A-Za-z0-9_\-]+",
                        "_",
                        name,
                    ).strip("_")
                    or "Candidate"
                )

                out_file = (
                    f"Submission_FreddieMac_{safe_name}.docx"
                )

                process_freddie_word_doc(
                    TEMPLATE_FILENAME,
                    mapping,
                    vetting_pairs,
                    out_file,
                )

                # ====================================================
                # VNDLY SUBMISSION PACKAGE — READY TO COPY
                # ====================================================

                st.subheader("📨 VNDLY Submission Details")

                if vndly_package_error:
                    st.warning(
                        "⚠️ The resume was generated successfully, but the "
                        "VNDLY summary/skill package could not be created. "
                        f"Package error: {vndly_package_error}"
                    )

                st.markdown("**VNDLY Submission Summary**")
                if vndly_package.get("vndly_summary", ""):
                    st.code(
                        vndly_package.get("vndly_summary", ""),
                        language=None,
                    )
                else:
                    st.caption("No VNDLY submission summary was generated.")

                ranked_vndly_skills = vndly_package.get(
                    "final_recommended_vndly_skills",
                    [],
                ) or []

                st.markdown(
                    f"**Recommended VNDLY Skills — Top {FREDDIE_VNDLY_SKILL_CAP} Maximum**"
                )
                st.caption(
                    "Ranked from highest to lowest submission value. "
                    "Select these exact values in VNDLY."
                )

                source_labels = {
                    "exact_structured_must": "Exact Structured Must",
                    "exact_structured_nice": "Exact Structured Nice",
                    "required_job_intelligence": "Required",
                    "preferred_job_intelligence": "Preferred",
                    "additional_high_value": "Additional",
                }

                if ranked_vndly_skills:
                    for index, item in enumerate(
                        ranked_vndly_skills,
                        start=1,
                    ):
                        skill_name = str(
                            item.get("skill_name", "") or ""
                        ).strip()
                        category = str(
                            item.get("source_category", "") or ""
                        ).strip().lower()
                        label = source_labels.get(
                            category,
                            "Recommended",
                        )
                        st.markdown(
                            f"{index}. `{skill_name}` — {label}"
                        )
                else:
                    st.caption(
                        "No VNDLY catalogue skills were supported strongly enough "
                        "to recommend."
                    )

                # ====================================================
                # AUTOMATIC FREDDIE SUBMISSION EMAIL
                # ====================================================

                if not vndly_package_error and vndly_package.get(
                    "vndly_summary",
                    "",
                ):
                    email_key = fred_generation_email_key(
                        name,
                        selected_vndly_context.get("job_id", ""),
                        raw_text,
                        vndly_package.get("vndly_summary", ""),
                    )

                    sent_keys = st.session_state.setdefault(
                        "fred_sent_email_keys",
                        set(),
                    )

                    if email_key not in sent_keys:
                        try:
                            fred_send_submission_email(
                                name,
                                selected_vndly_context,
                                vndly_package,
                            )
                            sent_keys.add(email_key)
                            st.success(
                                "📧 VNDLY submission details were emailed automatically."
                            )
                        except Exception as email_error:
                            st.warning(
                                "⚠️ The resume was generated successfully, but the "
                                "automatic Freddie submission email could not be sent. "
                                f"Email error: {email_error}"
                            )
                    else:
                        st.info(
                            "📧 This exact Freddie candidate/job submission was already "
                            "emailed during the current app session, so a duplicate "
                            "email was not sent."
                        )
                else:
                    st.warning(
                        "⚠️ Automatic email was skipped because the VNDLY "
                        "submission package was not available."
                    )

                with open(
                    out_file,
                    "rb",
                ) as file:
                    st.download_button(
                        label="⬇️ Download Generated Document",
                        data=file,
                        file_name=out_file,
                        mime=(
                            "application/vnd.openxmlformats-"
                            "officedocument.wordprocessingml.document"
                        ),
                        type="primary",
                    )

                st.success(
                    "✅ Freddie Mac submission is ready."
                )

            except Exception as e:
                st.error(
                    f"❌ Freddie Mac processing failed: {str(e)}"
                )

            finally:
                if (
                    resume_path
                    and os.path.exists(resume_path)
                ):
                    try:
                        os.remove(resume_path)
                    except Exception:
                        pass


# ====================================================================
# --- 🔵 CLIENT APP 2: PERATON 🔵 ---
# ====================================================================
def peraton_app():
    TEMPLATE_FILENAME = "Peraton_Template.docx"
    st.title("Peraton Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("🎤 Supplier Technical Interview Results")
    
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="per_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="per_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="per_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="per_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="per_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="per_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="per_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="per_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="per_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="per_a5")

    st.header("📝 Job Description & Notes")
    Job_Description_Notes_etc = st.text_area("Paste JD, Manager Notes, etc. here:", height=200, key="per_jd")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="per_res")

    if st.button("🚀 Generate Peraton Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line.
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Responsible, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title.
                        - For 'Responsible', write exactly 1 sentence summarizing what the candidate was responsible for at this specific job, based on their bullets. Do NOT use the candidate's name or pronouns (he/she). Start the sentence with the exact words "Responsible for ".
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role.
                        - For 'Dates', preserve the date precision actually shown on the original resume. If month + year are shown, format as "MM/YYYY to MM/YYYY" (or "MM/YYYY to Present"). If ONLY years are shown, preserve them as "YYYY to YYYY" (or "YYYY to Present"). NEVER invent January, "01", or any other month when the source provides only a year. Mixed-precision dates must remain mixed rather than filling in a missing month.
                    5. Certifications: Extract active certifications into an ARRAY of strings. Do not include classes or courses taken. Keep it strictly to the certification name.

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": [],
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Responsible": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title().strip()
                    name_parts = name.split(' ', 1)
                    first_name = name_parts[0] if len(name_parts) > 0 else ""
                    last_name = name_parts[1] if len(name_parts) > 1 else ""

                    mapping = {
                        "Firstname": first_name, "Lastname": last_name, 
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY": "", 
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", "SUMMARY5": "",
                        "CERTIFICATION1": "", "CERTIFICATION2": "", "CERTIFICATION3": "", "CERTIFICATION4": "",
                        "SKILL1": "", "SKILL2": "", "SKILL3": "", "SKILL4": ""
                    }

                    # Handle Certifications Array
                    certs = data.get("Certifications", [])
                    if isinstance(certs, str): 
                        certs = [c.strip() for c in certs.split(',') if c.strip()]
                    for i in range(1, 5):
                        mapping[f"CERTIFICATION{i}"] = certs[i-1] if i <= len(certs) else ""

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"RESPONSIBLE{i}"] = exp[i-1].get('Responsible', '')
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            
                            # USE NEW PERATON DATE FORMATTER HERE
                            mapping[f"Dates{i}"] = format_peraton_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"RESPONSIBLE{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description_Notes_etc.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Peraton Fieldglass portal. 
                            Your goal is to make an evidence-based business case for why this candidate will succeed, strictly following Peraton's specific submission guidelines.

                            ========================
                            OVERALL SUMMARY RULES (STRICT)
                            ========================
                            You must output the "SUMMARY" field as a single paragraph of 150 words or less.
                            - FIRST PERSON: You MUST write this summary entirely in the 1st person perspective (e.g., "I am a Data Analyst with...", "My background includes..."). Do NOT use the candidate's name or 3rd person pronouns (he/she/they).
                            - It must highlight exactly why they are the best candidate for this specific position.
                            - You MUST explicitly state their total number of years of professional experience.
                            - Be sure the paragraph seamlessly highlights the specific skills required for the position.

                            ========================
                            SUMMARY OF QUALIFICATIONS RULES (STRICT)
                            ========================
                            You must output EXACTLY 5 distinct strings mapped to "SUMMARY1" through "SUMMARY5".
                            - TAILOR TO THE JD: Only include qualifications specifically tailored to the position being applied for. Omit irrelevant history.
                            - ACTION VERBS: Start EVERY point with a strong action verb (e.g., Engineered, Managed, Architected).
                            - QUANTIFY EVERYTHING: You MUST substantiate/quantify each point with a number derived from the resume/Q&A (e.g., number of systems, applications, users, bandwidth, endpoints, data scale, or team size). 
                            - PRIORITIZATION: Place the most relevant points first.
                            - You MUST provide all 5 bullets. Do not leave any blank. 
                            - DO NOT include the physical bullet character (•) in the text. The Word document template will apply the bullet automatically.

                            ========================
                            TECHNICAL SKILLS RULES (STRICT)
                            ========================
                            - ONLY include technical skills and tools directly relevant to the position being applied for.
                            - DO NOT include basic/universal skills (e.g., MS Word, Excel, PowerPoint, Outlook).
                            - Divide the relevant skills into exactly 4 logical buckets mapped to SKILL1 through SKILL4. 

                            Format:
                            "Skill Area (Tool1, Tool2, Tool3, Tool4)"

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth.
                            - The technical interview Q&A is SECONDARY and should only be used to clarify depth or extract quantifiable metrics (numbers/scale) to support the Summary bullets.
                            - If there is any conflict: 👉 prioritize the resume over Q&A.

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME.
                            3. FACT AUDIT: Before outputting the final JSON, verify every tool and NUMBER you used. If a metric/number doesn't exist in the resume/Q&A, do not invent one. 

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY": "150 word paragraph highlighting fit and total years of experience...",
                              "SUMMARY1": "Engineered [X] resulting in [Y]...",
                              "SUMMARY2": "Architected [X] for [Y] users...",
                              "SUMMARY3": "Managed [X] endpoints...",
                              "SUMMARY4": "Developed [X] applications...",
                              "SUMMARY5": "Optimized [X] systems...",
                              "SKILL1": "Skill Area (tools)",
                              "SKILL2": "Skill Area (tools)",
                              "SKILL3": "Skill Area (tools)",
                              "SKILL4": "Skill Area (tools)"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Job Description / Notes:
                            {Job_Description_Notes_etc}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY"] = summary_data.get("SUMMARY", "")
                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SUMMARY5"] = summary_data.get("SUMMARY5", "")
                            
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")

                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_Peraton_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- 🟢 CLIENT APP 3: CAPITAL ONE 🟢 ---
# ====================================================================
def capital_one_app():
    TEMPLATE_FILENAME = "CapitalOne_Template.docx"
    st.title("Capital One Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="co_loc")
        
        # --- NEW REMOTE/HYBRID LOGIC ---
        Remote_or_Onsite = st.selectbox("Remote or Onsite/Hybrid", ["Remote", "Onsite/Hybrid"], index=0, key="co_rem")
        co_location_name = ""
        if Remote_or_Onsite == "Onsite/Hybrid":
            co_location_name = st.text_input("Location Name (e.g., HQ2)", key="co_loc_name")
            
        # --- NEW START DATE LOGIC ---
        co_start = st.text_input("Availability to Start", key="co_start")

    with col2:
        # --- NEW FORMER EMPLOYEE LOGIC ---
        Former_CO = st.selectbox("Former Capital One", ["Y", "N"], index=1, key="co_form")
        co_mgr = ""
        co_dates = ""
        if Former_CO == "Y":
            co_mgr = st.text_input("Manager Name", key="co_mgr")
            co_dates = st.text_input("Dates Worked", key="co_dates")
            
        LinkedIn_GitHub_Portfolio_Link = st.text_input("LinkedIn/GitHub/Portfolio Link", key="co_link")

    # Mapping logic variables based on above inputs
    if Remote_or_Onsite == "Onsite/Hybrid" and co_location_name.strip():
        remote_val = f"Onsite/Hybrid - {co_location_name.strip()}"
    else:
        remote_val = Remote_or_Onsite
        
    if Former_CO == "Y":
        former_co_val = f"Y - {co_mgr.strip()}, {co_dates.strip()}"
    else:
        former_co_val = "N"

    st.header("🎤 Supplier Technical Interview Results")
    
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="co_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="co_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="co_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="co_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="co_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="co_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="co_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="co_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="co_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="co_a5")

    st.header("📝 Job Description & Notes")
    Job_Description_Notes_etc = st.text_area("Paste JD, Manager Notes, etc. here:", height=200, key="co_jd")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="co_res")

    if st.button("🚀 Generate Capital One Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line.
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title (e.g., remove '- Contract', '(Contract)', or '- Consultant').
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title()
                    
                    mapping = {
                        "FullName": name, "Location": Current_Location_City_ST, "Remote": remote_val,
                        "Start": co_start,
                        "Certifications": data.get("Certifications", ""),
                        "FormerCapitalOne": former_co_val, "Links": LinkedIn_GitHub_Portfolio_Link,
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", 
                        "SKILL1": "", "YEARS1": "", "SKILL2": "", "YEARS2": "",
                        "SKILL3": "", "YEARS3": "", "SKILL4": "", "YEARS4": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description_Notes_etc.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The Hiring Manager has 30 seconds to read this. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed.

                            ========================
                            THE NARRATIVE BLUEPRINT (4 Bullet Points)
                            ========================
                            Generate exactly 4 distinct bullet points mapped to SUMMARY1 through SUMMARY4. 
                            CRITICAL RULE: DO NOT use the candidate's name or 3rd person pronouns (he/she/they) in ANY of the bullet points. Start each bullet with a strong action verb or direct statement.
                            
                            - SUMMARY1 (The Anchor): What is their TOTAL progressive professional experience, and what is their dominant expertise that solves the PRIMARY technical "must-have" of the Job Description?
                            - SUMMARY2 (The Alignment): Explicitly name their CURRENT or most recent employer. What is the most impressive, relevant project they recently delivered that proves they can handle this specific job?
                            - SUMMARY3 (The Execution & Impact): How did they build it, and why does it matter? Weave specific tools/methodologies into an "Execution Statement" that highlights complexity, scale, or business impact.
                            - SUMMARY4 (The Closer): Based on their past execution, what specific value will they deliver on Day 1 in THIS new role?

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - LEADERSHIP VERBS: Use high-authority active verbs (e.g., 'Engineered', 'Optimized', 'Architected', 'Spearheaded').
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely."
                            - Do NOT repeat or restate the job description.

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth for experience.
                            - The technical interview Q&A is SECONDARY.
                            If there is any conflict: 👉 prioritize the resume over Q&A.

                            ========================
                            SKILLS SECTION
                            ========================
                            - EXACTLY 4 items
                            - Prioritize the specific "Must-Have" technologies.
                            - Use only tools explicitly mentioned in resume/Q&A/notes
                            Format: "Skill Area (Tool1, Tool2, Tool3, Tool4)"
                            Years format: "X+ years, current" OR "X+ years, 2026"

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME. 
                            3. DO NOT INFLATE TO MATCH THE JD. 
                            4. FACT AUDIT: Before outputting the final JSON, verify every single tool.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY1": "Bullet 1 following the blueprint",
                              "SUMMARY2": "Bullet 2 following the blueprint",
                              "SUMMARY3": "Bullet 3 following the blueprint",
                              "SUMMARY4": "Bullet 4 following the blueprint",
                              "SKILL1": "Skill Area (tools)",
                              "YEARS1": "X+ years, current",
                              "SKILL2": "Skill Area (tools)",
                              "YEARS2": "X+ years, current",
                              "SKILL3": "Skill Area (tools)",
                              "YEARS3": "X+ years, current",
                              "SKILL4": "Skill Area (tools)",
                              "YEARS4": "X+ years, current"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Job Description / Notes:
                            {Job_Description_Notes_etc}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["YEARS1"] = summary_data.get("YEARS1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["YEARS2"] = summary_data.get("YEARS2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["YEARS3"] = summary_data.get("YEARS3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")
                            mapping["YEARS4"] = summary_data.get("YEARS4", "")
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    # 1. Generate the .docx file first
                    docx_file = f"Submission_CapitalOne_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, docx_file)
                    
                    # 2. Convert the .docx to .pdf using LibreOffice (Linux compatible)
                    pdf_file = f"Submission_CapitalOne_{name.replace(' ', '_')}.pdf"
                    
                    try:
                        subprocess.run([
                            'libreoffice', '--headless', '--convert-to', 'pdf', 
                            docx_file, '--outdir', os.getcwd()
                        ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    except Exception as e:
                        st.error(f"⚠️ PDF Conversion failed: {e}. Ensure 'libreoffice' is in packages.txt.")
                    
                    # 3. Provide the PDF for download (fallback to docx if PDF fails to generate)
                    if os.path.exists(pdf_file):
                        with open(pdf_file, "rb") as file:
                            btn = st.download_button(
                                label="⬇️ Download Generated PDF",
                                data=file,
                                file_name=pdf_file,
                                mime="application/pdf",
                                type="primary"
                            )
                        st.success(f"✅ Success! PDF Document is ready for download.")
                    else:
                        with open(docx_file, "rb") as file:
                            btn = st.download_button(
                                label="⬇️ Download Generated Word Doc (PDF Failed)",
                                data=file,
                                file_name=docx_file,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                type="primary"
                            )
                        st.warning("⚠️ Success, but PDF conversion failed. Provided Word Document instead.")

                    # 4. Cleanup the temporary files
                    try:
                        os.remove(resume_path)
                        if os.path.exists(docx_file):
                            os.remove(docx_file)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- 🟣 CLIENT APP 4: ADUSA 🟣 ---
# ====================================================================
def adusa_app():
    TEMPLATE_FILENAME = "ADUSA_Template.docx"
    st.title("ADUSA Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="ad_loc")
        
        Remote_or_Onsite = st.selectbox("Remote or Onsite/Hybrid", ["Remote", "Onsite/Hybrid"], index=0, key="ad_rem")
        ad_location_name = ""
        if Remote_or_Onsite == "Onsite/Hybrid":
            ad_location_name = st.text_input("Location Name (e.g., HQ2)", key="ad_loc_name")
            
        ad_start = st.text_input("Availability to Start", key="ad_start")

    with col2:
        Former_ADUSA = st.selectbox("Former ADUSA", ["Y", "N"], index=1, key="ad_form")
        ad_mgr = ""
        ad_dates = ""
        if Former_ADUSA == "Y":
            ad_mgr = st.text_input("Manager Name", key="ad_mgr")
            ad_dates = st.text_input("Dates Worked", key="ad_dates")
        # LinkedIn Removed for ADUSA

    if Remote_or_Onsite == "Onsite/Hybrid" and ad_location_name.strip():
        remote_val = f"Onsite/Hybrid - {ad_location_name.strip()}"
    else:
        remote_val = Remote_or_Onsite
        
    if Former_ADUSA == "Y":
        former_adusa_val = f"Y - {ad_mgr.strip()}, {ad_dates.strip()}"
    else:
        former_adusa_val = "N"

    st.header("🎤 Supplier Technical Interview Results")
    
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="ad_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="ad_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="ad_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="ad_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="ad_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="ad_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="ad_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="ad_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="ad_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="ad_a5")

    st.header("📝 Job Description & Notes")
    Job_Description_Notes_etc = st.text_area("Paste JD, Manager Notes, etc. here:", height=200, key="ad_jd")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="ad_res")

    if st.button("🚀 Generate ADUSA Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line.
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title (e.g., remove '- Contract', '(Contract)', or '- Consultant').
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title()
                    
                    mapping = {
                        "FullName": name, "Location": Current_Location_City_ST, "Remote": remote_val,
                        "Start": ad_start,
                        "Certifications": data.get("Certifications", ""),
                        "FormerADUSA": former_adusa_val, 
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", 
                        "SKILL1": "", "YEARS1": "", "SKILL2": "", "YEARS2": "",
                        "SKILL3": "", "YEARS3": "", "SKILL4": "", "YEARS4": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description_Notes_etc.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The Hiring Manager has 30 seconds to read this. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed.

                            ========================
                            THE NARRATIVE BLUEPRINT (4 Bullet Points)
                            ========================
                            Generate exactly 4 distinct bullet points mapped to SUMMARY1 through SUMMARY4. 
                            CRITICAL RULE: DO NOT use the candidate's name or 3rd person pronouns (he/she/they) in ANY of the bullet points. Start each bullet with a strong action verb or direct statement.
                            
                            - SUMMARY1 (The Anchor): What is their TOTAL progressive professional experience, and what is their dominant expertise that solves the PRIMARY technical "must-have" of the Job Description?
                            - SUMMARY2 (The Alignment): Explicitly name their CURRENT or most recent employer. What is the most impressive, relevant project they recently delivered that proves they can handle this specific job?
                            - SUMMARY3 (The Execution & Impact): How did they build it, and why does it matter? Weave specific tools/methodologies into an "Execution Statement" that highlights complexity, scale, or business impact.
                            - SUMMARY4 (The Closer): Based on their past execution, what specific value will they deliver on Day 1 in THIS new role?

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - LEADERSHIP VERBS: Use high-authority active verbs (e.g., 'Engineered', 'Optimized', 'Architected', 'Spearheaded').
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely."
                            - Do NOT repeat or restate the job description.

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth for experience.
                            - The technical interview Q&A is SECONDARY.
                            If there is any conflict: 👉 prioritize the resume over Q&A.

                            ========================
                            SKILLS SECTION
                            ========================
                            - EXACTLY 4 items
                            - Prioritize the specific "Must-Have" technologies.
                            - Use only tools explicitly mentioned in resume/Q&A/notes
                            Format: "Skill Area (Tool1, Tool2, Tool3, Tool4)"
                            Years format: "X+ years, current" OR "X+ years, 2026"

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME. 
                            3. DO NOT INFLATE TO MATCH THE JD. 
                            4. FACT AUDIT: Before outputting the final JSON, verify every single tool.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY1": "Bullet 1 following the blueprint",
                              "SUMMARY2": "Bullet 2 following the blueprint",
                              "SUMMARY3": "Bullet 3 following the blueprint",
                              "SUMMARY4": "Bullet 4 following the blueprint",
                              "SKILL1": "Skill Area (tools)",
                              "YEARS1": "X+ years, current",
                              "SKILL2": "Skill Area (tools)",
                              "YEARS2": "X+ years, current",
                              "SKILL3": "Skill Area (tools)",
                              "YEARS3": "X+ years, current",
                              "SKILL4": "Skill Area (tools)",
                              "YEARS4": "X+ years, current"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Job Description / Notes:
                            {Job_Description_Notes_etc}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["YEARS1"] = summary_data.get("YEARS1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["YEARS2"] = summary_data.get("YEARS2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["YEARS3"] = summary_data.get("YEARS3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")
                            mapping["YEARS4"] = summary_data.get("YEARS4", "")
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_ADUSA_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- 🟡 CLIENT APP 5: CBRE 🟡 ---
# ====================================================================
def cbre_app():
    TEMPLATE_FILENAME = "CBRE_Template.docx"
    st.title("CBRE Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="cb_loc")
        
        Remote_or_Onsite = st.selectbox("Remote or Onsite/Hybrid", ["Remote", "Onsite/Hybrid"], index=0, key="cb_rem")
        cb_location_name = ""
        if Remote_or_Onsite == "Onsite/Hybrid":
            cb_location_name = st.text_input("Location Name (e.g., HQ2)", key="cb_loc_name")
            
        cb_start = st.text_input("Availability to Start", key="cb_start")

    with col2:
        Former_CBRE = st.selectbox("Former CBRE", ["Y", "N"], index=1, key="cb_form")
        cb_mgr = ""
        cb_dates = ""
        if Former_CBRE == "Y":
            cb_mgr = st.text_input("Manager Name", key="cb_mgr")
            cb_dates = st.text_input("Dates Worked", key="cb_dates")
            
        LinkedIn_GitHub_Portfolio_Link = st.text_input("LinkedIn/GitHub/Portfolio Link", key="cb_link")

    if Remote_or_Onsite == "Onsite/Hybrid" and cb_location_name.strip():
        remote_val = f"Onsite/Hybrid - {cb_location_name.strip()}"
    else:
        remote_val = Remote_or_Onsite
        
    if Former_CBRE == "Y":
        former_cbre_val = f"Y - {cb_mgr.strip()}, {cb_dates.strip()}"
    else:
        former_cbre_val = "N"

    st.header("🎤 Supplier Technical Interview Results")
    
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="cb_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="cb_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="cb_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="cb_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="cb_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="cb_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="cb_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="cb_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="cb_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="cb_a5")

    st.header("📝 Job Description & Notes")
    Job_Description_Notes_etc = st.text_area("Paste JD, Manager Notes, etc. here:", height=200, key="cb_jd")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="cb_res")

    if st.button("🚀 Generate CBRE Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line.
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title (e.g., remove '- Contract', '(Contract)', or '- Consultant').
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title()
                    
                    mapping = {
                        "FullName": name, "Location": Current_Location_City_ST, "Remote": remote_val,
                        "Start": cb_start,
                        "Certifications": data.get("Certifications", ""),
                        "FormerCBRE": former_cbre_val, "Links": LinkedIn_GitHub_Portfolio_Link,
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", 
                        "SKILL1": "", "YEARS1": "", "SKILL2": "", "YEARS2": "",
                        "SKILL3": "", "YEARS3": "", "SKILL4": "", "YEARS4": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description_Notes_etc.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The Hiring Manager has 30 seconds to read this. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed.

                            ========================
                            THE NARRATIVE BLUEPRINT (4 Bullet Points)
                            ========================
                            Generate exactly 4 distinct bullet points mapped to SUMMARY1 through SUMMARY4. 
                            CRITICAL RULE: DO NOT use the candidate's name or 3rd person pronouns (he/she/they) in ANY of the bullet points. Start each bullet with a strong action verb or direct statement.
                            
                            - SUMMARY1 (The Anchor): What is their TOTAL progressive professional experience, and what is their dominant expertise that solves the PRIMARY technical "must-have" of the Job Description?
                            - SUMMARY2 (The Alignment): Explicitly name their CURRENT or most recent employer. What is the most impressive, relevant project they recently delivered that proves they can handle this specific job?
                            - SUMMARY3 (The Execution & Impact): How did they build it, and why does it matter? Weave specific tools/methodologies into an "Execution Statement" that highlights complexity, scale, or business impact.
                            - SUMMARY4 (The Closer): Based on their past execution, what specific value will they deliver on Day 1 in THIS new role?

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - LEADERSHIP VERBS: Use high-authority active verbs (e.g., 'Engineered', 'Optimized', 'Architected', 'Spearheaded').
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely."
                            - Do NOT repeat or restate the job description.

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth for experience.
                            - The technical interview Q&A is SECONDARY.
                            If there is any conflict: 👉 prioritize the resume over Q&A.

                            ========================
                            SKILLS SECTION
                            ========================
                            - EXACTLY 4 items
                            - Prioritize the specific "Must-Have" technologies.
                            - Use only tools explicitly mentioned in resume/Q&A/notes
                            Format: "Skill Area (Tool1, Tool2, Tool3, Tool4)"
                            Years format: "X+ years, current" OR "X+ years, 2026"

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME. 
                            3. DO NOT INFLATE TO MATCH THE JD. 
                            4. FACT AUDIT: Before outputting the final JSON, verify every single tool.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY1": "Bullet 1 following the blueprint",
                              "SUMMARY2": "Bullet 2 following the blueprint",
                              "SUMMARY3": "Bullet 3 following the blueprint",
                              "SUMMARY4": "Bullet 4 following the blueprint",
                              "SKILL1": "Skill Area (tools)",
                              "YEARS1": "X+ years, current",
                              "SKILL2": "Skill Area (tools)",
                              "YEARS2": "X+ years, current",
                              "SKILL3": "Skill Area (tools)",
                              "YEARS3": "X+ years, current",
                              "SKILL4": "Skill Area (tools)",
                              "YEARS4": "X+ years, current"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Job Description / Notes:
                            {Job_Description_Notes_etc}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["YEARS1"] = summary_data.get("YEARS1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["YEARS2"] = summary_data.get("YEARS2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["YEARS3"] = summary_data.get("YEARS3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")
                            mapping["YEARS4"] = summary_data.get("YEARS4", "")
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_CBRE_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- 🟠 CLIENT APP 6: BNSF 🟠 ---
# ====================================================================
def bnsf_app():
    TEMPLATE_FILENAME = "BNSF_Template.docx"
    st.title("BNSF Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="bn_loc")
        
        Remote_or_Onsite = st.selectbox("Remote or Onsite/Hybrid", ["Remote", "Onsite/Hybrid"], index=0, key="bn_rem")
        bn_location_name = ""
        if Remote_or_Onsite == "Onsite/Hybrid":
            bn_location_name = st.text_input("Location Name (e.g., HQ2)", key="bn_loc_name")
            
        bn_start = st.text_input("Availability to Start", key="bn_start")
        
        # --- NEW BNSF SPECIFIC FIELDS ---
        Bill_Rate = st.text_input("Bill Rate", key="bn_rate")
        Work_Status = st.selectbox("Work Status", ["USC", "GC", "Visa"], index=0, key="bn_status")

    with col2:
        Former_BNSF = st.selectbox("Former BNSF", ["Y", "N"], index=1, key="bn_form")
        bn_mgr = ""
        bn_dates = ""
        if Former_BNSF == "Y":
            bn_mgr = st.text_input("Manager Name", key="bn_mgr")
            bn_dates = st.text_input("Dates Worked", key="bn_dates")
            
        LinkedIn_GitHub_Portfolio_Link = st.text_input("LinkedIn/GitHub/Portfolio Link", key="bn_link")

    if Remote_or_Onsite == "Onsite/Hybrid" and bn_location_name.strip():
        remote_val = f"Onsite/Hybrid - {bn_location_name.strip()}"
    else:
        remote_val = Remote_or_Onsite
        
    if Former_BNSF == "Y":
        former_bnsf_val = f"Y - {bn_mgr.strip()}, {bn_dates.strip()}"
    else:
        former_bnsf_val = "N"

    st.header("🎤 Supplier Technical Interview Results")
    
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="bn_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="bn_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="bn_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="bn_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="bn_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="bn_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="bn_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="bn_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="bn_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="bn_a5")

    st.header("📝 Job Description & Notes")
    Job_Description_Notes_etc = st.text_area("Paste JD, Manager Notes, etc. here:", height=200, key="bn_jd")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="bn_res")

    if st.button("🚀 Generate BNSF Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line.
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title (e.g., remove '- Contract', '(Contract)', or '- Consultant').
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title()
                    
                    mapping = {
                        "FullName": name, "Location": Current_Location_City_ST, "Remote": remote_val,
                        "Start": bn_start,
                        "Certifications": data.get("Certifications", ""),
                        "FormerBNSF": former_bnsf_val, "Links": LinkedIn_GitHub_Portfolio_Link,
                        "BillRate": Bill_Rate, "WorkStatus": Work_Status,
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", 
                        "SKILL1": "", "YEARS1": "", "SKILL2": "", "YEARS2": "",
                        "SKILL3": "", "YEARS3": "", "SKILL4": "", "YEARS4": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description_Notes_etc.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The Hiring Manager has 30 seconds to read this. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed.

                            ========================
                            THE NARRATIVE BLUEPRINT (4 Bullet Points)
                            ========================
                            Generate exactly 4 distinct bullet points mapped to SUMMARY1 through SUMMARY4. 
                            CRITICAL RULE: DO NOT use the candidate's name or 3rd person pronouns (he/she/they) in ANY of the bullet points. Start each bullet with a strong action verb or direct statement.
                            
                            - SUMMARY1 (The Anchor): What is their TOTAL progressive professional experience, and what is their dominant expertise that solves the PRIMARY technical "must-have" of the Job Description?
                            - SUMMARY2 (The Alignment): Explicitly name their CURRENT or most recent employer. What is the most impressive, relevant project they recently delivered that proves they can handle this specific job?
                            - SUMMARY3 (The Execution & Impact): How did they build it, and why does it matter? Weave specific tools/methodologies into an "Execution Statement" that highlights complexity, scale, or business impact.
                            - SUMMARY4 (The Closer): Based on their past execution, what specific value will they deliver on Day 1 in THIS new role?

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - LEADERSHIP VERBS: Use high-authority active verbs (e.g., 'Engineered', 'Optimized', 'Architected', 'Spearheaded').
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely."
                            - Do NOT repeat or restate the job description.

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth for experience.
                            - The technical interview Q&A is SECONDARY.
                            If there is any conflict: 👉 prioritize the resume over Q&A.

                            ========================
                            SKILLS SECTION
                            ========================
                            - EXACTLY 4 items
                            - Prioritize the specific "Must-Have" technologies.
                            - Use only tools explicitly mentioned in resume/Q&A/notes
                            Format: "Skill Area (Tool1, Tool2, Tool3, Tool4)"
                            Years format: "X+ years, current" OR "X+ years, 2026"

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME. 
                            3. DO NOT INFLATE TO MATCH THE JD. 
                            4. FACT AUDIT: Before outputting the final JSON, verify every single tool.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY1": "Bullet 1 following the blueprint",
                              "SUMMARY2": "Bullet 2 following the blueprint",
                              "SUMMARY3": "Bullet 3 following the blueprint",
                              "SUMMARY4": "Bullet 4 following the blueprint",
                              "SKILL1": "Skill Area (tools)",
                              "YEARS1": "X+ years, current",
                              "SKILL2": "Skill Area (tools)",
                              "YEARS2": "X+ years, current",
                              "SKILL3": "Skill Area (tools)",
                              "YEARS3": "X+ years, current",
                              "SKILL4": "Skill Area (tools)",
                              "YEARS4": "X+ years, current"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Job Description / Notes:
                            {Job_Description_Notes_etc}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["YEARS1"] = summary_data.get("YEARS1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["YEARS2"] = summary_data.get("YEARS2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["YEARS3"] = summary_data.get("YEARS3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")
                            mapping["YEARS4"] = summary_data.get("YEARS4", "")
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_BNSF_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- ⬜ CLIENT APP 7: DALLAS GENERIC ⬜ ---
# ====================================================================
def dallas_generic_app():
    TEMPLATE_FILENAME = "DallasGeneric_Template.docx"
    st.title("Dallas Generic Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="dg_loc")
        
        Remote_or_Onsite = st.selectbox("Remote or Onsite/Hybrid", ["Remote", "Onsite/Hybrid"], index=0, key="dg_rem")
        dg_location_name = ""
        if Remote_or_Onsite == "Onsite/Hybrid":
            dg_location_name = st.text_input("Location Name (e.g., HQ2)", key="dg_loc_name")
            
        dg_start = st.text_input("Availability to Start", key="dg_start")

    with col2:
        LinkedIn_GitHub_Portfolio_Link = st.text_input("LinkedIn/GitHub/Portfolio Link", key="dg_link")

    if Remote_or_Onsite == "Onsite/Hybrid" and dg_location_name.strip():
        remote_val = f"Onsite/Hybrid - {dg_location_name.strip()}"
    else:
        remote_val = Remote_or_Onsite

    st.header("🎤 Supplier Technical Interview Results")
    
    # Row 1
    row1_c1, row1_c2 = st.columns(2)
    Question_1 = row1_c1.text_input("Question 1", key="dg_q1")
    Answer_1 = row1_c2.text_area("Answer 1", height=68, key="dg_a1")

    # Row 2
    row2_c1, row2_c2 = st.columns(2)
    Question_2 = row2_c1.text_input("Question 2", key="dg_q2")
    Answer_2 = row2_c2.text_area("Answer 2", height=68, key="dg_a2")

    # Row 3
    row3_c1, row3_c2 = st.columns(2)
    Question_3 = row3_c1.text_input("Question 3", key="dg_q3")
    Answer_3 = row3_c2.text_area("Answer 3", height=68, key="dg_a3")

    # Row 4
    row4_c1, row4_c2 = st.columns(2)
    Question_4 = row4_c1.text_input("Question 4", key="dg_q4")
    Answer_4 = row4_c2.text_area("Answer 4", height=68, key="dg_a4")

    # Row 5
    row5_c1, row5_c2 = st.columns(2)
    Question_5 = row5_c1.text_input("Question 5", key="dg_q5")
    Answer_5 = row5_c2.text_area("Answer 5", height=68, key="dg_a5")

    st.header("📝 Job Description & Notes")
    Job_Description_Notes_etc = st.text_area("Paste JD, Manager Notes, etc. here:", height=200, key="dg_jd")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="dg_res")

    if st.button("🚀 Generate Dallas Generic Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found. Please make sure it is uploaded to your GitHub repository.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors (e.g., "TCS", "Cognizant") that share the same line.
                    3. Education: Extract School, Degree, and DegreeStatus.
                        - DegreeStatus: Evaluate if the degree is completed or in progress.
                        - Return "Yes" if the resume indicates the degree is finished.
                        - Return "Pursuing" if the resume indicates ongoing study, contains the word "Pursuing", "In-progress", or lists a graduation date in the future (relative to July 2026).
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title (e.g., remove '- Contract', '(Contract)', or '- Consultant').
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": "", "DegreeStatus": "Yes"}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = [
                        'gemini-2.5-flash', 
                        'gemini-3-flash-preview', 
                        'gemini-3.1-flash-lite-preview', 
                        'gemini-pro-latest'
                    ]
                    data = {}
                    
                    for attempt in range(6):
                        if attempt == 0:
                            current_model = models_to_try[0]
                        elif attempt in [1, 2]:
                            current_model = models_to_try[1]
                        elif attempt in [3, 4]:
                            current_model = models_to_try[2]
                        else:
                            current_model = models_to_try[3]
                        
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(
                                    response_mime_type="application/json"
                                )
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    name = data.get('FullName', '').title()
                    
                    mapping = {
                        "FullName": name, "Location": Current_Location_City_ST, "Remote": remote_val,
                        "Start": dg_start,
                        "Certifications": data.get("Certifications", ""),
                        "Links": LinkedIn_GitHub_Portfolio_Link, 
                        "Q1": Question_1, "A1": Answer_1, "Q2": Question_2, "A2": Answer_2,
                        "Q3": Question_3, "A3": Answer_3, "Q4": Question_4, "A4": Answer_4, "Q5": Question_5, "A5": Answer_5,
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", 
                        "SKILL1": "", "YEARS1": "", "SKILL2": "", "YEARS2": "",
                        "SKILL3": "", "YEARS3": "", "SKILL4": "", "YEARS4": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""
                        mapping[f"DegreeStatus{i}"] = edu[i-1].get('DegreeStatus', 'Yes') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description_Notes_etc.strip():
                        try:
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The Hiring Manager has 30 seconds to read this. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed.

                            ========================
                            THE NARRATIVE BLUEPRINT (4 Bullet Points)
                            ========================
                            Generate exactly 4 distinct bullet points mapped to SUMMARY1 through SUMMARY4. 
                            CRITICAL RULE: DO NOT use the candidate's name or 3rd person pronouns (he/she/they) in ANY of the bullet points. Start each bullet with a strong action verb or direct statement.
                            
                            - SUMMARY1 (The Anchor): What is their TOTAL progressive professional experience, and what is their dominant expertise that solves the PRIMARY technical "must-have" of the Job Description?
                            - SUMMARY2 (The Alignment): Explicitly name their CURRENT or most recent employer. What is the most impressive, relevant project they recently delivered that proves they can handle this specific job?
                            - SUMMARY3 (The Execution & Impact): How did they build it, and why does it matter? Weave specific tools/methodologies into an "Execution Statement" that highlights complexity, scale, or business impact.
                            - SUMMARY4 (The Closer): Based on their past execution, what specific value will they deliver on Day 1 in THIS new role?

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - LEADERSHIP VERBS: Use high-authority active verbs (e.g., 'Engineered', 'Optimized', 'Architected', 'Spearheaded').
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely."
                            - Do NOT repeat or restate the job description.

                            ========================
                            🔴 Q&A USAGE RULES (CRITICAL)
                            ========================
                            - The resume is the PRIMARY source of truth for experience.
                            - The technical interview Q&A is SECONDARY.
                            If there is any conflict: 👉 prioritize the resume over Q&A.

                            ========================
                            SKILLS SECTION
                            ========================
                            - EXACTLY 4 items
                            - Prioritize the specific "Must-Have" technologies.
                            - Use only tools explicitly mentioned in resume/Q&A/notes
                            Format: "Skill Area (Tool1, Tool2, Tool3, Tool4)"
                            Years format: "X+ years, current" OR "X+ years, 2026"

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume or Q&A.
                            2. DO NOT INFER OR ASSUME. 
                            3. DO NOT INFLATE TO MATCH THE JD. 
                            4. FACT AUDIT: Before outputting the final JSON, verify every single tool.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON:

                            {{
                              "SUMMARY1": "Bullet 1 following the blueprint",
                              "SUMMARY2": "Bullet 2 following the blueprint",
                              "SUMMARY3": "Bullet 3 following the blueprint",
                              "SUMMARY4": "Bullet 4 following the blueprint",
                              "SKILL1": "Skill Area (tools)",
                              "YEARS1": "X+ years, current",
                              "SKILL2": "Skill Area (tools)",
                              "YEARS2": "X+ years, current",
                              "SKILL3": "Skill Area (tools)",
                              "YEARS3": "X+ years, current",
                              "SKILL4": "Skill Area (tools)",
                              "YEARS4": "X+ years, current"
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Job Description / Notes:
                            {Job_Description_Notes_etc}

                            Technical Interview Q&A:
                            Q1: {Question_1}
                            A1: {Answer_1}
                            Q2: {Question_2}
                            A2: {Answer_2}
                            Q3: {Question_3}
                            A3: {Answer_3}
                            Q4: {Question_4}
                            A4: {Answer_4}
                            Q5: {Question_5}
                            A5: {Answer_5}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                if attempt == 0:
                                    current_model = models_to_try[0]
                                elif attempt in [1, 2]:
                                    current_model = models_to_try[1]
                                elif attempt in [3, 4]:
                                    current_model = models_to_try[2]
                                else:
                                    current_model = models_to_try[3]
                                
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(
                                            response_mime_type="application/json"
                                        )
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SKILL1"] = summary_data.get("SKILL1", "")
                            mapping["YEARS1"] = summary_data.get("YEARS1", "")
                            mapping["SKILL2"] = summary_data.get("SKILL2", "")
                            mapping["YEARS2"] = summary_data.get("YEARS2", "")
                            mapping["SKILL3"] = summary_data.get("SKILL3", "")
                            mapping["YEARS3"] = summary_data.get("YEARS3", "")
                            mapping["SKILL4"] = summary_data.get("SKILL4", "")
                            mapping["YEARS4"] = summary_data.get("YEARS4", "")
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_DallasGeneric_{name.replace(' ', '_')}.docx"
                    process_word_doc(TEMPLATE_FILENAME, mapping, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- 📄 NEW APP: PDF TO WORD CONVERTER 📄 ---
# ====================================================================
def pdf_to_word_app():
    st.title("PDF to Word Converter")
    st.markdown("Upload a PDF resume below to convert it to a Word Document (.docx) without any AI formatting or content changes.")

    st.header("📂 File Upload")
    resume_file = st.file_uploader("📤 Upload PDF Resume...", type=['pdf'], key="pdf2word_res")

    if st.button("🔄 Convert to Word (Adobe High Fidelity)", type="primary"):
        if not resume_file:
            st.error("❌ Error: Please upload a PDF file.")
        else:
            with st.spinner("Converting via Adobe PDF Services... this may take 15-30 seconds."):
                try:
                    # Pull Adobe credentials from Streamlit secrets
                    client_id = st.secrets["PDF_SERVICES_CLIENT_ID"]
                    client_secret = st.secrets["PDF_SERVICES_CLIENT_SECRET"]

                    # Read PDF bytes directly from uploader (no temp file needed)
                    pdf_bytes = resume_file.read()

                    # Call Adobe converter
                    from adobe_converter import convert_pdf_to_docx
                    docx_bytes = convert_pdf_to_docx(pdf_bytes, client_id, client_secret)

                    docx_filename = f"Converted_{resume_file.name.replace('.pdf', '')}.docx"

                    st.download_button(
                        label="⬇️ Download Word Document",
                        data=docx_bytes,
                        file_name=docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )
                    st.success("✅ Success! Document converted with full formatting preserved.")

                except KeyError:
                    st.error("❌ Adobe API credentials not found in Streamlit Secrets. Please add PDF_SERVICES_CLIENT_ID and PDF_SERVICES_CLIENT_SECRET.")
                except RuntimeError as e:
                    st.error(f"❌ Adobe Conversion Failed: {str(e)}")
                except Exception as e:
                    st.error(f"❌ Unexpected Error: {str(e)}")


# ====================================================================
# --- 🔵 CLIENT APP 8: DELOITTE 🔵 ---
# ====================================================================
def process_deloitte_doc(template_path, mapping, resume_path, output_path):
    import docx
    import re
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_TAB_ALIGNMENT
    from copy import deepcopy

    doc = docx.Document(template_path)
    orig_doc = docx.Document(resume_path)

    # 1. Delete empty skill rows BEFORE replacing text
    for tbl in doc.tables:
        rows_to_delete = []
        for row in tbl.rows:
            cell_text = row.cells[0].text
            for i in range(1, 9):
                if f"{{{{SKILL{i}}}}}" in cell_text:
                    if not mapping.get(f"SKILL{i}", "").strip():
                        rows_to_delete.append(row)
        for row in rows_to_delete:
            # Safer XML removal method for table rows
            row._element.getparent().remove(row._element)

    # 2. Replace placeholders while preserving formatting
    for p in doc.paragraphs:
        # Explicitly rebuild the header lines to maintain the bold prefix
        if "Candidate Legal Name:" in p.text and "{{FullName}}" in p.text:
            p.text = ""
            r1 = p.add_run("Candidate Legal Name: ")
            r1.bold = True
            p.add_run(str(mapping.get("FullName", "")))
            continue
        if "Preferred Name:" in p.text and "{{PreferredName}}" in p.text:
            p.text = ""
            r1 = p.add_run("Preferred Name: ")
            r1.bold = True
            p.add_run(str(mapping.get("PreferredName", "")))
            continue
        if "Current Location:" in p.text and "{{Location}}" in p.text:
            p.text = ""
            r1 = p.add_run("Current Location: ")
            r1.bold = True
            p.add_run(str(mapping.get("Location", "")))
            continue
        if "Upcoming Scheduled Time off:" in p.text and "{{TimeOff}}" in p.text:
            p.text = ""
            r1 = p.add_run("Upcoming Scheduled Time off: ")
            r1.bold = True
            p.add_run(str(mapping.get("TimeOff", "")))
            continue

        # For all other paragraphs (like Summary points)
        for k, v in mapping.items():
            placeholder = f"{{{{{k}}}}}"
            if placeholder in p.text:
                replaced = False
                for run in p.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(v))
                        replaced = True
                # Fallback if the placeholder got split across multiple runs by Word
                if not replaced:
                    p.text = p.text.replace(placeholder, str(v))

    # Replace placeholders in table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for k, v in mapping.items():
                        placeholder = f"{{{{{k}}}}}"
                        if placeholder in p.text:
                            replaced = False
                            for run in p.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, str(v))
                                    replaced = True
                            if not replaced:
                                p.text = p.text.replace(placeholder, str(v))

    # --- XML SCRUBBER FUNCTION ---
    def clean_xml_element(element):
        for hl in element.xpath('.//w:hyperlink'):
            parent = hl.getparent()
            if parent is not None:
                idx = parent.index(hl)
                for child in list(hl):
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(hl)
        for drawing in element.xpath('.//w:drawing | .//w:pict | .//w:object'):
            parent = drawing.getparent()
            if parent is not None:
                parent.remove(drawing)
        for bm in element.xpath('.//w:bookmarkStart | .//w:bookmarkEnd'):
            parent = bm.getparent()
            if parent is not None:
                parent.remove(bm)
        for sectPr in element.xpath('.//w:sectPr'):
            parent = sectPr.getparent()
            if parent is not None:
                parent.remove(sectPr)
        return element
    # ----------------------------------

    # 3. Extract the "Rest of Resume"
    body_elements = orig_doc._body._body
    start_copying = False
    elements_to_copy = []

    for element in body_elements:
        try:
            text = "".join(element.itertext())
        except Exception:
            text = ""
            
        if "Links:" in text or "Technical Skills" in text:
            start_copying = True
        
        if start_copying:
            clean_el = clean_xml_element(deepcopy(element))
            elements_to_copy.append(clean_el)

    # 4. Insert extracted elements perfectly at {{RESUME_BODY}}
    for p in doc.paragraphs:
        if "{{RESUME_BODY}}" in p.text:
            parent = p._element.getparent()
            index = parent.index(p._element)
            for el in elements_to_copy:
                parent.insert(index, el)
                index += 1
            parent.remove(p._element)
            break

    # 5. STRICT FORMATTING & TAB STOP FIX
    def apply_strict_formatting(paragraph):
        # A. Set Tab Stop to 7.5" Right-Aligned and remove right indent margins
        try:
            paragraph.paragraph_format.tab_stops.clear_all()
            paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
            paragraph.paragraph_format.right_indent = Inches(0)
        except Exception:
            pass

        # B. Clean up runs: Set Arial 10pt and collapse multiple tabs/spaces
        for run in paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            if run.text:
                # Convert 3 or more spaces to a single tab
                run.text = re.sub(r' {3,}', '\t', run.text)
                # Collapse multiple tabs into a single tab
                run.text = re.sub(r'\t{2,}', '\t', run.text)
                # Remove spaces that are directly next to tabs
                run.text = re.sub(r' \t|\t ', '\t', run.text)

        # C. Cross-Run Cleanup: Catch tabs that are split across formatting boundaries
        for i in range(len(paragraph.runs) - 1):
            if paragraph.runs[i].text and paragraph.runs[i].text.endswith('\t'):
                j = i + 1
                while j < len(paragraph.runs) and paragraph.runs[j].text.startswith('\t'):
                    paragraph.runs[j].text = paragraph.runs[j].text.lstrip('\t')
                    j += 1

    # Apply the formatting sweep across the entire generated document
    for p in doc.paragraphs:
        apply_strict_formatting(p)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    apply_strict_formatting(p)

    # Collapse extra blank body paragraphs left by removed/empty sections.
    collapse_extra_blank_paragraphs(doc)
    
    doc.save(output_path)
def deloitte_app():
    TEMPLATE_FILENAME = "Deloitte_Template.docx"
    st.title("Deloitte Precision Extractor")

    with st.sidebar:
        st.header("🔑 API Configuration")
        if "API_KEY" in st.secrets:
            API_KEY = st.secrets["API_KEY"]
            st.success("✅ API Key loaded from Secrets")
        else:
            API_KEY = st.text_input("Gemini API Key", type="password")
            st.info("Paste your Gemini API key here to run the tool.")

    st.header("📋 Candidate Information")
    col1, col2 = st.columns(2)
    with col1:
        Legal_Name = st.text_input("Candidate Legal Name", key="del_legal")
        Current_Location_City_ST = st.text_input("Current Location (City, ST)", key="del_loc")
    with col2:
        Preferred_Name = st.text_input("Preferred Name", key="del_pref")
        Time_Off = st.text_input("Upcoming Scheduled Time off", key="del_toff")

    st.header("📝 Job Description & Notes")

    # --- GOOGLE SHEET LOGIC FOR DELOITTE ---
    SHEET_URL = "https://docs.google.com/spreadsheets/d/1swKaDhGdlmO0ILv7tngkRCbsK-7jF89cdCpf5bypzwM/edit?gid=2029588481#gid=2029588481"

    @st.cache_data(ttl=600)
    def load_deloitte_reqs(url):
        try:
            # Use the exact GID for the Deloitte sheet
            base_url = url.split("/edit")[0]
            csv_url = base_url + "/export?format=csv&gid=2029588481"
            df = pd.read_csv(csv_url)
            return df.fillna("")
        except Exception:
            return pd.DataFrame()

    df_reqs = load_deloitte_reqs(SHEET_URL)
    options = ["None"]
    req_mapping = {}

    if not df_reqs.empty:
        for idx, row in df_reqs.iterrows():
            # Column A (index 0) = ID, Column B (index 1) = VMS ID, Column C (index 2) = Job Title
            req_id = str(row.iloc[0]).replace(".0", "") if len(row) > 0 else ""
            vms_id = str(row.iloc[1]).replace(".0", "") if len(row) > 1 else ""
            title = str(row.iloc[2]) if len(row) > 2 else ""
            
            # Column I (index 8) = Job Description
            desc = str(row.iloc[8]) if len(row) > 8 else ""

            # Skills starting from Column J (index 9) to the end of the row
            skills = []
            if len(row) > 9:
                for val in row.iloc[9:]:
                    val_str = str(val).strip()
                    if val_str:
                        skills.append(val_str)
            skills_matrix_text = "\n".join(skills)

            option_parts = [req_id, vms_id, title]
            option_str = " - ".join([p.strip() for p in option_parts if p.strip()])
            
            if option_str and option_str != "-":
                options.append(option_str)
                req_mapping[option_str] = {"desc": desc, "skills": skills_matrix_text}

    if "del_jd" not in st.session_state:
        st.session_state.del_jd = ""
    if "del_notes" not in st.session_state:
        st.session_state.del_notes = ""
    if "del_skills" not in st.session_state:
        st.session_state.del_skills = ""

    def update_jd_text():
        if "del_req_selector" not in st.session_state:
            return
        selected = st.session_state.del_req_selector
        if selected != "None" and selected in req_mapping:
            st.session_state.del_jd = req_mapping[selected].get("desc", "").strip()
            st.session_state.del_skills = req_mapping[selected].get("skills", "").strip()
            st.session_state.del_notes = ""
        else:
            st.session_state.del_jd = ""
            st.session_state.del_skills = ""
            st.session_state.del_notes = ""
            
    selected_req = st.selectbox(
        "🔍 Select a Deloitte Requisition (Type to search):",
        options=options,
        key="del_req_selector",
        on_change=update_jd_text
    )

    Job_Description = st.text_area("1. Official Job Description (Paste generic JD here):", height=150, key="del_jd")
    Chat_Notes = st.text_area("2. Notes (Optional - Overrides JD):", height=100, key="del_notes")
    Skills_Matrix = st.text_area("3. Skills Matrix:", height=100, key="del_skills")

    st.header("📂 File Uploads")
    resume_file = st.file_uploader("📤 Upload the candidate's resume...", type=['pdf', 'docx', 'doc'], key="del_res")

    if st.button("🚀 Generate Deloitte Submission", type="primary"):
        if not API_KEY:
            st.error("❌ Error: Please enter your Gemini API Key in the sidebar.")
        elif not resume_file:
            st.error("❌ Error: Please upload a resume.")
        elif not os.path.exists(TEMPLATE_FILENAME):
            st.error(f"❌ Error: The preloaded template '{TEMPLATE_FILENAME}' was not found.")
        else:
            with st.spinner("Processing document and querying AI..."):
                try:
                    resume_path = f"temp_{resume_file.name}"
                    with open(resume_path, "wb") as f:
                        f.write(resume_file.getbuffer())

                    raw_text = extract_text(resume_path)
                    client = genai.Client(api_key=API_KEY)

                    prompt = f"""
                    Return a valid JSON object ONLY.

                    RULES:
                    1. Name: Pull from top/header (Title Case). If the name is completely missing or unreadable in the text, extract it from the provided FILENAME.
                    2. Company Name Cleaning: Extract ONLY the primary end-client name. You MUST physically strip out any geographic locations (e.g., ", DELAWARE", ", MD") and strip out any contracting agencies/vendors.
                    3. Education: Extract School and Degree.
                    4. Experience: Company, Title, Bullets (LIST), Environment (String, optional), Dates.
                        - For 'Dates', preserve the date precision from the original resume. If a role lists only years, return only those years (for example, "2018 - 2020"). NEVER invent January, "01", or any other month. If month + year are present, preserve them; mixed-precision ranges must stay mixed.
                        - For 'Title', clean the string by physically stripping out any employment type modifiers, hyphens, or parentheses at the end of the title.
                        - For 'Environment', you may ONLY extract this if the original resume explicitly uses the word "Environment:" or "Technologies:" at the bottom of the role. If those exact words are not there, you MUST leave it blank "". Do NOT auto-generate or compile an environment list from the bullet points.
                    5. Certifications: Extract any certifications listed into a single comma-separated string. If none are found, leave it blank "".

                    JSON Structure:
                    {{
                        "FullName": "",
                        "Certifications": "",
                        "Education": [{{"School": "", "Degree": ""}}],
                        "Experience": [{{"Company": "", "Title": "", "Bullets": [], "Environment": "", "Dates": ""}}]
                    }}

                    RESUME: {raw_text}
                    """

                    models_to_try = ['gemini-2.5-flash', 'gemini-3-flash-preview', 'gemini-3.1-flash-lite-preview', 'gemini-pro-latest']
                    data = {}
                    
                    for attempt in range(6):
                        current_model = models_to_try[0] if attempt == 0 else (models_to_try[1] if attempt in [1, 2] else (models_to_try[2] if attempt in [3, 4] else models_to_try[3]))
                        try:
                            response = client.models.generate_content(
                                model=current_model,
                                contents=prompt,
                                config=types.GenerateContentConfig(response_mime_type="application/json")
                            )
                            data = repair_and_load_json(response.text)
                            break
                        except Exception as e:
                            if "503" in str(e) and attempt < 5:
                                time.sleep(2 ** ((attempt % 2) + 1))
                                continue
                            else:
                                raise e

                    ai_name = data.get('FullName', '').title()
                    final_name = Legal_Name.strip() if Legal_Name.strip() else ai_name

                    # --- UPDATED MAPPING TO SUPPORT 8 SKILLS + COMPANIES + 5 SUMMARIES ---
                    mapping = {
                        "FullName": final_name, "PreferredName": Preferred_Name, 
                        "TimeOff": Time_Off, "Location": Current_Location_City_ST,
                        "SUMMARY1": "", "SUMMARY2": "", "SUMMARY3": "", "SUMMARY4": "", "SUMMARY5": "",
                        "SKILL1": "", "YEARS1": "", "SKILL1COMPANIES": "",
                        "SKILL2": "", "YEARS2": "", "SKILL2COMPANIES": "",
                        "SKILL3": "", "YEARS3": "", "SKILL3COMPANIES": "",
                        "SKILL4": "", "YEARS4": "", "SKILL4COMPANIES": "",
                        "SKILL5": "", "YEARS5": "", "SKILL5COMPANIES": "",
                        "SKILL6": "", "YEARS6": "", "SKILL6COMPANIES": "",
                        "SKILL7": "", "YEARS7": "", "SKILL7COMPANIES": "",
                        "SKILL8": "", "YEARS8": "", "SKILL8COMPANIES": ""
                    }

                    edu = data.get('Education', [])
                    for i in range(1, 4):
                        mapping[f"School{i}"] = clean_school(edu[i-1].get('School', '')) if i <= len(edu) else ""
                        mapping[f"Degree{i}"] = edu[i-1].get('Degree', '') if i <= len(edu) else ""

                    exp = data.get('Experience', [])
                    for i in range(1, max(7, len(exp)) + 1):
                        if i <= len(exp):
                            mapping[f"Company{i}"] = exp[i-1].get('Company', '')
                            raw_title = exp[i-1].get('Title', '')
                            clean_title = re.sub(r'\s*\(.*$', '', raw_title).strip()
                            mapping[f"Title{i}"] = clean_title
                            mapping[f"Bullets{i}"] = clean_bullets(exp[i-1].get('Bullets', []))
                            mapping[f"Environment{i}"] = exp[i-1].get('Environment', '')
                            mapping[f"Dates{i}"] = standardize_dates(exp[i-1].get('Dates', ''))
                        else:
                            mapping[f"Company{i}"] = ""
                            mapping[f"Title{i}"] = ""
                            mapping[f"Bullets{i}"] = ""
                            mapping[f"Environment{i}"] = ""
                            mapping[f"Dates{i}"] = ""

                    if Job_Description.strip():
                        try:
                            # --- UPDATED SUMMARY PROMPT FOR DELOITTE ---
                            summary_prompt = f"""
                            You are an elite, no-nonsense Senior Technical Recruiter writing an executive submission summary for a Fieldglass portal. 
                            The Hiring Manager has 30 seconds to read this. Your goal is to make a punchy, evidence-based business case for why this candidate will succeed, optimized for both human reading and ATS exact-match parsing.

                            ========================
                            🚨 HIERARCHY OF TRUTH: JD vs. MANAGER NOTES (CRITICAL)
                            ========================
                            You are receiving two primary sources of information for this role:
                            1. Notes (HIGH PRIORITY)
                            2. Official Job Description (BASELINE)

                            - OVERRIDE RULE: The Notes represent the Hiring Manager's actual, immediate needs. They are the ABSOLUTE SOURCE OF TRUTH. If they contradict or update the official Job Description, you MUST strictly follow the Notes.
                            - ATS COMPROMISE: To satisfy the Fieldglass parsing algorithm, you must still weave in high-level, generalized keywords from the baseline Job Description where applicable, but NEVER highlight technical skills or responsibilities that the manager explicitly negated in their notes.

                            ========================
                            THE NARRATIVE BLUEPRINT (5 Bullet Points)
                            ========================
                            Follow this exact structure for the SUMMARY. Generate exactly 5 distinct bullet points mapped to SUMMARY1 through SUMMARY5. Every point MUST sell the candidate's fit for the role:
                            - SUMMARY1 (The Anchor): Use their FIRST NAME. State their total progressive professional experience, their core title, and dominant expertise solving the PRIMARY technical "must-have" of the Job Description. (Format years strictly rounded DOWN to the nearest whole year as 'X+ years').
                            - SUMMARY2 (The Alignment): NO PRONOUNS. Start directly with an action verb or adverb (e.g., "Recently architected..."). Explicitly name their CURRENT or most recent employer and frame their most relevant project as a 1:1 match for the manager's challenge.
                            - SUMMARY3 (The Execution): NO PRONOUNS. Start directly with an action verb (e.g., "Engineered..."). Weave the specific tools/methodologies into an "Execution Statement" highlighting complexity, scale, or business impact.
                            - SUMMARY4 (Secondary Expertise): NO PRONOUNS. Start directly with a verb (e.g., "Possesses secondary expertise in..."). Highlight a secondary skill, architecture, or methodology requested in the JD that they possess.
                            - SUMMARY5 (The Closer): NO PRONOUNS. Start directly with a verb (e.g., "Will deliver immediate value by..."). State what specific value they will deliver on Day 1 based on their past execution track record.

                            CRITICAL ATS HACK: Across these 5 bullets, you MUST seamlessly embed exact phrases from the Job Description to maximize the Fieldglass match score.

                            ========================
                            STYLE & TONE RULES (STRICT)
                            ========================
                            - TELEGRAPHIC STYLE: Bullet 1 MUST use the candidate's name. Bullets 2, 3, 4, and 5 MUST strictly drop all pronouns (he/she/they) and names, starting directly with an action verb.
                            - Write like a human pitching to a colleague. Confident, direct, and factual.
                            - TECH MATCHING: Strictly align the tools you highlight with the JD. 
                            - LOCATION NEUTRAL: Never mention the physical location (e.g., Reston, VA, onsite, hybrid) in the summary.
                            - LEADERSHIP VERBS: Use high-authority active verbs (e.g., 'Engineered', 'Optimized', 'Architected').
                            - Do NOT use generic filler: "strong background," "highly experienced," "positions them uniquely."
                            - Do NOT repeat or restate the job description.

                            ========================
                            SKILLS SECTION (INTELLIGENT SEMANTIC MATCHING)
                            ========================
                            - You are evaluating the candidate against the exact skills listed in the "Skills Matrix:" input.
                            - INTELLIGENT MATCHING: You MUST look for semantic equivalents, synonyms, or related frameworks in the resume (e.g., if the matrix asks for "React.js", you MUST credit the candidate if they have "React Native", "React", etc.).
                            - Output the EXACT skill name as it is written in the Skills Matrix input, but calculate the Years and Companies based on your intelligent matching of their resume.
                            - If the provided Skills Matrix only lists 4 skills, you MUST ONLY output those 4 skills and leave SKILL5 through SKILL8 completely empty ("").
                            - You MUST extract and populate the 'Years' and 'Companies' fields for EVERY single skill you evaluate. Do not leave 'Years' or 'Companies' blank if the candidate has experience with the tool/concept.

                            Format: "Skill Area"
                            Years format: "X+ years" ONLY. (CRITICAL: DO NOT add ", current" or any years/dates to this string. Only output the number of years).
                            Companies format: "Company A, Company B" (Extract exact company names from the resume where the skill was used)

                            ========================
                            🚨 ZERO-TOLERANCE HALLUCINATION RULES (CRITICAL)
                            ========================
                            1. THE RESUME IS THE ONLY SOURCE OF TRUTH. You are strictly forbidden from copying a skill, tool, or technology from the Job Description and assigning it to the candidate unless it physically appears in their Resume.
                            2. DO NOT INFER OR ASSUME. 
                            3. DO NOT INFLATE TO MATCH THE JD. If the candidate lacks a requested skill, omit it completely.

                            ========================
                            YEARS ACCURACY RULES (REALISTIC RECRUITER MODE)
                            ========================
                            - Use July 2026 as the current date.
                            - STRICT MATH (DATE-DRIVEN ONLY): Calculate years strictly based on the earliest chronological date provided in the 'Professional Experience'. Round DOWN to the nearest whole year.

                            ========================
                            OUTPUT FORMAT (STRICT)
                            ========================
                            Return ONLY valid JSON. If less than 8 skills are found, leave the remaining skill strings empty ("").

                            {{
                              "SUMMARY1": "Bullet 1...",
                              "SUMMARY2": "Bullet 2...",
                              "SUMMARY3": "Bullet 3...",
                              "SUMMARY4": "Bullet 4...",
                              "SUMMARY5": "Bullet 5...",
                              "SKILL1": "Skill Area",
                              "YEARS1": "X+ years",
                              "SKILL1COMPANIES": "Company A, Company B",
                              "SKILL2": "Skill Area",
                              "YEARS2": "X+ years",
                              "SKILL2COMPANIES": "Company C",
                              "SKILL3": "",
                              "YEARS3": "",
                              "SKILL3COMPANIES": "",
                              "SKILL4": "",
                              "YEARS4": "",
                              "SKILL4COMPANIES": "",
                              "SKILL5": "",
                              "YEARS5": "",
                              "SKILL5COMPANIES": "",
                              "SKILL6": "",
                              "YEARS6": "",
                              "SKILL6COMPANIES": "",
                              "SKILL7": "",
                              "YEARS7": "",
                              "SKILL7COMPANIES": "",
                              "SKILL8": "",
                              "YEARS8": "",
                              "SKILL8COMPANIES": ""
                            }}

                            ========================
                            INPUT DATA
                            ========================
                            Official Job Description:
                            {Job_Description}

                            Notes:
                            {Chat_Notes}

                            Skills Matrix:
                            {Skills_Matrix}

                            Resume:
                            {raw_text}
                            """
                            time.sleep(2) 
                            summary_data = {}
                            
                            for attempt in range(6):
                                current_model = models_to_try[0] if attempt == 0 else (models_to_try[1] if attempt in [1, 2] else (models_to_try[2] if attempt in [3, 4] else models_to_try[3]))
                                try:
                                    summary_response = client.models.generate_content(
                                        model=current_model,
                                        contents=summary_prompt,
                                        config=types.GenerateContentConfig(response_mime_type="application/json")
                                    )
                                    summary_data = repair_and_load_json(summary_response.text)
                                    break
                                except Exception as api_e:
                                    if "503" in str(api_e) and attempt < 5:
                                        time.sleep(2 ** ((attempt % 2) + 1))
                                        continue
                                    else:
                                        raise api_e

                            mapping["SUMMARY1"] = summary_data.get("SUMMARY1", "")
                            mapping["SUMMARY2"] = summary_data.get("SUMMARY2", "")
                            mapping["SUMMARY3"] = summary_data.get("SUMMARY3", "")
                            mapping["SUMMARY4"] = summary_data.get("SUMMARY4", "")
                            mapping["SUMMARY5"] = summary_data.get("SUMMARY5", "")
                            
                            for i in range(1, 9):
                                mapping[f"SKILL{i}"] = summary_data.get(f"SKILL{i}", "")
                                mapping[f"YEARS{i}"] = summary_data.get(f"YEARS{i}", "")
                                mapping[f"SKILL{i}COMPANIES"] = summary_data.get(f"SKILL{i}COMPANIES", "")
                                
                        except Exception as e:
                            st.warning(f"⚠️ Warning: Summary generation failed. Proceeding without it. ({e})")

                    out_file = f"Submission_Deloitte_{final_name.replace(' ', '_')}.docx"
                    # Pass resume_path to the new function so it can extract the raw XML elements
                    process_deloitte_doc(TEMPLATE_FILENAME, mapping, resume_path, out_file)
                    
                    with open(out_file, "rb") as file:
                        btn = st.download_button(
                            label="⬇️ Download Generated Document",
                            data=file,
                            file_name=out_file,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary"
                        )
                    
                    st.success(f"✅ Success! Document is ready for download.")

                    try:
                        os.remove(resume_path)
                    except:
                        pass

                except Exception as e:
                    st.error(f"❌ Process Failed: {str(e)}")


# ====================================================================
# --- MAIN ROUTING LOGIC (The Dropdown Page) ---
# ====================================================================

st.title("Resume Formatter")
st.markdown("Please select your client account below to access the customized formatter.")

client_selection = st.selectbox(
    "Select Client Account or Tool:", 
    ["Fannie Mae", "Freddie Mac", "Deloitte", "Peraton", "Capital One", "ADUSA", "CBRE", "BNSF", "Dallas Generic", "PDF to Word"]
)

st.divider()

if client_selection == "Fannie Mae":
    fannie_mae_app()
elif client_selection == "Freddie Mac":
    freddie_mac_app()
elif client_selection == "Deloitte":
    deloitte_app()
elif client_selection == "Peraton":
    peraton_app()
elif client_selection == "Capital One":
    capital_one_app()
elif client_selection == "ADUSA":
    adusa_app()
elif client_selection == "CBRE":
    cbre_app()
elif client_selection == "BNSF":
    bnsf_app()
elif client_selection == "Dallas Generic":
    dallas_generic_app()
elif client_selection == "PDF to Word":
    pdf_to_word_app()
